"""
DuckGrade Canvas Integration - PyQt6 Modern Interface
Complete implementation matching Tkinter canvas_gui.py structure

This version provides a professional PyQt6 interface that exactly matches
the functionality and layout of the original Tkinter Canvas GUI.
"""

import sys
from pathlib import Path
import threading
import json
import os
import mimetypes
from PyQt6.QtWidgets import (QApplication, QMainWindow, QTabWidget, QWidget, 
                            QVBoxLayout, QHBoxLayout, QGroupBox, QLabel, 
                            QLineEdit, QPushButton, QComboBox, QTextEdit,
                            QCheckBox, QProgressBar, QMessageBox, QFileDialog,
                            QSpacerItem, QSizePolicy, QInputDialog, QScrollArea,
                            QRadioButton, QButtonGroup, QSplitter, QSpinBox, QDoubleSpinBox,
                            QStackedWidget, QDialog, QFormLayout, QGridLayout,
                            QListWidget, QListWidgetItem, QTreeWidget, QTreeWidgetItem)
from PyQt6.QtCore import Qt, QSize, QEvent, QTimer, QObject, pyqtSignal, QUrl, QMimeData
from PyQt6.QtGui import QIcon, QFont, QDragEnterEvent, QDropEvent

# Token counting for cost estimation
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    print("Note: tiktoken not available. Token counting will use approximations.")
    TIKTOKEN_AVAILABLE = False

# Try to import QWebEngineView for PDF/document rendering
try:
    from PyQt6.QtWebEngineWidgets import QWebEngineView
    WEB_ENGINE_AVAILABLE = True
except ImportError:
    print("Note: QWebEngineView not available. Will use text extraction for all documents.")
    QWebEngineView = None
    WEB_ENGINE_AVAILABLE = False

# Import Canvas integration components
try:
    from canvas_integration import CanvasAPI, TwoStepCanvasGrading
except ImportError:
    print("Warning: Canvas integration module not found. Some functionality may be limited.")
    CanvasAPI = None
    TwoStepCanvasGrading = None

# Additional imports for student anonymization
import datetime
import re


class StudentAnonymizer:
    """Handles student identity anonymization for privacy protection"""
    
    def __init__(self):
        self.name_map = {}  # Maps anonymous IDs to real names
        self.reverse_map = {}  # Maps real names to anonymous IDs
        
    def anonymize_name(self, real_name: str, user_id: int) -> str:
        """Convert real name to anonymous identifier"""
        if real_name in self.reverse_map:
            return self.reverse_map[real_name]
        
        # Create anonymous ID
        anon_id = f"Student_{len(self.name_map) + 1:03d}"
        
        # Store mappings
        self.name_map[anon_id] = {
            'real_name': real_name,
            'user_id': user_id
        }
        self.reverse_map[real_name] = anon_id
        
        return anon_id
    
    def get_real_name(self, anon_id: str) -> str:
        """Get real name from anonymous ID"""
        return self.name_map.get(anon_id, {}).get('real_name', anon_id)
    
    def get_user_id(self, anon_id: str) -> int:
        """Get Canvas user ID from anonymous ID"""
        return self.name_map.get(anon_id, {}).get('user_id')
    
    def anonymize_text(self, text: str) -> str:
        """Remove/anonymize any student names that might appear in text"""
        # This could be enhanced to detect and replace names in paper content
        # For now, we'll focus on file names and basic patterns
        for real_name, anon_id in self.reverse_map.items():
            # Replace name variations (case insensitive)
            text = re.sub(re.escape(real_name), anon_id, text, flags=re.IGNORECASE)
            
            # Also replace common name patterns
            first_name = real_name.split()[0] if ' ' in real_name else real_name
            text = re.sub(re.escape(first_name), anon_id.split('_')[1], text, flags=re.IGNORECASE)
        
        return text
    
    def save_mapping(self, file_path: str):
        """Save name mapping to file for later use"""
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump({
                'name_map': self.name_map,
                'reverse_map': self.reverse_map,
                'created_at': datetime.datetime.now().isoformat()
            }, f, indent=2)
    
    def load_mapping(self, file_path: str):
        """Load name mapping from file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.name_map = data.get('name_map', {})
                self.reverse_map = data.get('reverse_map', {})
        except Exception as e:
            print(f"Warning: Could not load name mapping: {e}")

# Models that do not support file uploads (static list of older models)
MODELS_WITHOUT_FILE_SUPPORT = [
    'gpt-3.5-turbo',
    'gpt-3.5-turbo-1106', 
    'gpt-3.5-turbo-0125',
    'gpt-3.5-turbo-16k',
    'gpt-3.5-turbo-instruct',
    'text-davinci-003',
    'text-davinci-002',
    'davinci-002',
    'babbage-002'
]

# Import existing configuration system
try:
    from duckworks_core import DuckWorksConfig
except ImportError:
    # Fallback configuration class if not available
    class DuckWorksConfig:
        def __init__(self):
            self.config = {}
        def set(self, key, value):
            self.config[key] = value
        def get(self, key, default=None):
            return self.config.get(key, default)


class DownloadOnlyWorker(QObject):
    """Worker class for running Step 1 (Download Only) in background thread with proper signals"""
    progress_updated = pyqtSignal(int, str)  # percent, description
    log_message = pyqtSignal(str)  # message
    completed = pyqtSignal(dict)  # results
    error_occurred = pyqtSignal(str)  # error message
    
    def __init__(self, canvas_api, course_id, course_name, assignment_id, assignment_name, 
                 use_canvas_rubric, rubric_path, instructor_config_path, openai_key, selected_model,
                 course_materials_files=None, course_materials_instructions=""):
        super().__init__()
        self.canvas_api = canvas_api
        self.course_id = course_id
        self.course_name = course_name
        self.assignment_id = assignment_id
        self.assignment_name = assignment_name
        self.use_canvas_rubric = use_canvas_rubric
        self.rubric_path = rubric_path
        self.instructor_config_path = instructor_config_path
        self.openai_key = openai_key
        self.selected_model = selected_model
        self.course_materials_files = course_materials_files or []
        self.course_materials_instructions = course_materials_instructions
    
    def run(self):
        """Run Step 1 (Download Only) process"""
        try:
            # Log start of Step 1
            self.progress_updated.emit(1, "Initializing download...")
            self.log_message.emit("Starting Step 1: Download Submissions Only")
            self.log_message.emit(f"Course: {self.course_name}")
            self.log_message.emit(f"Assignment: {self.assignment_name}")
            
            if self.use_canvas_rubric:
                self.log_message.emit("Will download Canvas rubric automatically")
            else:
                self.log_message.emit(f"Using local rubric: {Path(self.rubric_path).name}")
            
            if self.instructor_config_path:
                self.log_message.emit(f"Using instructor config: {Path(self.instructor_config_path).name}")
            else:
                self.log_message.emit("No instructor config - will use default for Step 2")
            
            self.log_message.emit("Student names will be anonymized for file organization")
            
            # Check if we have the TwoStepCanvasGrading class
            if TwoStepCanvasGrading is None:
                raise Exception("Canvas integration module not available. Please ensure canvas_integration.py is present.")
            
            # Initialize grading agent (needed for TwoStepCanvasGrading constructor, but not used in download-only)
            try:
                from grading_agent import GradingAgent
                grading_agent = GradingAgent(self.openai_key, model=self.selected_model)
                two_step_grading = TwoStepCanvasGrading(self.canvas_api, grading_agent)
            except ImportError:
                raise Exception("Grading agent module not found. Please ensure grading_agent.py is available.")
            
            # Update progress
            self.progress_updated.emit(2, "Preparing to download submissions...")
            
            # Create progress callback that emits signals
            def progress_callback(percent, description):
                self.progress_updated.emit(percent, description)
                
            # Create log callback that emits signals
            def log_callback(message):
                self.log_message.emit(message)
            
            # Run Step 1 (Download Only)
            results = two_step_grading.step1_download_only(
                course_id=self.course_id,
                assignment_id=self.assignment_id,
                assignment_name=self.assignment_name,
                rubric_path=self.rubric_path,
                instructor_config_path=self.instructor_config_path,
                use_canvas_rubric=self.use_canvas_rubric,
                progress_callback=progress_callback,
                log_callback=log_callback
            )
            
            # Log results for debugging
            self.log_message.emit(f"Step 1 (Download) results: success={results.get('success', False)}")
            if results.get('success'):
                self.log_message.emit(f"Downloaded {results.get('submission_count', 0)} submissions")
                self.log_message.emit(f"Folder: {results.get('folder_path', 'Unknown')}")
            
            # Update metadata file with Canvas IDs
            if results['success']:
                metadata_file = os.path.join(results['folder_path'], "assignment_metadata.json")
                metadata = {
                    'course_id': self.course_id,
                    'assignment_id': self.assignment_id,
                    'course_name': self.course_name,
                    'assignment_name': self.assignment_name,
                    'download_timestamp': datetime.datetime.now().isoformat(),
                    'step1_complete': True,
                    'step2_ready': True
                }
                with open(metadata_file, 'w') as f:
                    json.dump(metadata, f, indent=2)
                self.log_message.emit(f"Saved metadata to: {metadata_file}")
            
            # Emit completion signal
            self.completed.emit(results)
            
        except Exception as e:
            error_msg = f"Step 1 (Download) failed: {str(e)}"
            print(f"ERROR: {error_msg}")
            import traceback
            traceback.print_exc()
            self.error_occurred.emit(error_msg)

class GradingWorker(QObject):
    """Worker class for running Step 2 (Grading) in background thread with proper signals"""
    progress_updated = pyqtSignal(int, str)  # percent, description
    log_message = pyqtSignal(str)  # message
    completed = pyqtSignal(dict)  # results
    error_occurred = pyqtSignal(str)  # error message
    
    def __init__(self, downloaded_submission_data, download_folder, rubric_path, instructor_config_path,
                 use_canvas_rubric, openai_key, selected_model, assignment_name, course_name, canvas_api,
                 course_materials_files=None, course_materials_instructions="", additional_grading_instructions=""):
        super().__init__()
        self.downloaded_submission_data = downloaded_submission_data
        self.download_folder = download_folder
        self.rubric_path = rubric_path
        self.instructor_config_path = instructor_config_path
        self.use_canvas_rubric = use_canvas_rubric
        self.openai_key = openai_key
        self.selected_model = selected_model
        self.assignment_name = assignment_name
        self.course_name = course_name
        self.canvas_api = canvas_api
        self.course_materials_files = course_materials_files or []
        self.course_materials_instructions = course_materials_instructions
        self.additional_grading_instructions = additional_grading_instructions
        
    def run(self):
        """Run Step 2 (Grading) process"""
        try:
            import datetime
            from pathlib import Path
            
            self.progress_updated.emit(10, "Initializing grading...")
            self.log_message.emit("🤖 Starting AI grading of downloaded submissions...")
            
            # Use existing download folder structure instead of creating new GRADED folder
            if not self.download_folder.exists():
                self.error_occurred.emit("Download folder not found")
                return
            
            # Use the results folder from the existing download structure
            grading_folder = self.download_folder / "results"
            grading_folder.mkdir(exist_ok=True)
            
            # Load student anonymizer mapping if available
            mapping_file = self.download_folder / "student_mapping.json"
            student_anonymizer = None
            if mapping_file.exists():
                try:
                    from grading_agent import StudentAnonymizer
                    student_anonymizer = StudentAnonymizer()
                    student_anonymizer.load_mapping(mapping_file)
                    self.log_message.emit(f"Loaded student mapping from: {mapping_file}")
                except Exception as e:
                    self.log_message.emit(f"⚠️ Could not load student mapping: {e}")
            
            self.log_message.emit(f"📁 Using results folder: {grading_folder}")
            
            # Load rubric and instructor config once before grading
            try:
                if self.use_canvas_rubric:
                    self.log_message.emit("📝 Using Canvas rubric")
                else:
                    self.log_message.emit(f"📝 Using local rubric: {Path(self.rubric_path).name if self.rubric_path else 'None'}")
                
                if self.instructor_config_path and Path(self.instructor_config_path).exists():
                    self.log_message.emit(f"⚙️ Using instructor config: {Path(self.instructor_config_path).name}")
                else:
                    self.log_message.emit("⚙️ No instructor config specified")
                    
            except Exception as config_error:
                self.error_occurred.emit(f"Error loading configuration: {config_error}")
                return
            
            # Grade each submission
            self.progress_updated.emit(15, "Grading submissions...")
            
            total_submissions = len(self.downloaded_submission_data)
            graded_submissions = []
            
            # Import grading agent
            try:
                from grading_agent import GradingAgent
                
                # Initialize grading agent
                grading_agent = GradingAgent(
                    api_key=self.openai_key,
                    model=self.selected_model
                )
                
                self.log_message.emit(f"🤖 Using model: {self.selected_model}")
                
                # Load course materials if provided (optional)
                if self.course_materials_files:
                    self.log_message.emit(f"📚 Loading {len(self.course_materials_files)} course material files...")
                    if self.course_materials_instructions:
                        self.log_message.emit(f"📝 Using custom instructions: {self.course_materials_instructions[:80]}..." if len(self.course_materials_instructions) > 80 else f"📝 Using custom instructions: {self.course_materials_instructions}")
                    try:
                        grading_agent.load_course_materials(
                            self.course_materials_files, 
                            self.course_materials_instructions
                        )
                        self.log_message.emit(f"✅ Course materials loaded successfully for grading context")
                    except Exception as e:
                        self.log_message.emit(f"⚠️ Warning: Could not load course materials: {e}")
                        self.log_message.emit("📚 Proceeding with grading without course materials")
                else:
                    self.log_message.emit("📚 No course materials provided")
                
                # Log additional grading instructions
                if self.additional_grading_instructions.strip():
                    self.log_message.emit(f"📝 Using additional grading instructions: {self.additional_grading_instructions.strip()[:80]}..." if len(self.additional_grading_instructions.strip()) > 80 else f"📝 Using additional grading instructions: {self.additional_grading_instructions.strip()}")
                else:
                    self.log_message.emit("📝 No additional grading instructions provided")
                
                # Load rubric
                if self.use_canvas_rubric:
                    # Look for saved Canvas rubric in download folder
                    canvas_rubric_path = self.download_folder / "canvas_rubric.json"
                    if canvas_rubric_path.exists():
                        grading_agent.load_rubric(str(canvas_rubric_path))
                        self.log_message.emit(f"📝 Loaded Canvas rubric: {canvas_rubric_path.name}")
                    else:
                        self.error_occurred.emit("Canvas rubric not found in download folder")
                        return
                else:
                    if self.rubric_path and Path(self.rubric_path).exists():
                        grading_agent.load_rubric(self.rubric_path)
                        self.log_message.emit(f"📝 Loaded local rubric: {Path(self.rubric_path).name}")
                    else:
                        self.error_occurred.emit("Local rubric file not found")
                        return
                
                # Load instructor config if provided
                if self.instructor_config_path and Path(self.instructor_config_path).exists():
                    grading_agent.load_instructor_config(self.instructor_config_path)
                    self.log_message.emit(f"⚙️ Loaded instructor config: {Path(self.instructor_config_path).name}")
                else:
                    self.log_message.emit("⚙️ No instructor config provided - using default settings")
                
            except ImportError as e:
                self.error_occurred.emit(f"Could not import grading agent: {e}")
                return
            except Exception as e:
                self.error_occurred.emit(f"Error initializing grading agent: {e}")
                return
            
            for i, submission_data in enumerate(self.downloaded_submission_data):
                try:
                    progress = 15 + int((i / total_submissions) * 75)  # 15% to 90%
                    student_name = submission_data.get('name', f'Student_{i}')
                    
                    # Debug: Log submission data structure
                    submission_files = submission_data.get('files', [])
                    file_count = len(submission_files)
                    self.log_message.emit(f"🔍 Processing {student_name}: {file_count} file(s) detected")
                    
                    self.progress_updated.emit(progress, f"Grading {student_name}...")
                    self.log_message.emit(f"📝 Grading submission {i+1}/{total_submissions}: {student_name}")
                    
                    # Grade the submission using the grading agent
                    try:
                        # Get submission files (multi-file support)
                        submission_files = submission_data.get('files', [])
                        
                        # Multi-file vs Single-file submission routing
                        if len(submission_files) > 1:
                            # Multi-file submission - prepare structured data for grade_paper
                            file_names = [os.path.basename(f) for f in submission_files if f]
                            self.log_message.emit(f"📁 Multi-file submission detected: {len(submission_files)} files for {student_name}")
                            self.log_message.emit(f"   Files: {', '.join(file_names[:3])}{'...' if len(file_names) > 3 else ''}")
                            
                            # Create structured files data for grade_paper multi-file support
                            files_structured = []
                            for file_path in submission_files:
                                if file_path and os.path.exists(file_path):
                                    filename = os.path.basename(file_path)
                                    files_structured.append({
                                        'filename': filename,
                                        'file_path': file_path,
                                        'content': submission_data.get('content', '')  # Use extracted content as fallback
                                    })
                            
                            # Use enhanced grade_paper with files array - it handles multi-file internally
                            submission_file_data = {
                                'name': student_name,
                                'content': submission_data.get('content', ''),
                                'files': files_structured  # This triggers multi-file processing in grade_paper
                            }
                            grading_result = grading_agent.grade_paper(submission_file_data, self.additional_grading_instructions.strip())
                        elif len(submission_files) == 1:
                            # Single file submission - use existing grade_paper method
                            submission_file_data = {
                                'name': student_name,
                                'content': submission_data.get('content', ''),
                                'file_path': submission_files[0]
                            }
                            file_name = os.path.basename(submission_files[0]) if submission_files[0] else 'Unknown'
                            self.log_message.emit(f"📄 Single-file submission for {student_name}: {file_name}")
                            grading_result = grading_agent.grade_paper(submission_file_data, self.additional_grading_instructions.strip())
                        else:
                            # No files found - fallback to legacy content-based grading
                            submission_file_data = {
                                'name': student_name,
                                'content': submission_data.get('content', ''),
                                'file_path': submission_data.get('file_path', '')
                            }
                            self.log_message.emit(f"� Text-only submission for {student_name} (no files detected)")
                            grading_result = grading_agent.grade_paper(submission_file_data, self.additional_grading_instructions.strip())
                        
                        # Create a graded submission entry with actual results
                        graded_submission = {
                            'name': student_name,
                            'id': submission_data.get('id', ''),
                            'status': 'graded',
                            'score': grading_result.get('overall_score', grading_result.get('total_score', 0)),
                            'feedback': grading_result.get('overall_feedback', grading_result.get('feedback', '')),
                            'detailed_scores': grading_result.get('detailed_scores', {}),
                            'grading_result': grading_result
                        }
                        graded_submissions.append(graded_submission)
                        
                        self.log_message.emit(f"✓ Completed grading for {student_name} - Score: {grading_result.get('overall_score', grading_result.get('total_score', 'N/A'))}")
                        
                    except Exception as grading_error:
                        self.log_message.emit(f"❌ Grading failed for {student_name}: {grading_error}")
                        
                        # Create entry for failed grading
                        graded_submission = {
                            'name': student_name,
                            'id': submission_data.get('id', ''),
                            'status': 'error',
                            'error': str(grading_error)
                        }
                        graded_submissions.append(graded_submission)
                        continue
                    
                except Exception as submission_error:
                    error_msg = f"Error grading {student_name}: {submission_error}"
                    self.log_message.emit(f"❌ {error_msg}")
                    print(f"Submission error: {submission_error}")
                    continue
            
            # Complete grading process
            self.progress_updated.emit(95, "Finalizing results...")
            
            # Save graded results to JSON for review tab
            graded_results_file = grading_folder / "graded_results.json"
            try:
                import json
                with open(graded_results_file, 'w', encoding='utf-8') as f:
                    json.dump(graded_submissions, f, indent=2, ensure_ascii=False)
                self.log_message.emit(f"💾 Graded results saved to: {graded_results_file.name}")
            except Exception as save_error:
                self.log_message.emit(f"⚠️ Warning: Could not save graded results: {save_error}")
            
            # Store graded results in memory for review tab
            self.graded_results = graded_submissions
            self.log_message.emit(f"📝 Graded results stored in memory for review")
            
            # Create summary report
            summary_file = grading_folder / "grading_summary.txt"
            summary_content = (
                f"Grading Summary\n"
                f"===============\n"
                f"Assignment: {self.assignment_name}\n"
                f"Course: {self.course_name}\n"
                f"Graded: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"Model: {self.selected_model}\n"
                f"Total Submissions: {total_submissions}\n"
                f"Successfully Graded: {len(graded_submissions)}\n"
                f"Errors: {total_submissions - len(graded_submissions)}\n\n"
                f"Graded Students:\n"
            )
            
            for graded_sub in graded_submissions:
                summary_content += f"- {graded_sub['name']}\n"
            
            summary_file.write_text(summary_content, encoding='utf-8')
            
            self.progress_updated.emit(100, "Grading complete")
            
            self.log_message.emit(f"✅ Step 2 completed successfully!")
            self.log_message.emit(f"📁 Grading results saved to: {grading_folder}")
            self.log_message.emit(f"📊 {len(graded_submissions)}/{total_submissions} submissions graded successfully")
            self.log_message.emit(f"🔒 Student identities remain anonymized in file system")
            
            # Create Excel review spreadsheet
            print("DEBUG: *** ENTERING SPREADSHEET CREATION SECTION ***")
            print(f"DEBUG: About to try importing pandas...")
            excel_path = None
            try:
                print("DEBUG: Inside try block for Excel creation")
                import pandas as pd
                print("DEBUG: pandas imported successfully")
                
                # Filter successful submissions for Excel export
                successful_submissions = [sub for sub in graded_submissions if sub.get('status') == 'graded']
                
                print(f"DEBUG: Total graded_submissions: {len(graded_submissions)}")
                print(f"DEBUG: Successful submissions (status='graded'): {len(successful_submissions)}")
                if graded_submissions:
                    print(f"DEBUG: Sample submission status: {graded_submissions[0].get('status', 'NO_STATUS')}")
                    print(f"DEBUG: Sample submission keys: {list(graded_submissions[0].keys())}")
                    print(f"DEBUG: Sample submission score (old): {graded_submissions[0].get('score', 'NO_SCORE')}")
                    print(f"DEBUG: Sample submission score (new): {graded_submissions[0].get('grading_result', {}).get('overall_score', 'NO_OVERALL_SCORE')}")
                    print(f"DEBUG: Sample submission feedback (old): {graded_submissions[0].get('feedback', 'NO_FEEDBACK')[:50]}...")
                    print(f"DEBUG: Sample submission feedback (new): {graded_submissions[0].get('grading_result', {}).get('overall_feedback', 'NO_OVERALL_FEEDBACK')[:50]}...")
                    print(f"DEBUG: Sample submission name: {graded_submissions[0].get('name', 'NO_NAME')}")
                
                if successful_submissions:
                    # Load student mapping for real names and user IDs
                    student_mapping = {}
                    try:
                        import json
                        print(f"DEBUG: Looking for student mapping in grading_folder: {grading_folder}")
                        # The mapping files are in the parent folder of results, not in results itself
                        root_folder = grading_folder.parent
                        print(f"DEBUG: Looking in root folder instead: {root_folder}")
                        mapping_file = root_folder / "student_mapping.json"
                        print(f"DEBUG: Full mapping file path: {mapping_file}")
                        print(f"DEBUG: Mapping file exists: {mapping_file.exists()}")
                        if mapping_file.exists():
                            with open(mapping_file, 'r', encoding='utf-8') as f:
                                mapping_data = json.load(f)
                            name_map = mapping_data.get('name_map', mapping_data) if isinstance(mapping_data, dict) else mapping_data
                            for anon_name, data in name_map.items():
                                if isinstance(data, dict) and 'real_name' in data:
                                    student_mapping[anon_name] = {
                                        'real_name': data['real_name'],
                                        'user_id': data.get('user_id', '')
                                    }
                        else:
                            print("DEBUG: student_mapping.json not found, trying fallback")
                            # Fallback: try loading from submission_data.json in root folder
                            submission_data_file = root_folder / "submission_data.json"
                            print(f"DEBUG: Fallback file path: {submission_data_file}")
                            print(f"DEBUG: Fallback file exists: {submission_data_file.exists()}")
                            if submission_data_file.exists():
                                with open(submission_data_file, 'r', encoding='utf-8') as f:
                                    submission_data = json.load(f)
                                for sub in submission_data:
                                    anon_name = sub.get('name', '')
                                    student_mapping[anon_name] = {
                                        'real_name': sub.get('real_name', anon_name),
                                        'user_id': sub.get('user_id', '')
                                    }
                        print(f"DEBUG: Loaded student mapping with {len(student_mapping)} entries for spreadsheet")
                    except Exception as mapping_error:
                        print(f"DEBUG: Could not load student mapping: {mapping_error}")
                        student_mapping = {}
                    
                    # Prepare data for Excel using proper format to match working example
                    excel_data = []
                    for sub in successful_submissions:
                        # Get anonymized name from the submission
                        anon_name = sub.get('name', '')
                        
                        # Look up real name and user ID from mapping
                        mapping_info = student_mapping.get(anon_name, {})
                        real_name = mapping_info.get('real_name', anon_name)
                        user_id = mapping_info.get('user_id', '')
                        
                        # Extract grading data properly
                        score = sub.get('score', sub.get('grading_result', {}).get('overall_score', 0))
                        max_score = sub.get('max_score', 
                                           sub.get('grading_result', {}).get('max_possible_score', 
                                           sub.get('grading_result', {}).get('max_score', 100)))  # Look in grading_result for max_possible_score first
                        percentage = (score / max_score * 100) if max_score > 0 else 0
                        
                        # Format AI_Comments with rubric criteria breakdown
                        ai_comments = self.format_rubric_comments(sub.get('grading_result', {}))
                        
                        row = {
                            'Student_Name': real_name,  # Use real name from mapping
                            'Canvas_User_ID': user_id,  # Canvas user ID from mapping
                            'AI_Score': score,
                            'Max_Score': max_score,
                            'AI_Percentage': f"{percentage:.1f}%",
                            'Final_Grade': f"{percentage:.1f}%",  # Editable in review
                            'AI_Comments': ai_comments,
                            'Final_Comments': ai_comments,  # Copy formatted comments for editing
                            'Notes': '',  # For manual notes
                            'Upload_Status': 'PENDING'  # Status for Canvas upload
                        }
                        excel_data.append(row)
                    
                    # Create DataFrame and save to Excel
                    df = pd.DataFrame(excel_data)
                    excel_filename = f"{self.assignment_name.replace(' ', '_')}_REVIEW.xlsx"
                    excel_path = grading_folder / excel_filename
                    
                    df.to_excel(excel_path, index=False, engine='openpyxl')
                    
                    self.log_message.emit(f"📊 Review spreadsheet created: {excel_filename}")
                    
                else:
                    self.log_message.emit("⚠️ No successful submissions to export to spreadsheet")
                    
            except ImportError as import_error:
                print(f"DEBUG: ImportError in Excel creation: {import_error}")
                self.log_message.emit("⚠️ pandas or openpyxl not available for Excel export")
            except Exception as excel_error:
                print(f"DEBUG: Exception in Excel creation: {excel_error}")
                print(f"DEBUG: Exception type: {type(excel_error)}")
                import traceback
                print(f"DEBUG: Traceback: {traceback.format_exc()}")
                self.log_message.emit(f"⚠️ Could not create Excel file: {excel_error}")
            
            # Prepare results
            results = {
                'success': True,
                'graded_submissions': graded_submissions,
                'grading_folder': str(grading_folder),
                'total_submissions': total_submissions,
                'successful_submissions': len([sub for sub in graded_submissions if sub.get('status') == 'graded']),
                'review_file': str(excel_path) if excel_path else None
            }
            
            # Emit completion signal
            self.completed.emit(results)
            
        except Exception as e:
            error_msg = f"Step 2 (Grading) failed: {str(e)}"
            print(f"ERROR: {error_msg}")
            import traceback
            traceback.print_exc()
            self.error_occurred.emit(error_msg)
    
    def format_rubric_comments(self, grading_result):
        """Format rubric comments in the same style as the working example"""
        if not grading_result or 'criteria_scores' not in grading_result:
            # Fallback to overall feedback if no criteria scores
            return grading_result.get('overall_feedback', '')
        
        criteria_scores = grading_result.get('criteria_scores', {})
        formatted_comments = []
        
        for criterion_name, criterion_data in criteria_scores.items():
            score = criterion_data.get('score', 0)
            max_score = criterion_data.get('max_score', 0)  
            feedback = criterion_data.get('feedback', '')
            
            # Format: "Criterion: score/max_score\n  - feedback"
            criterion_line = f"{criterion_name}: {score}/{max_score}"
            if feedback.strip():
                # Add indented feedback
                feedback_lines = feedback.strip().split('\n')
                feedback_formatted = '\n'.join([f"  - {line}" if i == 0 else f"    {line}" for i, line in enumerate(feedback_lines)])
                criterion_line += f"\n{feedback_formatted}"
            
            formatted_comments.append(criterion_line)
        
        # Join all criteria with double newlines
        result = '\n\n'.join(formatted_comments)
        
        # Add overall feedback at the end if present
        overall_feedback = grading_result.get('overall_feedback', '')
        if overall_feedback and overall_feedback.strip():
            result += f"\n\nOverall: {overall_feedback.strip()}"
        
        return result

class Step1Worker(QObject):
    """Worker class for running Step 1 in background thread with proper signals"""
    progress_updated = pyqtSignal(int, str)  # percent, description
    log_message = pyqtSignal(str)  # message
    completed = pyqtSignal(dict)  # results
    error_occurred = pyqtSignal(str)  # error message
    
    def __init__(self, canvas_api, course_id, course_name, assignment_id, assignment_name, 
                 use_canvas_rubric, rubric_path, instructor_config_path, openai_key, selected_model,
                 course_materials_files=None, course_materials_instructions=""):
        super().__init__()
        self.canvas_api = canvas_api
        self.course_id = course_id
        self.course_name = course_name
        self.assignment_id = assignment_id
        self.assignment_name = assignment_name
        self.use_canvas_rubric = use_canvas_rubric
        self.rubric_path = rubric_path
        self.instructor_config_path = instructor_config_path
        self.openai_key = openai_key
        self.selected_model = selected_model
        self.course_materials_files = course_materials_files or []
        self.course_materials_instructions = course_materials_instructions
    
    def run(self):
        """Run Step 1 process"""
        try:
            # Log start of Step 1
            self.progress_updated.emit(10, "Initializing...")
            self.log_message.emit("Starting Step 1: Download and Grade with Privacy Protection")
            self.log_message.emit(f"Course: {self.course_name}")
            self.log_message.emit(f"Assignment: {self.assignment_name}")
            
            if self.use_canvas_rubric:
                self.log_message.emit("Will download Canvas rubric automatically")
            else:
                self.log_message.emit(f"Using local rubric: {Path(self.rubric_path).name}")
            
            if self.instructor_config_path:
                self.log_message.emit(f"Using instructor config: {Path(self.instructor_config_path).name}")
            else:
                self.log_message.emit("Using default grading style")
            
            self.log_message.emit("Student names will be anonymized for AI processing")
            
            # Check if we have the TwoStepCanvasGrading class
            if TwoStepCanvasGrading is None:
                raise Exception("Canvas integration module not available. Please ensure canvas_integration.py is present.")
            
            # Initialize grading agent
            try:
                from grading_agent import GradingAgent
                grading_agent = GradingAgent(self.openai_key, model=self.selected_model)
                two_step_grading = TwoStepCanvasGrading(self.canvas_api, grading_agent)
            except ImportError:
                raise Exception("Grading agent module not found. Please ensure grading_agent.py is available.")
            
            # Update progress
            self.progress_updated.emit(30, "Downloading submissions...")
            
            # Create progress callback that emits signals
            def progress_callback(percent, description):
                self.progress_updated.emit(percent, description)
                
            # Create log callback that emits signals
            def log_callback(message):
                self.log_message.emit(message)
            
            # Run Step 1
            results = two_step_grading.step1_download_and_grade(
                course_id=self.course_id,
                assignment_id=self.assignment_id,
                assignment_name=self.assignment_name,
                rubric_path=self.rubric_path,
                instructor_config_path=self.instructor_config_path,
                use_canvas_rubric=self.use_canvas_rubric,
                progress_callback=progress_callback,
                log_callback=log_callback
            )
            
            # Log results for debugging
            self.log_message.emit(f"Step 1 results: success={results.get('success', False)}")
            if 'review_file' in results:
                self.log_message.emit(f"Review file reported: {results['review_file']}")
                if os.path.exists(results['review_file']):
                    self.log_message.emit("✓ Review file exists on disk")
                else:
                    self.log_message.emit("❌ Review file not found on disk")
            else:
                self.log_message.emit("❌ No review_file in results")
            
            # Update metadata file with Canvas IDs
            if results['success']:
                metadata_file = os.path.join(results['folder_path'], "assignment_metadata.json")
                if os.path.exists(metadata_file):
                    with open(metadata_file, 'r') as f:
                        metadata = json.load(f)
                    metadata['course_id'] = self.course_id
                    metadata['assignment_id'] = self.assignment_id
                    with open(metadata_file, 'w') as f:
                        json.dump(metadata, f, indent=2)
            
            # Emit completion signal
            self.completed.emit(results)
            
        except Exception as e:
            self.error_occurred.emit(str(e))


class ScrollFriendlyComboBox(QComboBox):
    """
    Custom ComboBox that doesn't capture wheel events when not focused.
    This allows users to scroll the parent widget without interruption.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self._wheel_enabled = False
    
    def wheelEvent(self, event):
        # Only process wheel events if explicitly enabled (after clicking)
        if self._wheel_enabled and self.hasFocus():
            super().wheelEvent(event)
        else:
            # Ignore the wheel event and let it propagate to parent
            event.ignore()
    
    def mousePressEvent(self, event):
        # Enable wheel events when combo box is clicked
        self._wheel_enabled = True
        super().mousePressEvent(event)
    
    def focusOutEvent(self, event):
        # Disable wheel events when focus is lost
        self._wheel_enabled = False
        super().focusOutEvent(event)


class DuckGradeCanvasGUI(QMainWindow):
    """
    Main Canvas GUI window matching the exact structure of canvas_gui.py
    
    Features 4 tabs matching Tkinter version:
    - 🔗 Canvas Connection (API settings, privacy, status)
    - 🎯 1-2-3 Grading (workflow with cost transparency and review)
    - ⚡ Single-Step Grading (direct grading without review)
    - 📊 Results (session history and statistics)
    """
    
    def __init__(self):
        super().__init__()
        # Cache for decrypted credentials to avoid repeated password prompts
        self._cached_openai_key = None
        self._cached_canvas_url = None
        self._cached_canvas_token = None
        
        # Initialize submission token tracking
        self.submission_tokens = 0
        self.submission_cost = 0.0
        
        # Load general configuration
        self.general_config = self.load_general_config()
        self.init_ui()
        self.setup_window()
    
    def load_general_config(self):
        """Load general configuration from config/general_config.json"""
        config_path = Path("config/general_config.json")
        default_config = {
            "ui_options": {
                "show_test_mode_button": False
            }
        }
        
        try:
            if config_path.exists():
                with open(config_path, 'r') as f:
                    config = json.load(f)
                # Merge with defaults to ensure all keys exist
                for key, value in default_config.items():
                    if key not in config:
                        config[key] = value
                    elif isinstance(value, dict) and isinstance(config[key], dict):
                        for sub_key, sub_value in value.items():
                            if sub_key not in config[key]:
                                config[key][sub_key] = sub_value
                return config
            else:
                # Create config file with defaults
                config_path.parent.mkdir(exist_ok=True)
                with open(config_path, 'w') as f:
                    json.dump(default_config, f, indent=4)
                return default_config
        except Exception as e:
            print(f"Warning: Could not load general config: {e}")
            return default_config
    
    def init_ui(self):
        """Initialize the user interface"""
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        
        # Create tab widget
        self.tab_widget = QTabWidget()
        
        # Create all tabs but only add Home initially
        self.home_tab = self.create_home_tab()
        self.two_step_tab = self.create_two_step_tab()
        self.single_step_tab = self.create_single_step_tab()
        self.results_tab = self.create_results_tab()
        self.ducktest_two_step_tab = self.create_ducktest_two_step_tab()
        self.ducktest_one_step_tab = self.create_ducktest_one_step_tab()
        
        # Review Tab (initially not created, only created when needed)
        self.review_tab = None
        self.review_tab_visible = False
        
        # Initially show only Home tab
        self.tab_widget.addTab(self.home_tab, "🏠 Home")
        
        # Track current tool
        self.current_tool = None
        
        main_layout.addWidget(self.tab_widget)
        
        # Initialize default state - no tool selected initially
    
    def create_home_tab(self) -> QWidget:
        """Create Home tab with API connections and tool selection"""
        # Create main tab widget
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Create scroll area for the main content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Create content widget that will go inside scroll area
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        
        # OpenAI Configuration Group
        openai_group = QGroupBox("🤖 OpenAI Configuration")
        openai_layout = QVBoxLayout(openai_group)
        
        # API Key row
        key_layout = QHBoxLayout()
        key_layout.addWidget(QLabel("API Key:"))
        self.openai_key_entry = QLineEdit()
        self.openai_key_entry.setEchoMode(QLineEdit.EchoMode.Password)
        self.openai_key_entry.setPlaceholderText("Enter your OpenAI API key")
        self.openai_key_entry.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        key_layout.addWidget(self.openai_key_entry)
        
        openai_layout.addLayout(key_layout)
        
        # Model selection row
        model_layout = QHBoxLayout()
        model_layout.addWidget(QLabel("Model:"))
        
        self.model_combo = ScrollFriendlyComboBox()
        self.model_combo.addItem("gpt-4o-mini")
        self.model_combo.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.model_combo.currentIndexChanged.connect(self.on_model_changed)
        self.model_combo.currentTextChanged.connect(self.on_model_text_changed)
        
        # Style the ComboBox with custom arrow using PNG
        self.model_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
                min-width: 350px;
                padding-right: 25px;
            }
            QComboBox:hover {
                border-color: #adb5bd;
            }
            QComboBox:focus {
                border-color: #80bdff;
                outline: 0;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        
        model_layout.addWidget(self.model_combo)
        
        refresh_models_btn = QPushButton("Refresh Models")
        refresh_models_btn.clicked.connect(self.refresh_models)
        model_layout.addWidget(refresh_models_btn)
        
        openai_layout.addLayout(model_layout)
        layout.addWidget(openai_group)
        
        # Canvas Configuration Group
        canvas_group = QGroupBox("🎨 Canvas Configuration")
        canvas_layout = QVBoxLayout(canvas_group)
        
        # Canvas URL
        url_layout = QHBoxLayout()
        url_layout.addWidget(QLabel("Canvas URL:"))
        self.canvas_url_entry = QLineEdit()
        self.canvas_url_entry.setPlaceholderText("https://your-school.instructure.com")
        url_layout.addWidget(self.canvas_url_entry)
        canvas_layout.addLayout(url_layout)
        
        # API Token
        token_layout = QHBoxLayout()
        token_layout.addWidget(QLabel("API Token:"))
        self.canvas_token_entry = QLineEdit()
        self.canvas_token_entry.setEchoMode(QLineEdit.EchoMode.Password)
        self.canvas_token_entry.setPlaceholderText("Enter your Canvas API token")
        token_layout.addWidget(self.canvas_token_entry)
        canvas_layout.addLayout(token_layout)
        
        # Connection test and configuration buttons
        test_layout = QHBoxLayout()
        test_connection_btn = QPushButton("Connect")
        
        # Add network icon if available
        if Path("assets/network_outlined.png").exists():
            test_connection_btn.setIcon(QIcon("assets/network_outlined.png"))
            test_connection_btn.setIconSize(QSize(16, 16))
        
        test_connection_btn.clicked.connect(self.test_connection)
        test_layout.addWidget(test_connection_btn)
        
        load_config_btn = QPushButton("📂 Load Configuration")
        load_config_btn.clicked.connect(self.load_configuration)
        test_layout.addWidget(load_config_btn)
        
        save_config_btn = QPushButton("💾 Save Configuration")
        save_config_btn.clicked.connect(self.save_configuration)
        test_layout.addWidget(save_config_btn)
        test_layout.addStretch()
        
        canvas_layout.addLayout(test_layout)
        layout.addWidget(canvas_group)
        
        # Privacy and Safety Group
        privacy_group = QGroupBox("🔒 Privacy and Safety")
        privacy_layout = QVBoxLayout(privacy_group)
        
        self.anonymize_checkbox = QCheckBox("Anonymize student names in grading feedback")
        self.anonymize_checkbox.setChecked(True)
        privacy_layout.addWidget(self.anonymize_checkbox)
        
        self.backup_checkbox = QCheckBox("Create backup of submissions before grading")
        self.backup_checkbox.setChecked(True)
        privacy_layout.addWidget(self.backup_checkbox)
        
        layout.addWidget(privacy_group)
        
        # Connection Status Group
        status_group = QGroupBox("📊 Connection Status")
        status_layout = QVBoxLayout(status_group)
        
        self.openai_status = QLabel("🔴 OpenAI API not configured")
        self.openai_status.setStyleSheet("color: red; font-weight: bold;")
        status_layout.addWidget(self.openai_status)
        
        self.connection_status = QLabel("🔴 Not connected to Canvas")
        self.connection_status.setStyleSheet("color: red; font-weight: bold;")
        status_layout.addWidget(self.connection_status)
        
        # Test mode button for development/testing (conditionally shown)
        if self.general_config.get("ui_options", {}).get("show_test_mode_button", False):
            test_mode_layout = QHBoxLayout()
            test_mode_btn = QPushButton("🧪 Enable Test Mode")
            test_mode_btn.setToolTip("Enable grading buttons for testing without Canvas connection")
            test_mode_btn.clicked.connect(self.enable_test_mode)
            test_mode_btn.setFixedHeight(21)
            test_mode_layout.addWidget(test_mode_btn)
            test_mode_layout.addStretch()
            status_layout.addLayout(test_mode_layout)
        
        layout.addWidget(status_group)
        layout.addStretch()
        
        # Set the content widget in the scroll area
        scroll_area.setWidget(content_widget)
        
        # Add scroll area to the main tab layout
        tab_layout.addWidget(scroll_area)
        
        # Add tool selection buttons at the bottom
        tool_selection_group = QGroupBox("🛠️ DuckWorks Tool Suite")
        tool_layout = QHBoxLayout(tool_selection_group)
        
        # Create button group for exclusive selection
        self.tool_button_group = QButtonGroup()
        
        # DuckGrade button
        self.duckgrade_button = QPushButton()
        self.duckgrade_button.setText("DuckGrade")
        self.duckgrade_button.setToolTip("DuckGrade: Automated grading with AI-powered feedback, Canvas integration, and privacy protection")
        self.duckgrade_button.setCheckable(True)
        self.duckgrade_button.setMinimumHeight(80)
        self.duckgrade_button.setMinimumWidth(200)
        if Path("assets/mallard_icon.png").exists():
            self.duckgrade_button.setIcon(QIcon("assets/mallard_icon.png"))
            self.duckgrade_button.setIconSize(QSize(32, 32))
        self.duckgrade_button.clicked.connect(lambda: self.switch_to_tool("duckgrade"))
        self.tool_button_group.addButton(self.duckgrade_button)
        
        # DuckTest button
        self.ducktest_button = QPushButton()
        self.ducktest_button.setText("DuckTest")
        self.ducktest_button.setToolTip("DuckTest: Intelligent assessment creation with automated question generation and rubric design")
        self.ducktest_button.setCheckable(True)
        self.ducktest_button.setMinimumHeight(80)
        self.ducktest_button.setMinimumWidth(200)
        if Path("assets/duckling_icon.png").exists():
            self.ducktest_button.setIcon(QIcon("assets/duckling_icon.png"))
            self.ducktest_button.setIconSize(QSize(32, 32))
        self.ducktest_button.clicked.connect(lambda: self.switch_to_tool("ducktest"))
        self.tool_button_group.addButton(self.ducktest_button)
        
        # Style the tool buttons
        tool_button_style = """
            QPushButton {
                border: 2px solid #dee2e6;
                border-radius: 12px;
                background-color: #ffffff;
                padding: 15px;
                font-size: 11pt;
                font-weight: bold;
                text-align: center;
                color: #495057;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border-color: #007bff;
            }
            QPushButton:checked {
                background-color: #007bff;
                color: white;
                border-color: #0056b3;
            }
            QPushButton:pressed {
                background-color: #0056b3;
            }
            QPushButton:checked:hover {
                background-color: #0056b3;
            }
        """
        self.duckgrade_button.setStyleSheet(tool_button_style)
        self.ducktest_button.setStyleSheet(tool_button_style)
        
        # Add buttons to layout with stretch
        tool_layout.addStretch()
        tool_layout.addWidget(self.duckgrade_button)
        tool_layout.addWidget(self.ducktest_button)
        tool_layout.addStretch()
        
        # Add tool selection group to main tab layout
        tab_layout.addWidget(tool_selection_group)
        
        return tab
    
    def switch_to_tool(self, tool_name):
        """Switch to the specified tool and show its tabs"""
        # Clear existing tabs except Home (index 0)
        while self.tab_widget.count() > 1:
            self.tab_widget.removeTab(1)
        
        # Update current tool
        self.current_tool = tool_name
        
        if tool_name == "duckgrade":
            # Add DuckGrade tabs
            self.tab_widget.addTab(self.two_step_tab, "🎯 1-2-3 Grading")
            self.tab_widget.addTab(self.single_step_tab, "⚡ Single-Step Grading")
            self.tab_widget.addTab(self.results_tab, "📊 Results")
            
            # Update button states
            self.duckgrade_button.setChecked(True)
            self.ducktest_button.setChecked(False)
            
        elif tool_name == "ducktest":
            # Add DuckTest tabs
            self.tab_widget.addTab(self.ducktest_two_step_tab, "📝 Two-Step Assessment")
            self.tab_widget.addTab(self.ducktest_one_step_tab, "⚡ One-Step Assessment")
            
            # Update button states
            self.duckgrade_button.setChecked(False)
            self.ducktest_button.setChecked(True)
        
        # Switch to the first non-home tab if we just added some
        if self.tab_widget.count() > 1:
            self.tab_widget.setCurrentIndex(1)
    
    def create_two_step_tab(self) -> QWidget:
        """Create 1-2-3 Grading tab matching new workflow structure"""
        # Create main tab widget
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Create scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Create content widget that will go inside scroll area
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        
        # Workflow description
        workflow_group = QGroupBox("🎯 1-2-3 Grading Workflow")
        workflow_layout = QVBoxLayout(workflow_group)
        
        workflow_text = QLabel(
            "1-2-3 Grading provides transparent cost control and quality assurance:\n\n"
            "Step 1: Download submissions and calculate exact grading costs\n"
            "Step 2: AI grades submissions and saves results for review\n"
            "Step 3: Review results, make adjustments, then upload to Canvas\n\n"
            "This approach shows total costs upfront and allows verification before upload."
        )
        workflow_text.setStyleSheet("padding: 10px; background-color: #f8f9fa; border-radius: 4px;")
        workflow_layout.addWidget(workflow_text)
        
        layout.addWidget(workflow_group)
        
        # Assignment Selection Group
        assignment_group = QGroupBox("📚 Assignment Selection")
        assignment_layout = QVBoxLayout(assignment_group)
        
        # Course selection input
        course_layout = QHBoxLayout()
        course_layout.addWidget(QLabel("Course:"))
        
        self.course_id_combo = ScrollFriendlyComboBox()
        self.course_id_combo.addItem("No courses loaded")
        self.course_id_combo.setEditable(True)  # Allow manual entry
        self.course_id_combo.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        
        # Style the ComboBox with custom arrow using PNG (same as Model dropdown)
        self.course_id_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
                padding-right: 25px;
            }
            QComboBox:hover {
                border-color: #adb5bd;
            }
            QComboBox:focus {
                border-color: #80bdff;
                outline: 0;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        
        course_layout.addWidget(self.course_id_combo)
        
        # Connect course selection to assignment refresh
        self.course_id_combo.currentIndexChanged.connect(self.on_course_selected)
        
        assignment_layout.addLayout(course_layout)
        
        # Assignment dropdown
        assignment_select_layout = QHBoxLayout()
        assignment_select_layout.addWidget(QLabel("Select Assignment:"))
        self.assignment_combo = ScrollFriendlyComboBox()
        self.assignment_combo.addItem("No assignments loaded")
        self.assignment_combo.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        
        # Style the assignment ComboBox with custom arrow (same as Model and Course dropdowns)
        self.assignment_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
                padding-right: 25px;
            }
            QComboBox:hover {
                border-color: #adb5bd;
            }
            QComboBox:focus {
                border-color: #80bdff;
                outline: 0;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        
        assignment_select_layout.addWidget(self.assignment_combo)
        
        refresh_assignments_btn = QPushButton("↻ Refresh")
        refresh_assignments_btn.clicked.connect(self.refresh_assignments)
        assignment_select_layout.addWidget(refresh_assignments_btn)
        
        assignment_layout.addLayout(assignment_select_layout)
        
        # Assignment info (removed - no details functionality implemented)
        # self.assignment_info = QLabel("Select an assignment to view details")
        # self.assignment_info.setStyleSheet("padding: 10px; background-color: #f8f9fa; border-radius: 4px;")
        # assignment_layout.addWidget(self.assignment_info)
        
        layout.addWidget(assignment_group)
        
        # Grading Configuration Group
        config_group = QGroupBox("⚙️ Grading Configuration")
        config_layout = QVBoxLayout(config_group)
        
        # Rubric source selection
        rubric_source_layout = QHBoxLayout()
        rubric_source_layout.addWidget(QLabel("Rubric Source:"))
        
        # Create button group for radio buttons (ensures mutual exclusivity)
        self.rubric_source_group = QButtonGroup(self)
        
        self.local_rubric_radio = QRadioButton("Local File")
        self.local_rubric_radio.setChecked(True)  # Default to local file
        self.local_rubric_radio.toggled.connect(self.on_rubric_source_changed)
        self.rubric_source_group.addButton(self.local_rubric_radio, 0)
        rubric_source_layout.addWidget(self.local_rubric_radio)
        
        self.canvas_rubric_radio = QRadioButton("Canvas Rubric")
        self.canvas_rubric_radio.toggled.connect(self.on_rubric_source_changed)
        self.rubric_source_group.addButton(self.canvas_rubric_radio, 1)
        rubric_source_layout.addWidget(self.canvas_rubric_radio)
        
        rubric_source_layout.addStretch()
        config_layout.addLayout(rubric_source_layout)
        
        # Local rubric file selection (initially visible)
        self.local_rubric_widget = QWidget()
        local_rubric_layout = QHBoxLayout(self.local_rubric_widget)
        local_rubric_layout.setContentsMargins(0, 0, 0, 0)
        local_rubric_layout.addWidget(QLabel("Rubric File:"))
        self.rubric_path_entry = QLineEdit("rubrics/sample_rubric.json")
        self.rubric_path_entry.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        local_rubric_layout.addWidget(self.rubric_path_entry)
        
        browse_rubric_btn = QPushButton("Browse")
        browse_rubric_btn.setFixedHeight(21)
        browse_rubric_btn.clicked.connect(self.browse_rubric)
        local_rubric_layout.addWidget(browse_rubric_btn)
        
        config_layout.addWidget(self.local_rubric_widget)
        
        # Canvas rubric info (initially hidden)
        self.canvas_rubric_widget = QWidget()
        canvas_rubric_layout = QHBoxLayout(self.canvas_rubric_widget)
        canvas_rubric_layout.setContentsMargins(0, 0, 0, 0)
        canvas_rubric_info = QLabel("✓ Canvas rubric will be downloaded automatically from the selected assignment")
        canvas_rubric_info.setStyleSheet("color: #28a745; font-weight: bold; padding: 5px;")
        canvas_rubric_layout.addWidget(canvas_rubric_info)
        canvas_rubric_layout.addStretch()
        
        config_layout.addWidget(self.canvas_rubric_widget)
        self.canvas_rubric_widget.hide()  # Initially hidden
        
        # Instructor config
        instructor_layout = QHBoxLayout()
        instructor_layout.addWidget(QLabel("Instructor Config (Optional):"))
        self.instructor_config_entry = QLineEdit("config/grading_instructor_config.json")
        instructor_layout.addWidget(self.instructor_config_entry)
        
        browse_instructor_btn = QPushButton("Browse")
        browse_instructor_btn.clicked.connect(self.browse_instructor_config)
        instructor_layout.addWidget(browse_instructor_btn)
        
        new_instructor_btn = QPushButton("New")
        new_instructor_btn.setToolTip("Create a new instructor configuration with guided setup")
        new_instructor_btn.clicked.connect(self.create_new_instructor_config)
        instructor_layout.addWidget(new_instructor_btn)

        config_layout.addLayout(instructor_layout)
        
        # Additional Grading Instructions
        grading_instructions_layout = QVBoxLayout()
        grading_instructions_layout.addWidget(QLabel("📝 Additional Grading Instructions (Optional):"))
        
        self.additional_grading_instructions = QTextEdit()
        self.additional_grading_instructions.setMaximumHeight(90)
        self.additional_grading_instructions.setPlaceholderText(
            "Enter assignment-specific instructions for the AI grader\n\n"
            "Multi-file example: Grade the research paper and verify calculations match the Excel data"
        )
        self.additional_grading_instructions.setStyleSheet("""
            QTextEdit {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
            QTextEdit:focus {
                border-color: #80bdff;
                outline: 0;
            }
        """)
        
        grading_instructions_layout.addWidget(self.additional_grading_instructions)
        config_layout.addLayout(grading_instructions_layout)

        layout.addWidget(config_group)        # Step 1: Download Submissions Group (moved here for better workflow order)
        step1_group = QGroupBox("📥 Step 1: Download Submissions and Calculate Costs")
        step1_layout = QVBoxLayout(step1_group)
        
        step1_desc = QLabel("Download all submissions to calculate exact token usage and costs before grading")
        step1_layout.addWidget(step1_desc)
        
        step1_button_layout = QHBoxLayout()
        self.step1_download_button = QPushButton("📥 Start Step 1: Download Submissions")
        self.step1_download_button.setEnabled(False)
        self.step1_download_button.clicked.connect(self.start_step1_download)
        step1_button_layout.addWidget(self.step1_download_button)
        
        self.step1_download_status = QLabel("Ready")
        self.step1_download_status.setStyleSheet("color: blue;")
        self.step1_download_status.setMinimumWidth(120)  # Ensure enough space for "Downloading..."
        step1_button_layout.addWidget(self.step1_download_status)
        step1_button_layout.addStretch()
        
        step1_layout.addLayout(step1_button_layout)
        
        # Progress bar for step 1 download
        self.progress_step1_download = QProgressBar()
        step1_layout.addWidget(self.progress_step1_download)
        
        self.progress_desc_step1_download = QLabel("Ready")
        self.progress_desc_step1_download.setStyleSheet("color: blue; font-size: 10px;")
        step1_layout.addWidget(self.progress_desc_step1_download)
        
        layout.addWidget(step1_group)
        
        # Course Materials Group
        course_materials_group = QGroupBox("📚 Course Materials (Optional)")
        course_materials_layout = QVBoxLayout(course_materials_group)
        
        # Description
        materials_desc = QLabel("Upload course materials (syllabus, textbooks, slides, etc.) to provide context for AI grading. "
                               "The AI will reference these materials when evaluating student submissions.")
        materials_desc.setWordWrap(True)
        materials_desc.setStyleSheet("color: #666; margin-bottom: 10px; font-size: 10pt;")
        course_materials_layout.addWidget(materials_desc)
        
        # File upload area
        self.setup_course_materials_widget(course_materials_layout)
        
        layout.addWidget(course_materials_group)
        
        # Budget Section - consolidates all cost-related controls and displays
        budget_group = QGroupBox("💰 Budget and Cost Calculation")
        budget_layout = QVBoxLayout(budget_group)
        
        # Budget description
        budget_desc = QLabel("Set your budget and see exact costs before grading begins. "
                            "Step 1 downloads submissions to calculate precise token usage.")
        budget_desc.setWordWrap(True)
        budget_desc.setStyleSheet("color: #666; margin-bottom: 10px; font-size: 10pt;")
        budget_layout.addWidget(budget_desc)
        
        # Move budget controls here (populate from setup_budget_widget)
        self.setup_budget_widget(budget_layout)
        
        # Course Materials List (integrated with budget display)
        course_materials_widget = QWidget()
        course_materials_layout = QVBoxLayout(course_materials_widget)
        course_materials_layout.addWidget(QLabel("📚 Course Materials:"))
        
        # Create the course materials file list directly here
        self.course_files_list = QTreeWidget()
        self.course_files_list.setHeaderLabels(["📄 File", "🔢 Tokens", "💰 Cost", "📊 Budget Impact"])
        self.course_files_list.setMaximumHeight(125)  # Average of 100 and 150
        self.course_files_list.setRootIsDecorated(False)  # No expand/collapse icons
        self.course_files_list.setAlternatingRowColors(True)
        
        # Set column widths to match downloaded submissions list exactly
        header = self.course_files_list.header()
        header.setStretchLastSection(False)  # Don't auto-stretch last column
        
        # File name column takes remaining space
        self.course_files_list.setColumnWidth(0, 200)  # Same as downloaded submissions
        header.setSectionResizeMode(0, header.ResizeMode.Stretch)  # File name stretches
        
        # Match the column widths from downloaded submissions list exactly
        self.course_files_list.setColumnWidth(1, 100)   # Tokens - same as submissions
        self.course_files_list.setColumnWidth(2, 100)   # Cost - same as submissions  
        self.course_files_list.setColumnWidth(3, 120)   # Budget Impact - wider to prevent cutoff
        
        course_materials_layout.addWidget(self.course_files_list)
        
        # File list buttons
        file_btn_layout = QHBoxLayout()
        remove_btn = QPushButton("Remove Selected")
        remove_btn.clicked.connect(self.remove_selected_course_file)
        file_btn_layout.addWidget(remove_btn)
        
        clear_all_btn = QPushButton("Clear All")
        clear_all_btn.clicked.connect(self.clear_all_course_files)
        file_btn_layout.addWidget(clear_all_btn)
        file_btn_layout.addStretch()
        
        course_materials_layout.addLayout(file_btn_layout)
        
        # Store as instance variable for potential future reference
        self.course_materials_budget_container = course_materials_widget
        
        budget_layout.addWidget(course_materials_widget)
        
        # Downloaded Submissions List (initially empty)
        submissions_widget = QWidget()
        submissions_layout = QVBoxLayout(submissions_widget)
        submissions_layout.addWidget(QLabel("📥 Downloaded Submissions:"))
        
        self.downloaded_submissions_list = QTreeWidget()
        self.downloaded_submissions_list.setHeaderLabels(["👨‍🎓 Student", "📄 Files", "🔢 Tokens", "💰 Cost", "📊 Budget Impact"])
        self.downloaded_submissions_list.setMaximumHeight(125)  # Average of 100 and 150  
        self.downloaded_submissions_list.setRootIsDecorated(False)
        self.downloaded_submissions_list.setAlternatingRowColors(True)
        
        # Initially show placeholder
        placeholder_item = QTreeWidgetItem(self.downloaded_submissions_list)
        placeholder_item.setText(0, "No submissions downloaded yet")
        placeholder_item.setText(1, "Use Step 1 to download")
        placeholder_item.setText(2, "-")
        placeholder_item.setText(3, "-")
        placeholder_item.setText(4, "-")
        
        # Set column widths for downloaded submissions
        header = self.downloaded_submissions_list.header()
        header.setStretchLastSection(False)
        self.downloaded_submissions_list.setColumnWidth(0, 200)  # Student name
        self.downloaded_submissions_list.setColumnWidth(1, 100)  # Files
        self.downloaded_submissions_list.setColumnWidth(2, 100)  # Tokens  
        self.downloaded_submissions_list.setColumnWidth(3, 100)  # Cost
        self.downloaded_submissions_list.setColumnWidth(4, 120)  # Budget Impact - wider to prevent cutoff
        header.setSectionResizeMode(0, header.ResizeMode.Stretch)
        
        submissions_layout.addWidget(self.downloaded_submissions_list)
        
        # Add buttons for downloaded submissions management
        submissions_btn_layout = QHBoxLayout()
        
        remove_selected_submission_btn = QPushButton("Remove Selected")
        remove_selected_submission_btn.clicked.connect(self.remove_selected_submission)
        submissions_btn_layout.addWidget(remove_selected_submission_btn)
        
        clear_all_submissions_btn = QPushButton("Clear All")
        clear_all_submissions_btn.clicked.connect(self.clear_all_submissions)
        submissions_btn_layout.addWidget(clear_all_submissions_btn)
        submissions_btn_layout.addStretch()
        
        submissions_layout.addLayout(submissions_btn_layout)
        budget_layout.addWidget(submissions_widget)
        
        # Total token usage summary at bottom of Budget section with integrated "Within Budget" indicator
        self.budget_summary_label = QLabel("📊 Total Usage: 0 tokens ($0.00) | Budget: $1.00 | Remaining: $1.00 | 🟢 Within budget")
        self.budget_summary_label.setStyleSheet("font-size: 11pt; font-weight: bold; padding: 8px; "
                                              "background-color: #e3f2fd; border-radius: 4px; margin-top: 10px;")
        budget_layout.addWidget(self.budget_summary_label)
        
        # Initialize the labels for use in update_token_display (but don't display them)
        self.token_counter_label = QLabel("0 / 25,000 tokens (0%)")
        self.token_counter_label.setStyleSheet("font-size: 11pt; font-weight: bold;")
        self.token_counter_label.hide()  # Hide but keep for compatibility
        
        self.cost_estimate_label = QLabel("Course materials cost: $0.00")
        self.cost_estimate_label.setStyleSheet("font-size: 10pt; color: #6c757d;")
        self.cost_estimate_label.hide()  # Hide but keep for compatibility
        
        # Keep the token status label for internal use but hide it
        self.token_status_label = QLabel("🟢 Within budget")
        self.token_status_label.hide()  # Hide but keep for compatibility
        
        layout.addWidget(budget_group)
        
        # Step 2: AI Grading Group (renamed and updated)
        step2_group = QGroupBox("🤖 Step 2: AI Grading")
        step2_layout = QVBoxLayout(step2_group)
        
        step2_desc = QLabel("Run AI grading on downloaded submissions and save results for review")
        step2_layout.addWidget(step2_desc)
        
        step2_button_layout = QHBoxLayout()
        self.step2_button = QPushButton("🚀 Start Step 2: Grade Submissions")
        self.step2_button.setEnabled(False)  # Initially disabled until Step 1 completes
        self.step2_button.clicked.connect(self.start_step2_grading)
        step2_button_layout.addWidget(self.step2_button)
        
        self.step2_status = QLabel("Ready (Complete Step 1 first)")
        self.step2_status.setStyleSheet("color: #6c757d;")
        self.step2_status.setMinimumWidth(120)  # Ensure enough space for longer status text
        step2_button_layout.addWidget(self.step2_status)
        step2_button_layout.addStretch()
        
        step2_layout.addLayout(step2_button_layout)
        
        # Progress bar for step 2
        self.progress_step2 = QProgressBar()
        step2_layout.addWidget(self.progress_step2)
        
        self.progress_desc_step2 = QLabel("Ready")
        self.progress_desc_step2.setStyleSheet("color: blue; font-size: 10px;")
        step2_layout.addWidget(self.progress_desc_step2)
        
        layout.addWidget(step2_group)
        
        # Step 3: Review and Upload Group (renamed from old Step 2)
        step3_group = QGroupBox("👁️ Step 3: Review and Upload")
        step3_layout = QVBoxLayout(step3_group)
        
        step3_desc = QLabel("Review AI grading results, make adjustments, upload final grades")
        step3_layout.addWidget(step3_desc)
        
        step3_button_layout = QHBoxLayout()
        self.step3_button = QPushButton("📋 Start Step 3: Upload")
        self.step3_button.setEnabled(False)
        self.step3_button.clicked.connect(self.start_step3)
        step3_button_layout.addWidget(self.step3_button)
        
        review_folder_btn = QPushButton("📁 Open Review Folder")
        review_folder_btn.clicked.connect(self.open_review_folder)
        step3_button_layout.addWidget(review_folder_btn)
        
        # Store as instance variable for enabling/disabling
        self.review_folder_btn = review_folder_btn
        
        self.step3_status = QLabel("Ready")
        self.step3_status.setStyleSheet("color: blue;")
        self.step3_status.setMinimumWidth(120)  # Ensure enough space for longer status text
        step3_button_layout.addWidget(self.step3_status)
        step3_button_layout.addStretch()
        
        step3_layout.addLayout(step3_button_layout)
        
        # Progress bar for step 3
        self.progress_step3 = QProgressBar()
        step3_layout.addWidget(self.progress_step3)
        
        self.progress_desc_step3 = QLabel("Ready")
        self.progress_desc_step3.setStyleSheet("color: blue; font-size: 10px;")
        step3_layout.addWidget(self.progress_desc_step3)
        
        layout.addWidget(step3_group)
        
        # Log Group
        log_group = QGroupBox("📋 1-2-3 Grading Log")
        log_layout = QVBoxLayout(log_group)
        
        self.two_step_log = QTextEdit()
        self.two_step_log.setReadOnly(True)
        self.two_step_log.setMaximumHeight(500)  # Increased from 200 to 500 (2.5x)
        self.two_step_log.setPlainText("1-2-3 grading log will appear here...")
        log_layout.addWidget(self.two_step_log)
        
        layout.addWidget(log_group)
        layout.addStretch()
        
        # Set the content widget in the scroll area
        scroll_area.setWidget(content_widget)
        
        # Add scroll area to the main tab layout
        tab_layout.addWidget(scroll_area)
        
        return tab
    
    def create_single_step_tab(self) -> QWidget:
        """Create Single-Step Grading tab matching Tkinter structure"""
        # Create main tab widget
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Create scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Create content widget that will go inside scroll area
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        
        # Warning Group
        warning_group = QGroupBox("⚠️ Single-Step Grading Warning")
        warning_layout = QVBoxLayout(warning_group)
        
        warning_text = QLabel(
            "⚠️ WARNING: Single-step grading uploads results immediately!\n\n"
            "This mode downloads submissions, grades them with AI, and uploads results "
            "directly to Canvas without human review. Use only when you fully trust "
            "the AI grading configuration.\n\n"
            "For safety, consider using Two-Step Grading instead."
        )
        warning_text.setStyleSheet("color: red; font-weight: bold; padding: 10px; background-color: #fff3cd; border-radius: 4px;")
        warning_layout.addWidget(warning_text)
        
        layout.addWidget(warning_group)
        
        # Assignment Selection Group
        single_assignment_group = QGroupBox("📚 Assignment Selection")
        single_assignment_layout = QVBoxLayout(single_assignment_group)
        
        # Course ID input for single step
        single_course_layout = QHBoxLayout()
        single_course_layout.addWidget(QLabel("Course ID:"))
        self.single_course_id_entry = QLineEdit()
        self.single_course_id_entry.setPlaceholderText("Enter Canvas course ID")
        single_course_layout.addWidget(self.single_course_id_entry)
        single_assignment_layout.addLayout(single_course_layout)
        
        # Assignment dropdown
        single_assignment_select_layout = QHBoxLayout()
        single_assignment_select_layout.addWidget(QLabel("Select Assignment:"))
        self.single_assignment_combo = ScrollFriendlyComboBox()
        self.single_assignment_combo.addItem("No assignments loaded")
        self.single_assignment_combo.setMinimumWidth(280)  # Standard width for dropdown
        single_assignment_select_layout.addWidget(self.single_assignment_combo)
        
        refresh_single_assignments_btn = QPushButton("↻ Refresh")
        refresh_single_assignments_btn.clicked.connect(self.refresh_assignments)
        single_assignment_select_layout.addWidget(refresh_single_assignments_btn)
        
        single_assignment_layout.addLayout(single_assignment_select_layout)
        
        # Assignment info for single step (removed - no details functionality implemented)
        # self.single_assignment_info = QLabel("Select an assignment to view details")
        # self.single_assignment_info.setStyleSheet("padding: 10px; background-color: #f8f9fa; border-radius: 4px;")
        # single_assignment_layout.addWidget(self.single_assignment_info)
        
        layout.addWidget(single_assignment_group)
        
        # Configuration Group
        single_config_group = QGroupBox("⚙️ Grading Configuration")
        single_config_layout = QVBoxLayout(single_config_group)
        
        # Rubric source selection for single step
        single_rubric_source_layout = QHBoxLayout()
        single_rubric_source_layout.addWidget(QLabel("Rubric Source:"))
        
        # Create button group for radio buttons (ensures mutual exclusivity)
        self.single_rubric_source_group = QButtonGroup(self)
        
        self.single_local_rubric_radio = QRadioButton("Local File")
        self.single_local_rubric_radio.setChecked(True)  # Default to local file
        self.single_local_rubric_radio.toggled.connect(self.on_single_rubric_source_changed)
        self.single_rubric_source_group.addButton(self.single_local_rubric_radio, 0)
        single_rubric_source_layout.addWidget(self.single_local_rubric_radio)
        
        self.single_canvas_rubric_radio = QRadioButton("Canvas Rubric")
        self.single_canvas_rubric_radio.toggled.connect(self.on_single_rubric_source_changed)
        self.single_rubric_source_group.addButton(self.single_canvas_rubric_radio, 1)
        single_rubric_source_layout.addWidget(self.single_canvas_rubric_radio)
        
        single_rubric_source_layout.addStretch()
        single_config_layout.addLayout(single_rubric_source_layout)
        
        # Local rubric file selection (initially visible)
        self.single_local_rubric_widget = QWidget()
        single_local_rubric_layout = QHBoxLayout(self.single_local_rubric_widget)
        single_local_rubric_layout.setContentsMargins(0, 0, 0, 0)
        single_local_rubric_layout.addWidget(QLabel("Rubric File:"))
        self.single_rubric_entry = QLineEdit("rubrics/sample_rubric.json")
        self.single_rubric_entry.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        single_local_rubric_layout.addWidget(self.single_rubric_entry)
        
        browse_single_rubric_btn = QPushButton("Browse")
        browse_single_rubric_btn.setFixedHeight(21)
        browse_single_rubric_btn.clicked.connect(self.browse_single_rubric)
        single_local_rubric_layout.addWidget(browse_single_rubric_btn)
        
        single_config_layout.addWidget(self.single_local_rubric_widget)
        
        # Canvas rubric info for single step (initially hidden)
        self.single_canvas_rubric_widget = QWidget()
        single_canvas_rubric_layout = QHBoxLayout(self.single_canvas_rubric_widget)
        single_canvas_rubric_layout.setContentsMargins(0, 0, 0, 0)
        single_canvas_rubric_info = QLabel("✓ Canvas rubric will be downloaded automatically from the selected assignment")
        single_canvas_rubric_info.setStyleSheet("color: #28a745; font-weight: bold; padding: 5px;")
        single_canvas_rubric_layout.addWidget(single_canvas_rubric_info)
        single_canvas_rubric_layout.addStretch()
        
        single_config_layout.addWidget(self.single_canvas_rubric_widget)
        self.single_canvas_rubric_widget.hide()  # Initially hidden
        
        # Instructor config
        single_instructor_layout = QHBoxLayout()
        single_instructor_layout.addWidget(QLabel("Instructor Config (Optional):"))
        self.single_instructor_entry = QLineEdit("config/grading_instructor_config.json")
        single_instructor_layout.addWidget(self.single_instructor_entry)
        
        browse_single_instructor_btn = QPushButton("Browse")
        browse_single_instructor_btn.clicked.connect(self.browse_single_instructor_config)
        single_instructor_layout.addWidget(browse_single_instructor_btn)
        
        new_single_instructor_btn = QPushButton("New")
        new_single_instructor_btn.setToolTip("Create a new instructor configuration with guided setup")
        new_single_instructor_btn.clicked.connect(self.create_new_instructor_config)
        single_instructor_layout.addWidget(new_single_instructor_btn)
        
        single_config_layout.addLayout(single_instructor_layout)
        
        layout.addWidget(single_config_group)
        
        # Course Materials Group for Single-Step
        single_materials_group = QGroupBox("📚 Course Materials (Optional)")
        single_materials_layout = QVBoxLayout(single_materials_group)
        
        # Note that this shares the same course materials as two-step
        shared_materials_note = QLabel("📝 Note: Course materials are shared between Two-Step and Single-Step grading modes.")
        shared_materials_note.setStyleSheet("color: #6c757d; font-style: italic; font-size: 9pt;")
        single_materials_layout.addWidget(shared_materials_note)
        
        # Quick status display
        self.single_materials_status = QLabel("No course materials uploaded")
        self.single_materials_status.setStyleSheet("color: #6c757d; font-size: 10pt;")
        single_materials_layout.addWidget(self.single_materials_status)
        
        # Button to go to two-step tab for materials management
        goto_twostep_btn = QPushButton("📚 Manage Course Materials in Two-Step Tab")
        goto_twostep_btn.clicked.connect(lambda: self.tab_widget.setCurrentIndex(1))  # Switch to Two-Step tab
        single_materials_layout.addWidget(goto_twostep_btn)
        
        layout.addWidget(single_materials_group)
        
        # Direct Grading Action Group
        action_group = QGroupBox("⚡ Direct Grading")
        action_layout = QVBoxLayout(action_group)
        
        # Final warning
        final_warning = QLabel("⚠️ WARNING: This will grade and upload results immediately without review!")
        final_warning.setStyleSheet("color: red; font-weight: bold; font-size: 11px;")
        action_layout.addWidget(final_warning)
        
        # Start button
        button_layout = QHBoxLayout()
        self.single_grade_button = QPushButton("⚡ Start Single-Step Grading")
        self.single_grade_button.setEnabled(False)
        self.single_grade_button.setStyleSheet("QPushButton { background-color: #dc3545; color: white; font-weight: bold; }")
        self.single_grade_button.clicked.connect(self.start_single_grading)
        button_layout.addWidget(self.single_grade_button)
        
        self.single_status = QLabel("Ready")
        self.single_status.setStyleSheet("color: blue;")
        self.single_status.setMinimumWidth(120)  # Ensure enough space for longer status text
        button_layout.addWidget(self.single_status)
        button_layout.addStretch()
        
        action_layout.addLayout(button_layout)
        
        # Progress
        self.progress_single_step = QProgressBar()
        action_layout.addWidget(self.progress_single_step)
        
        self.progress_desc_single = QLabel("Ready")
        self.progress_desc_single.setStyleSheet("color: blue; font-size: 10px;")
        action_layout.addWidget(self.progress_desc_single)
        
        layout.addWidget(action_group)
        
        # Log Group
        log_group = QGroupBox("📝 Grading Log")
        log_layout = QVBoxLayout(log_group)
        
        self.single_step_log = QTextEdit()
        self.single_step_log.setReadOnly(True)
        self.single_step_log.setMaximumHeight(500)  # Increased from 200 to 500 (2.5x)
        self.single_step_log.setPlainText("Single-step grading log will appear here...")
        log_layout.addWidget(self.single_step_log)
        
        layout.addWidget(log_group)
        layout.addStretch()
        
        # Set the content widget in the scroll area
        scroll_area.setWidget(content_widget)
        
        # Add scroll area to the main tab layout
        tab_layout.addWidget(scroll_area)
        
        return tab
    
    def create_results_tab(self) -> QWidget:
        """Create Results tab matching Tkinter structure"""
        # Create main tab widget
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Create scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Create content widget that will go inside scroll area
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        
        # Results Summary Group
        summary_group = QGroupBox("📊 Grading Results Summary")
        summary_layout = QVBoxLayout(summary_group)
        
        self.results_summary = QLabel("No grading sessions completed yet.")
        self.results_summary.setStyleSheet("font-size: 12px; padding: 10px; background-color: #f8f9fa; border-radius: 4px;")
        summary_layout.addWidget(self.results_summary)
        
        layout.addWidget(summary_group)
        
        # Recent Sessions Group
        sessions_group = QGroupBox("📈 Recent Grading Sessions")
        sessions_layout = QVBoxLayout(sessions_group)
        
        # Session list
        self.sessions_list = QTextEdit()
        self.sessions_list.setReadOnly(True)
        self.sessions_list.setMaximumHeight(150)
        self.sessions_list.setPlainText("No recent sessions to display.")
        sessions_layout.addWidget(self.sessions_list)
        
        # Session management buttons
        session_buttons_layout = QHBoxLayout()
        view_session_btn = QPushButton("👁️ View Session Details")
        export_results_btn = QPushButton("📤 Export Results")
        clear_history_btn = QPushButton("🗑️ Clear History")
        
        session_buttons_layout.addWidget(view_session_btn)
        session_buttons_layout.addWidget(export_results_btn)
        session_buttons_layout.addWidget(clear_history_btn)
        session_buttons_layout.addStretch()
        
        sessions_layout.addLayout(session_buttons_layout)
        
        layout.addWidget(sessions_group)
        
        # Statistics Group
        stats_group = QGroupBox("📊 Grading Statistics")
        stats_layout = QVBoxLayout(stats_group)
        
        self.stats_display = QTextEdit()
        self.stats_display.setReadOnly(True)
        self.stats_display.setMaximumHeight(200)
        self.stats_display.setPlainText("Statistics will appear here after grading sessions.")
        stats_layout.addWidget(self.stats_display)
        
        layout.addWidget(stats_group)
        layout.addStretch()
        
        # Set the content widget in the scroll area
        scroll_area.setWidget(content_widget)
        
        # Add scroll area to the main tab layout
        tab_layout.addWidget(scroll_area)
        
        return tab
    
    def create_ducktest_two_step_tab(self) -> QWidget:
        """Create DuckTest Two-Step Assessment tab (placeholder for now)"""
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Add a placeholder message
        placeholder_group = QGroupBox("🚧 Two-Step Assessment Creation")
        placeholder_layout = QVBoxLayout(placeholder_group)
        
        placeholder_label = QLabel("This tab will contain the Two-Step Assessment creation interface.\n\n"
                                 "Features coming soon:\n"
                                 "• Create assessment rubrics\n"
                                 "• Generate question banks\n"
                                 "• Design multi-stage assessments\n"
                                 "• Preview and review assessment content")
        placeholder_label.setStyleSheet("font-size: 12pt; padding: 20px;")
        placeholder_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        placeholder_layout.addWidget(placeholder_label)
        tab_layout.addWidget(placeholder_group)
        tab_layout.addStretch()
        
        return tab
    
    def create_ducktest_one_step_tab(self) -> QWidget:
        """Create DuckTest One-Step Assessment tab (placeholder for now)"""
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Add a placeholder message
        placeholder_group = QGroupBox("⚡ One-Step Assessment Creation")
        placeholder_layout = QVBoxLayout(placeholder_group)
        
        placeholder_label = QLabel("This tab will contain the One-Step Assessment creation interface.\n\n"
                                 "Features coming soon:\n"
                                 "• Quick assessment generation\n"
                                 "• Automated question creation\n"
                                 "• Instant assessment deployment\n"
                                 "• Real-time assessment analytics")
        placeholder_label.setStyleSheet("font-size: 12pt; padding: 20px;")
        placeholder_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        placeholder_layout.addWidget(placeholder_label)
        tab_layout.addWidget(placeholder_group)
        tab_layout.addStretch()
        
        return tab
    
    def create_review_tab(self, review_spreadsheet_path, submission_folder_path) -> QWidget:
        """Create Review Tab with split-panel layout for reviewing graded submissions"""
        # Create main tab widget
        tab = QWidget()
        tab_layout = QVBoxLayout(tab)
        
        # Create scroll area for the entire tab content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Create content widget that will go inside scroll area
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        
        # Header with assignment info
        header_group = QGroupBox("👁️ Review Graded Submissions")
        header_layout = QVBoxLayout(header_group)
        
        # Assignment details
        self.review_assignment_info = QLabel("Loading assignment information...")
        self.review_assignment_info.setStyleSheet("font-weight: bold; padding: 5px;")
        header_layout.addWidget(self.review_assignment_info)
        
        content_layout.addWidget(header_group)
        
        # Main content: Navigation controls
        nav_group = QGroupBox("🧭 Navigation")
        nav_layout = QHBoxLayout(nav_group)
        
        # Previous button
        self.review_prev_btn = QPushButton("◀ Previous")
        self.review_prev_btn.clicked.connect(self.review_previous_submission)
        nav_layout.addWidget(self.review_prev_btn)
        
        # Submission selector dropdown (no label, full width)
        self.review_submission_combo = QComboBox()
        self.review_submission_combo.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.review_submission_combo.currentIndexChanged.connect(self.review_submission_changed)
        
        # Style the combo box to match Two-Step Grading tab
        self.review_submission_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
                padding-right: 25px;
            }
            QComboBox:hover {
                border-color: #adb5bd;
            }
            QComboBox:focus {
                border-color: #80bdff;
                outline: 0;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        
        nav_layout.addWidget(self.review_submission_combo)
        
        # Next button
        self.review_next_btn = QPushButton("Next ▶")
        self.review_next_btn.clicked.connect(self.review_next_submission)
        nav_layout.addWidget(self.review_next_btn)
        
        # Progress info
        self.review_progress_label = QLabel("0 of 0")
        self.review_progress_label.setStyleSheet("font-weight: bold; padding: 0 10px;")
        nav_layout.addWidget(self.review_progress_label)
        
        content_layout.addWidget(nav_group)
        
        # Main split panel
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)
        
        # Left panel: Submission content (2/3 width)
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        
        submission_group = QGroupBox("📄 Submission Content")
        submission_layout = QVBoxLayout(submission_group)
        
        # Submission file info with navigation buttons
        file_info_layout = QHBoxLayout()
        
        self.review_file_info = QLabel("No submission selected")
        self.review_file_info.setStyleSheet("font-weight: bold; color: #666; padding: 5px;")
        file_info_layout.addWidget(self.review_file_info)
        
        file_info_layout.addStretch()
        
        # View In Directory button (aligned right)
        self.review_view_directory_btn = QPushButton("📁 View In Directory")
        self.review_view_directory_btn.setToolTip("Open student's submission folder")
        self.review_view_directory_btn.clicked.connect(self.review_open_submission_directory)
        self.review_view_directory_btn.setEnabled(False)  # Initially disabled
        file_info_layout.addWidget(self.review_view_directory_btn)
        
        # File navigation buttons (initially hidden)
        self.review_prev_file_btn = QPushButton("◀")
        self.review_prev_file_btn.setToolTip("Previous file")
        self.review_prev_file_btn.setFixedSize(30, 25)
        self.review_prev_file_btn.clicked.connect(self.review_previous_submission_file)
        self.review_prev_file_btn.hide()
        file_info_layout.addWidget(self.review_prev_file_btn)
        
        self.review_file_counter = QLabel("")
        self.review_file_counter.setStyleSheet("font-weight: bold; color: #666; padding: 0 10px;")
        self.review_file_counter.hide()
        file_info_layout.addWidget(self.review_file_counter)
        
        self.review_next_file_btn = QPushButton("▶")
        self.review_next_file_btn.setToolTip("Next file")
        self.review_next_file_btn.setFixedSize(30, 25)
        self.review_next_file_btn.clicked.connect(self.review_next_submission_file)
        self.review_next_file_btn.hide()
        file_info_layout.addWidget(self.review_next_file_btn)
        
        submission_layout.addLayout(file_info_layout)
        
        # Submission content viewer - using stacked widget for different view types
        self.submission_viewer_stack = QStackedWidget()
        self.submission_viewer_stack.setMinimumHeight(400)
        
        # Text viewer for extracted text
        self.review_submission_viewer = QTextEdit()
        self.review_submission_viewer.setReadOnly(True)
        self.review_submission_viewer.setPlainText("Select a submission to view its content...")
        self.submission_viewer_stack.addWidget(self.review_submission_viewer)
        
        # Web viewer for rendered documents (PDFs, converted Word docs)
        if WEB_ENGINE_AVAILABLE:
            self.review_document_viewer = QWebEngineView()
            
            # Apply custom CSS to style the web view scrollbars
            scrollbar_css = """
            QWebEngineView {
                background-color: white;
            }
            """
            self.review_document_viewer.setStyleSheet(scrollbar_css)
            
            # Inject CSS for web content scrollbars using JavaScript
            web_scrollbar_js = """
            var style = document.createElement('style');
            style.textContent = `
                ::-webkit-scrollbar {
                    width: 12px;
                    height: 12px;
                }
                ::-webkit-scrollbar-track {
                    background: #f1f1f1;
                    border-radius: 6px;
                }
                ::-webkit-scrollbar-thumb {
                    background: #c1c1c1;
                    border-radius: 6px;
                    border: 1px solid #f1f1f1;
                }
                ::-webkit-scrollbar-thumb:hover {
                    background: #a8a8a8;
                }
                ::-webkit-scrollbar-corner {
                    background: #f1f1f1;
                }
            `;
            document.head.appendChild(style);
            """
            
            # Inject the scrollbar styling when the page loads
            self.review_document_viewer.loadFinished.connect(
                lambda success: self.on_document_load_finished(success, web_scrollbar_js)
            )
            
            # Add load started handler for debugging
            self.review_document_viewer.loadStarted.connect(
                lambda: print("DEBUG: QWebEngineView load started")
            )
            
            self.submission_viewer_stack.addWidget(self.review_document_viewer)
        else:
            self.review_document_viewer = None
        
        submission_layout.addWidget(self.submission_viewer_stack)
        
        # View mode selector
        view_mode_layout = QHBoxLayout()
        view_mode_layout.addWidget(QLabel("View Mode:"))
        
        if WEB_ENGINE_AVAILABLE:
            self.view_mode_rendered_btn = QPushButton("🖼️ Rendered")
            self.view_mode_rendered_btn.setCheckable(True)
            self.view_mode_rendered_btn.setChecked(True)  # Default to rendered view
            self.view_mode_rendered_btn.setToolTip("View rendered document (PDFs, converted Word docs)")
            self.view_mode_rendered_btn.clicked.connect(lambda: self.switch_view_mode("rendered"))
            view_mode_layout.addWidget(self.view_mode_rendered_btn)
        
        self.view_mode_text_btn = QPushButton("📄 Text")
        self.view_mode_text_btn.setCheckable(True)
        self.view_mode_text_btn.setChecked(False if WEB_ENGINE_AVAILABLE else True)  # Only checked if no web engine
        self.view_mode_text_btn.setToolTip("View extracted text content")
        self.view_mode_text_btn.clicked.connect(lambda: self.switch_view_mode("text"))
        view_mode_layout.addWidget(self.view_mode_text_btn)
        
        if WEB_ENGINE_AVAILABLE:
            # Create button group for exclusive selection
            self.view_mode_group = QButtonGroup()
            self.view_mode_group.addButton(self.view_mode_rendered_btn)
            self.view_mode_group.addButton(self.view_mode_text_btn)
        else:
            self.view_mode_rendered_btn = None
            self.view_mode_group = None
        
        view_mode_layout.addStretch()
        submission_layout.addLayout(view_mode_layout)
        
        left_layout.addWidget(submission_group)
        
        # Right panel: Editable comments and score (1/3 width)
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # Score editing section
        score_group = QGroupBox("📊 Score")
        score_layout = QVBoxLayout(score_group)
        
        score_input_layout = QHBoxLayout()
        score_input_layout.addWidget(QLabel("Points:"))
        
        # Create score input with separate max score display
        score_container = QWidget()
        score_container_layout = QHBoxLayout(score_container)
        score_container_layout.setContentsMargins(0, 0, 0, 0)
        score_container_layout.setSpacing(8)  # Small gap between score box and max score text
        
        # Use QLineEdit for the editable score part - standalone design
        self.review_score_entry = QLineEdit()
        self.review_score_entry.setPlaceholderText("0")
        self.review_score_entry.setText("0")
        self.review_score_entry.setMaximumWidth(80)  # Wider for standalone box
        self.review_score_entry.textChanged.connect(self.review_score_changed)
        
        # Create a separate label for the max score (to the right, not connected)
        self.review_max_score_display = QLabel("/ 100")
        self.review_max_score_display.setStyleSheet("""
            QLabel {
                color: #6c757d;
                font-size: 11pt;
                font-style: italic;
                padding: 0px;
                margin: 0px;
            }
        """)
        
        # Apply initial styling to the score entry
        self.review_score_entry.setStyleSheet("""
            QLineEdit {
                border: 2px solid #ced4da;
                border-radius: 6px;
                padding: 8px 10px;
                background-color: white;
                font-size: 11pt;
                font-weight: normal;
                color: #495057;
                min-height: 16px;
                max-height: 16px;
            }
            QLineEdit:focus {
                border-color: #80bdff;
                box-shadow: 0px 0px 0px 0.2rem rgba(0, 123, 255, 0.25);
                outline: 0;
                background-color: #fff;
            }
        """)
        
        # Connect focus events for dynamic styling
        self.review_score_entry.focusInEvent = lambda event: self.update_score_focus_style(True, event)
        self.review_score_entry.focusOutEvent = lambda event: self.update_score_focus_style(False, event)
        
        # Add both parts to the container with spacing
        score_container_layout.addWidget(self.review_score_entry)
        score_container_layout.addWidget(self.review_max_score_display)
        
        # Initially hide max score display until we know if max score is available
        self.review_max_score_display.hide()
        
        score_input_layout.addWidget(score_container)
        score_input_layout.addStretch()
        
        score_layout.addLayout(score_input_layout)
        right_layout.addWidget(score_group)
        
        # Comments editing section
        comments_group = QGroupBox("💭 Comments and Feedback")
        comments_layout = QVBoxLayout(comments_group)
        
        # Comment editor
        self.review_comments_editor = QTextEdit()
        self.review_comments_editor.setPlaceholderText("Enter feedback comments here...")
        self.review_comments_editor.setMinimumHeight(300)
        self.review_comments_editor.textChanged.connect(self.review_comments_changed)
        comments_layout.addWidget(self.review_comments_editor)
        
        # Comment editor controls
        comment_controls_layout = QHBoxLayout()
        
        self.review_clear_btn = QPushButton("🗑️ Clear")
        self.review_clear_btn.setToolTip("Clear all comments")
        self.review_clear_btn.clicked.connect(self.review_clear_comments)
        comment_controls_layout.addWidget(self.review_clear_btn)
        
        self.review_restore_btn = QPushButton("↻ Restore AI Comments")
        self.review_restore_btn.setToolTip("Restore original AI comments")
        self.review_restore_btn.clicked.connect(self.review_restore_ai_comments)
        comment_controls_layout.addWidget(self.review_restore_btn)
        
        comment_controls_layout.addStretch()
        
        comments_layout.addLayout(comment_controls_layout)
        right_layout.addWidget(comments_group)
        
        # Add panels to splitter
        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        
        # Set initial sizes (2/3 for left, 1/3 for right)
        splitter.setSizes([400, 200])
        
        content_layout.addWidget(splitter)
        
        # Save controls at bottom
        save_group = QGroupBox("💾 Save Changes")
        save_layout = QHBoxLayout(save_group)
        
        self.review_save_btn = QPushButton("💾 Save Current")
        self.review_save_btn.setToolTip("Save changes to current submission")
        self.review_save_btn.clicked.connect(self.review_save_current)
        self.review_save_btn.setEnabled(False)
        save_layout.addWidget(self.review_save_btn)
        
        self.review_save_all_btn = QPushButton("💾 Save All Changes")
        self.review_save_all_btn.setToolTip("Save all pending changes to spreadsheet")
        self.review_save_all_btn.clicked.connect(self.review_save_all_changes)
        save_layout.addWidget(self.review_save_all_btn)
        
        save_layout.addStretch()
        
        # Changes indicator
        self.review_changes_label = QLabel("No unsaved changes")
        self.review_changes_label.setStyleSheet("color: #666; font-style: italic;")
        save_layout.addWidget(self.review_changes_label)
        
        content_layout.addWidget(save_group)
        
        # Set the content widget in the scroll area
        scroll_area.setWidget(content_widget)
        
        # Add scroll area to the main tab layout
        tab_layout.addWidget(scroll_area)
        
        # Initialize review data
        self.review_data = {}
        self.review_current_index = 0
        self.review_spreadsheet_path = review_spreadsheet_path
        self.review_submission_folder = submission_folder_path
        self.review_unsaved_changes = set()  # Track submissions with unsaved changes
        self.review_original_data = {}  # Store original AI comments for restore function
        
        # Initialize default view mode - start with rendered view if available
        if WEB_ENGINE_AVAILABLE:
            print(f"DEBUG: Setting default view to rendered mode (stack index 1)")
            self.submission_viewer_stack.setCurrentIndex(1)  # Start with rendered view
        else:
            print(f"DEBUG: Web engine not available, using text mode (stack index 0)")
            self.submission_viewer_stack.setCurrentIndex(0)  # Fallback to text view
        
        # Don't load data here - it will be loaded after the tab is fully created
        
        return tab
    
    def setup_window(self):
        """Setup window properties"""
        self.setWindowTitle("DuckWorks Educational Suite - Multi-Tool Interface")
        
        # Set duck icon
        if Path("assets/icons8-flying-duck-48.png").exists():
            self.setWindowIcon(QIcon("assets/icons8-flying-duck-48.png"))
        
        # Home tab uses emoji house icon (no custom .png icon)
        
        self.center_window()
        
        # Apply professional styling matching modern design principles
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f8f9fa;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #dee2e6;
                border-radius: 8px;
                margin: 5px;
                padding-top: 10px;
                background-color: white;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #495057;
            }
            QPushButton {
                background-color: #007bff;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                height: 21px;
                min-height: 21px;
                max-height: 21px;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
            QPushButton:disabled {
                background-color: #6c757d;
            }
            QLineEdit, QTextEdit {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
            }
            QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
            }
            QComboBox:hover {
                border-color: #adb5bd;
            }
            QComboBox:focus {
                border-color: #80bdff;
                outline: 0;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
            QLineEdit:focus, QComboBox:focus, QTextEdit:focus {
                border-color: #80bdff;
                outline: 0;
                box-shadow: 0 0 0 0.2rem rgba(0, 123, 255, 0.25);
            }
            QTabWidget::pane {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #e9ecef;
                padding: 10px 15px;
                margin-right: 2px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background-color: white;
                border-bottom: 2px solid #007bff;
            }
            QProgressBar {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                text-align: center;
                background-color: #f8f9fa;
            }
            QProgressBar::chunk {
                background-color: #007bff;
                border-radius: 3px;
            }
            QCheckBox {
                spacing: 8px;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #6c757d;
                border-radius: 3px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #007bff;
                border-radius: 3px;
                background-color: #007bff;
            }
            QScrollArea {
                border: none;
                background-color: transparent;
            }
            QScrollBar:vertical {
                background-color: #f8f9fa;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #6c757d;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #495057;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                border: none;
                background: none;
            }
        """)
    
    def setup_budget_widget(self, parent_layout):
        """Setup the budget controls for the new Budget section"""
        
        # Initialize budget-related storage (if not already done)
        if not hasattr(self, 'token_budget'):
            self.token_budget = 25000  # Default token budget
            self.usd_budget = 1.0  # Default USD budget - much more reasonable!
            self.budget_mode = "USD"  # Current budget mode - default to USD with $1
            
            # Budget restrictions settings
            self.max_usd_cap = 100.0  # Static maximum USD cap for all models
            self.budget_restrictions_enabled = True  # Whether budget restrictions are active
        
        # Token budget setting
        budget_controls_layout = QHBoxLayout()
        budget_controls_layout.addWidget(QLabel("Budget:"))
        
        # Budget type selector
        self.budget_type_combo = QComboBox()
        self.budget_type_combo.addItems(["USD", "Tokens"])  # USD first as default
        self.budget_type_combo.currentTextChanged.connect(self.on_budget_type_changed)
        # Style to match other combo boxes on this tab
        self.budget_type_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background-color: white;
                font-size: 10pt;
                min-width: 80px;
                padding-right: 25px;
            }
            QComboBox:hover {
                border-color: #adb5bd;
            }
            QComboBox:focus {
                border-color: #80bdff;
                outline: 0;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        budget_controls_layout.addWidget(self.budget_type_combo)
        
        # Budget amount input box (custom QLineEdit instead of spinbox)
        self.budget_amount_input = QLineEdit()
        self.budget_amount_input.setText("1.000")  # Default USD value
        self.budget_amount_input.setPlaceholderText("Enter budget amount")
        
        # Custom container to show units
        budget_input_container = QWidget()
        budget_input_layout = QHBoxLayout(budget_input_container)
        budget_input_layout.setContentsMargins(0, 0, 0, 0)
        budget_input_layout.setSpacing(0)  # No spacing between elements
        
        budget_input_layout.addWidget(self.budget_amount_input)
        
        # Combined unit and equivalent label (single label to eliminate spacing)
        self.budget_unit_equivalent_label = QLabel("USD (6,666,666 tokens)")
        self.budget_unit_equivalent_label.setStyleSheet("color: #6c757d; font-weight: bold; font-size: 10pt; padding: 6px 6px 6px 6px;")
        budget_input_layout.addWidget(self.budget_unit_equivalent_label)
        
        # Set size constraints
        self.budget_amount_input.setMinimumWidth(100)
        self.budget_amount_input.setMaximumWidth(120)
        
        # Connect to validation and update handlers
        self.budget_amount_input.textChanged.connect(self.on_budget_input_changed)
        self.budget_amount_input.editingFinished.connect(self.validate_budget_input)
        
        # Styling to match other fields
        self.budget_amount_input.setStyleSheet("""
            QLineEdit {
                border: 1px solid #ddd;
                border-radius: 4px;
                padding: 6px 8px;
                background-color: white;
                font-size: 11pt;
            }
            QLineEdit:focus {
                border-color: #4a90e2;
            }
        """)
        
        budget_controls_layout.addWidget(budget_input_container)
        
        # Add some spacing before the checkbox
        budget_controls_layout.addSpacing(20)
        
        # Budget restriction control (moved inline)
        self.disable_budget_restrictions_checkbox = QCheckBox("Disable Budget Restrictions")
        self.disable_budget_restrictions_checkbox.setToolTip("Remove budget limits entirely (shows warning before enabling)")
        self.disable_budget_restrictions_checkbox.stateChanged.connect(self.on_budget_restrictions_changed)
        self.disable_budget_restrictions_checkbox.setStyleSheet("color: #dc3545; font-weight: bold;")
        budget_controls_layout.addWidget(self.disable_budget_restrictions_checkbox)
        
        budget_controls_layout.addStretch()
        parent_layout.addLayout(budget_controls_layout)
        
        # Initialize USD budget display
        QTimer.singleShot(100, self.update_initial_budget_display)
    
    def setup_course_materials_widget(self, parent_layout):
        """Setup the course materials upload widget with drag and drop and token counting"""
        
        # Initialize course materials storage (if not already done by setup_budget_widget)
        if not hasattr(self, 'course_materials'):
            self.course_materials = []  # List of file paths
            self.course_materials_tokens = 0
            self.course_materials_token_counts = {}  # Per-file token counts
            self.course_materials_cost = 0.0
        
        # File upload area with drag and drop - separate drop zone from buttons
        upload_container = QWidget()
        upload_container_layout = QVBoxLayout(upload_container)
        upload_container_layout.setSpacing(10)
        
        # Drop zone only (with dashed border)
        drop_zone = QWidget()
        drop_zone.setAcceptDrops(True)
        drop_zone.setStyleSheet("""
            QWidget {
                border: 2px dashed #dee2e6;
                border-radius: 8px;
                background-color: #f8f9fa;
                min-height: 80px;
            }
            QWidget:hover {
                border-color: #007bff;
                background-color: #e3f2fd;
            }
        """)
        
        # Override drag and drop events for the drop zone only
        def dragEnterEvent(event):
            if event.mimeData().hasUrls():
                event.acceptProposedAction()
            else:
                event.ignore()
                
        def dropEvent(event):
            files = [url.toLocalFile() for url in event.mimeData().urls()]
            self.add_course_materials(files)
            event.acceptProposedAction()
            
        drop_zone.dragEnterEvent = dragEnterEvent
        drop_zone.dropEvent = dropEvent
        
        drop_zone_layout = QVBoxLayout(drop_zone)
        
        # Drop zone text
        drop_label = QLabel("📁 Drop Course Files Here")
        drop_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_label.setStyleSheet("font-size: 14pt; color: #6c757d; font-weight: bold;")
        drop_zone_layout.addWidget(drop_label)
        
        upload_container_layout.addWidget(drop_zone)
        
        # "Or" separator text
        or_label = QLabel("or use the buttons below to browse")
        or_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        or_label.setStyleSheet("font-size: 10pt; color: #adb5bd; margin: 5px 0;")
        upload_container_layout.addWidget(or_label)
        
        # Browse buttons (outside the dashed border)
        button_layout = QHBoxLayout()
        
        browse_files_btn = QPushButton("📄 Browse Files")
        browse_files_btn.setStyleSheet("""
            QPushButton {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px 16px;
                background-color: white;
                color: #495057;
                font-weight: normal;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border-color: #adb5bd;
            }
            QPushButton:pressed {
                background-color: #dee2e6;
            }
        """)
        browse_files_btn.clicked.connect(self.browse_course_files)
        button_layout.addWidget(browse_files_btn)
        
        browse_folder_btn = QPushButton("📁 Browse Folder")
        browse_folder_btn.setStyleSheet("""
            QPushButton {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px 16px;
                background-color: white;
                color: #495057;
                font-weight: normal;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border-color: #adb5bd;
            }
            QPushButton:pressed {
                background-color: #dee2e6;
            }
        """)
        browse_folder_btn.clicked.connect(self.browse_course_folder)
        button_layout.addWidget(browse_folder_btn)
        
        upload_container_layout.addLayout(button_layout)
        parent_layout.addWidget(upload_container)
        
        # Move file list to Budget section instead of here
        # File list and token counter will be shown in the Budget section
        
        # Additional instructions field (keep this here in Course Materials)
        instructions_layout = QVBoxLayout()
        instructions_layout.addWidget(QLabel("📝 Additional Course Material Instructions:"))
        
        self.course_materials_instructions = QTextEdit()
        self.course_materials_instructions.setMaximumHeight(60)
        self.course_materials_instructions.setPlaceholderText(
            "e.g., Focus on chapters 1-4 of the textbook, ignore appendices, "
            "use terminology from the syllabus when grading..."
        )
        instructions_layout.addWidget(self.course_materials_instructions)
        
        parent_layout.addLayout(instructions_layout)

    def center_window(self):
        """Center the window on the screen"""
        frame_geometry = self.frameGeometry()
        screen = self.screen().availableGeometry()
        center_point = screen.center()
        frame_geometry.moveCenter(center_point)
        self.move(frame_geometry.topLeft())
    
    # ========================================
    # Course Materials Management Methods
    # ========================================
    
    def update_token_budget(self, value):
        """Update the token budget and refresh displays - legacy method for compatibility"""
        self.token_budget = value
        self.update_token_display()
    
    def update_budget_amount(self, value):
        """Update budget amount based on current mode"""
        # Mark budget as initialized by user interaction
        self._budget_initialized = True
        
        if self.budget_mode == "Tokens":
            self.token_budget = int(value)
            # Update equivalent USD display
            budget_cost = self.calculate_cost(self.token_budget)
            self.budget_unit_equivalent_label.setText(f"tokens (${budget_cost:.4f})")
        else:  # USD mode
            self.usd_budget = value
            # Update equivalent tokens display
            model_data = self.get_selected_model_data()
            if model_data and 'input_price' in model_data:
                # Calculate equivalent tokens for this USD amount
                price_per_token = model_data['input_price'] / 1000.0
                equivalent_tokens = int(value / price_per_token) if price_per_token > 0 else 0
                self.token_budget = equivalent_tokens  # Keep internal token budget in sync
                self.budget_unit_equivalent_label.setText(f"USD ({equivalent_tokens:,} tokens)")
            else:
                # Fallback if no model data
                self.budget_unit_equivalent_label.setText("USD (? tokens)")
        
        self.update_token_display()
    
    def on_budget_input_changed(self, text):
        """Handle budget input text changes with validation"""
        # Mark budget as initialized by user interaction
        self._budget_initialized = True
        
        # Try to parse the input
        try:
            value = float(text) if text else 0.0
            
            if self.budget_mode == "Tokens":
                self.token_budget = int(value)
                # Update equivalent USD display
                budget_cost = self.calculate_cost(self.token_budget)
                self.budget_unit_equivalent_label.setText(f"tokens (${budget_cost:.4f})")
            else:  # USD mode
                self.usd_budget = value
                # Update equivalent tokens display
                model_data = self.get_selected_model_data()
                if model_data and 'input_price' in model_data:
                    price_per_token = model_data['input_price'] / 1000.0
                    equivalent_tokens = int(value / price_per_token) if price_per_token > 0 else 0
                    self.token_budget = equivalent_tokens
                    self.budget_unit_equivalent_label.setText(f"USD ({equivalent_tokens:,} tokens)")
                else:
                    self.budget_unit_equivalent_label.setText("USD (calculating...)")
            
            self.update_token_display()
            
        except ValueError:
            # Invalid input - don't update anything, just let user continue typing
            pass
    
    def validate_budget_input(self):
        """Validate budget input when editing is finished"""
        text = self.budget_amount_input.text()
        try:
            value = float(text)
            
            # Check ranges if restrictions are enabled
            if self.budget_restrictions_enabled:
                if self.budget_mode == "USD":
                    if value > self.max_usd_cap:
                        self.budget_amount_input.setText(f"{self.max_usd_cap:.4f}")
                        QMessageBox.information(self, "Budget Adjusted", f"Budget capped at ${self.max_usd_cap}")
                        return
                else:  # Tokens mode
                    model_data = self.get_selected_model_data()
                    if model_data and 'input_price' in model_data:
                        price_per_token = model_data['input_price'] / 1000.0
                        max_tokens = int(self.max_usd_cap / price_per_token) if price_per_token > 0 else 1000000
                        if value > max_tokens:
                            self.budget_amount_input.setText(str(max_tokens))
                            QMessageBox.information(self, "Budget Adjusted", f"Budget capped at {max_tokens:,} tokens")
                            return
            
            # Format the input nicely
            if self.budget_mode == "USD":
                self.budget_amount_input.setText(f"{value:.4f}")
            else:
                self.budget_amount_input.setText(str(int(value)))
                
        except ValueError:
            # Reset to last valid value
            if self.budget_mode == "USD":
                self.budget_amount_input.setText(f"{self.usd_budget:.4f}")
            else:
                self.budget_amount_input.setText(str(self.token_budget))
            QMessageBox.warning(self, "Invalid Input", "Please enter a valid number.")
    
    def on_budget_type_changed(self, budget_type):
        """Handle budget type change between USD and Tokens"""
        # Get the current value from the input field before switching
        try:
            current_value = float(self.budget_amount_input.text())
        except ValueError:
            current_value = 1.0 if budget_type == "USD" else 25000  # Default values
        
        print(f"DEBUG: Budget type changed to {budget_type}, current input value: {current_value}")
        
        self.budget_mode = budget_type
        
        # Mark budget as initialized by user interaction
        self._budget_initialized = True
        
        if budget_type == "Tokens":
            # Converting FROM USD TO Tokens
            # Current value is in USD, calculate equivalent tokens
            model_data = self.get_selected_model_data()
            if model_data and 'input_price' in model_data:
                price_per_token = model_data['input_price'] / 1000.0
                equivalent_tokens = int(current_value / price_per_token) if price_per_token > 0 else 25000
                print(f"DEBUG: Converting ${current_value} to {equivalent_tokens:,} tokens (price: ${price_per_token:.6f}/token)")
            else:
                # Fallback if no model data - use a reasonable conversion
                equivalent_tokens = int(current_value * 25000)  # Assume ~$0.00004 per token
                print(f"DEBUG: Using fallback conversion: ${current_value} to {equivalent_tokens:,} tokens")
            
            # Update internal tracking
            self.token_budget = equivalent_tokens
            self.usd_budget = current_value
            
            # Switch to token mode
            self.budget_amount_input.setText(str(equivalent_tokens))
            
            # Update equivalent USD display
            budget_cost = self.calculate_cost(equivalent_tokens)
            self.budget_unit_equivalent_label.setText(f"tokens (${budget_cost:.4f})")
            
        else:  # USD mode
            # Converting FROM Tokens TO USD
            # Current value is in tokens, calculate equivalent USD
            model_data = self.get_selected_model_data()
            if model_data and 'input_price' in model_data:
                price_per_token = model_data['input_price'] / 1000.0
                equivalent_usd = current_value * price_per_token if price_per_token > 0 else current_value / 25000
                print(f"DEBUG: Converting {current_value:,} tokens to ${equivalent_usd:.3f} (price: ${price_per_token:.6f}/token)")
            else:
                # Fallback if no model data
                equivalent_usd = current_value / 25000  # Assume ~$0.00004 per token
                print(f"DEBUG: Using fallback conversion: {current_value:,} tokens to ${equivalent_usd:.3f}")
            
            # Update internal tracking
            self.usd_budget = equivalent_usd
            self.token_budget = int(current_value)
            
            # Switch to USD mode
            self.budget_amount_input.setText(f"{equivalent_usd:.4f}")
            
            # Update equivalent tokens display
            self.budget_unit_equivalent_label.setText(f"USD ({int(current_value):,} tokens)")
        
        self.update_token_display()
    
    def get_selected_model_data(self):
        """Get the currently selected model data"""
        current_index = self.model_combo.currentIndex()
        if current_index >= 0:
            return self.model_combo.itemData(current_index)
        return None
    
    def on_model_changed(self, index):
        """Handle model selection change by index"""
        print(f"DEBUG: Model changed to index {index}")
        
        if index < 0:
            return
        
        # Update budget display when model changes (for both USD and token modes)
        if hasattr(self, 'budget_mode') and self.budget_mode == "USD":
            # In USD mode, keep USD amount and recalculate equivalent tokens
            model_data = self.get_selected_model_data()
            if model_data and 'input_price' in model_data:
                price_per_token = model_data['input_price'] / 1000.0
                equivalent_tokens = int(self.usd_budget / price_per_token) if price_per_token > 0 else 0
                self.token_budget = equivalent_tokens
                self.budget_unit_equivalent_label.setText(f"USD ({equivalent_tokens:,} tokens)")
                
                # Enforce budget cap if restrictions are enabled
                if self.budget_restrictions_enabled and self.usd_budget > self.max_usd_cap:
                    self.usd_budget = self.max_usd_cap
                    self.budget_amount_input.setText(f"{self.max_usd_cap:.4f}")
                    # Recalculate tokens with capped USD
                    equivalent_tokens = int(self.max_usd_cap / price_per_token) if price_per_token > 0 else 0
                    self.token_budget = equivalent_tokens
                    self.budget_unit_equivalent_label.setText(f"USD ({equivalent_tokens:,} tokens)")
                    print(f"DEBUG: Budget capped at ${self.max_usd_cap} due to model change")
        else:
            # In token mode, keep token amount and recalculate equivalent USD
            budget_cost = self.calculate_cost(self.token_budget)
            self.budget_unit_equivalent_label.setText(f"tokens (${budget_cost:.4f})")
            
            # Enforce budget cap if restrictions are enabled
            if self.budget_restrictions_enabled:
                model_data = self.get_selected_model_data()
                if model_data and 'input_price' in model_data:
                    price_per_token = model_data['input_price'] / 1000.0
                    max_tokens = int(self.max_usd_cap / price_per_token) if price_per_token > 0 else 1000000
                    if self.token_budget > max_tokens:
                        self.token_budget = max_tokens
                        self.budget_amount_input.setText(str(max_tokens))
                        budget_cost = self.calculate_cost(self.token_budget)
                        self.budget_unit_equivalent_label.setText(f"tokens (${budget_cost:.4f})")
                        print(f"DEBUG: Token budget capped at {max_tokens:,} due to model change")
            
            print(f"DEBUG: Updated budget to ${budget_cost:.4f}")
        
        # Update token display for course materials (this recalculates costs)
        if hasattr(self, 'course_materials_tokens'):
            # Recalculate course materials cost with new model
            self.course_materials_cost = self.calculate_cost(self.course_materials_tokens)
            self.update_token_display()
            print(f"DEBUG: Updated course materials cost to ${self.course_materials_cost:.4f}")
            
        # Update per-file cost display
        if hasattr(self, 'course_materials_token_counts'):
            self.update_file_list_display()
            print(f"DEBUG: Updated per-file costs")
        
        # Update downloaded submissions costs when model changes
        if hasattr(self, 'downloaded_submission_data') and self.downloaded_submission_data:
            print(f"DEBUG: Recalculating costs for {len(self.downloaded_submission_data)} downloaded submissions")
            total_cost = 0.0
            
            # Recalculate cost for each submission with new model
            for sub_data in self.downloaded_submission_data:
                tokens = sub_data.get('tokens', 0)
                new_cost = self.calculate_token_cost(tokens)
                sub_data['cost'] = new_cost
                total_cost += new_cost
            
            # Update the UI list with new costs and budget impact icons
            for i in range(self.downloaded_submissions_list.topLevelItemCount()):
                item = self.downloaded_submissions_list.topLevelItem(i)
                if item.text(0) != "No submissions downloaded yet":
                    student_name = item.text(0)
                    # Find matching submission data
                    for sub_data in self.downloaded_submission_data:
                        if sub_data['name'] == student_name:
                            item.setText(3, f"${sub_data['cost']:.4f}")
                            # Update budget impact icon with new cost
                            impact_icon = self.get_budget_impact_icon_by_cost(sub_data['cost'])
                            item.setText(4, impact_icon)
                            break
            
            # Update budget summary with new total cost
            self.update_budget_summary()
            print(f"DEBUG: Updated downloaded submissions total cost to ${total_cost:.4f}")
        
        # Update budget ranges since model pricing may have changed
        if hasattr(self, 'budget_restrictions_enabled'):
            self.update_budget_ranges()
            print(f"DEBUG: Updated budget ranges for new model")
    
    def on_model_text_changed(self, model_text):
        """Handle model selection change by text (for backward compatibility)"""
        # This is called when text changes, but we primarily use index-based handling now
        pass
    
    def update_initial_budget_display(self):
        """Initialize the budget display on startup - only runs once"""
        # Skip if budget has already been initialized by user interaction
        if hasattr(self, '_budget_initialized') and self._budget_initialized:
            return
            
        # Make sure we have a valid model selected first
        if self.model_combo.count() > 0 and self.model_combo.currentIndex() >= 0:
            if hasattr(self, 'budget_mode'):
                # Calculate smart token budget based on $1 USD default
                model_data = self.get_selected_model_data()
                if model_data and 'input_price' in model_data:
                    price_per_token = model_data['input_price'] / 1000.0
                    usd_equivalent_tokens = int(self.usd_budget / price_per_token) if price_per_token > 0 else 25000
                    
                    # Update token budget to be equivalent to $1 for current model
                    self.token_budget = usd_equivalent_tokens
                    
                    # Update display based on current mode
                    if self.budget_mode == "USD":
                        # USD mode: show USD value and token equivalent
                        self.budget_amount_input.setText(f"{self.usd_budget:.4f}")
                        self.budget_unit_equivalent_label.setText(f"USD ({usd_equivalent_tokens:,} tokens)")
                    else:  # Tokens mode
                        # Token mode: show token value and USD equivalent
                        self.budget_amount_input.setText(str(self.token_budget))
                        budget_cost = self.calculate_cost(self.token_budget)
                        self.budget_unit_equivalent_label.setText(f"tokens (${budget_cost:.4f})")
                    
                    # Mark as initialized so this doesn't run again
                    self._budget_initialized = True
    
    def update_budget_ranges(self):
        """Update budget range checking - now just validates on input since we use QLineEdit"""
        # Since we're using QLineEdit instead of QSpinBox, we don't set ranges here
        # Range checking is done in validate_budget_input() method
        if not self.budget_restrictions_enabled:
            print("DEBUG: Budget restrictions disabled - no range limits")
            return
        
        # Log the current range limits for debugging
        if self.budget_mode == "USD":
            print(f"DEBUG: USD mode - max allowed: ${self.max_usd_cap}")
        else:  # Tokens mode
            model_data = self.get_selected_model_data()
            if model_data and 'input_price' in model_data:
                price_per_token = model_data['input_price'] / 1000.0
                max_tokens = int(self.max_usd_cap / price_per_token) if price_per_token > 0 else 1000000
                
                # Limit to reasonable maximum to avoid UI issues
                MAX_REASONABLE_TOKENS = 1000000000  # 1 billion tokens max
                if max_tokens > MAX_REASONABLE_TOKENS:
                    max_tokens = MAX_REASONABLE_TOKENS
                    print(f"DEBUG: Capped max tokens to {MAX_REASONABLE_TOKENS:,} for UI stability")
                
                print(f"DEBUG: Token mode - max allowed: {max_tokens:,} tokens (price: ${price_per_token:.6f}/token, cap: ${self.max_usd_cap})")
            else:
                print(f"DEBUG: Token mode - fallback max: 1,000,000 tokens")
    
    def on_budget_restrictions_changed(self, state):
        """Handle budget restrictions checkbox change"""
        if state == 2:  # Checked (Qt.CheckState.Checked)
            # Show warning before disabling restrictions
            reply = QMessageBox.warning(
                self,
                "Disable Budget Restrictions",
                "⚠️ WARNING: Disabling budget restrictions removes all spending limits!\n\n"
                "This could result in unexpectedly high API costs if you accidentally "
                "set a very large budget or process many files.\n\n"
                "Are you sure you want to proceed?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.budget_restrictions_enabled = False
                self.update_budget_ranges()
                print("DEBUG: Budget restrictions disabled")
            else:
                # User cancelled - uncheck the box
                self.disable_budget_restrictions_checkbox.setChecked(False)
        else:
            # Re-enable restrictions
            self.budget_restrictions_enabled = True
            current_value = self.budget_amount_spin.value()
            self.update_budget_ranges()
            
            # If current value exceeds new range, adjust it
            if self.budget_mode == "USD" and current_value > self.max_usd_cap:
                self.budget_amount_spin.setValue(self.max_usd_cap)
                QMessageBox.information(
                    self,
                    "Budget Adjusted",
                    f"Budget has been reduced to the maximum allowed value of ${self.max_usd_cap}."
                )
            elif self.budget_mode == "Tokens":
                # Check against dynamic token limit
                model_data = self.get_selected_model_data()
                if model_data and 'input_price' in model_data:
                    price_per_token = model_data['input_price'] / 1000.0
                    max_tokens = int(self.max_usd_cap / price_per_token) if price_per_token > 0 else 1000000
                    if current_value > max_tokens:
                        self.budget_amount_spin.setValue(max_tokens)
                        QMessageBox.information(
                            self,
                            "Budget Adjusted",
                            f"Budget has been reduced to the maximum allowed value of {max_tokens:,} tokens "
                            f"(equivalent to ${self.max_usd_cap})."
                        )
            
            print("DEBUG: Budget restrictions enabled")
    
    def browse_course_files(self):
        """Browse for individual course material files"""
        file_dialog = QFileDialog()
        files, _ = file_dialog.getOpenFileNames(
            self,
            "Select Course Material Files",
            "",
            "All Supported Files (*.pdf *.docx *.odt *.txt *.md *.rtf *.pptx *.xlsx *.csv *.html *.py *.java *.cpp *.c *.js);;All Files (*)"
        )
        if files:
            self.add_course_materials(files)
    
    def browse_course_folder(self):
        """Browse for a folder containing course materials"""
        folder = QFileDialog.getExistingDirectory(self, "Select Course Materials Folder")
        if folder:
            # Get all supported files from the folder
            supported_extensions = {'.pdf', '.docx', '.odt', '.txt', '.md', '.rtf', 
                                  '.pptx', '.xlsx', '.csv', '.html', '.py', '.java', 
                                  '.cpp', '.c', '.js', '.png', '.jpg', '.jpeg'}
            
            files = []
            folder_path = Path(folder)
            for file_path in folder_path.rglob('*'):
                if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                    files.append(str(file_path))
            
            if files:
                self.add_course_materials(files)
            else:
                QMessageBox.information(self, "No Files Found", 
                                      "No supported course material files found in the selected folder.")
    
    def add_course_materials(self, file_paths):
        """Add course material files and update token count"""
        added_files = []
        
        for file_path in file_paths:
            if os.path.exists(file_path) and file_path not in self.course_materials:
                self.course_materials.append(file_path)
                added_files.append(file_path)
                
                # Add to tree widget with initial display
                filename = os.path.basename(file_path)
                item = QTreeWidgetItem(self.course_files_list)
                item.setText(0, filename)  # File name
                item.setText(1, "calculating...")  # Tokens
                item.setText(2, "calculating...")  # Cost
                item.setText(3, "⚪")  # Budget Impact - calculating
                item.setData(0, Qt.ItemDataRole.UserRole, file_path)
                item.setToolTip(0, file_path)
        
        if added_files:
            # Update token count in background
            QTimer.singleShot(100, self.calculate_tokens_async)
            
            self.log_two_step(f"Added {len(added_files)} course material file(s)")
    
    def remove_selected_course_file(self):
        """Remove the selected course material file"""
        current_item = self.course_files_list.currentItem()
        if current_item:
            file_path = current_item.data(0, Qt.ItemDataRole.UserRole)
            if file_path in self.course_materials:
                self.course_materials.remove(file_path)
            
            # Remove from token counts
            if hasattr(self, 'course_materials_token_counts') and file_path in self.course_materials_token_counts:
                del self.course_materials_token_counts[file_path]
            
            # Remove from processing queue if present
            if hasattr(self, 'processing_queue') and file_path in self.processing_queue:
                self.processing_queue.remove(file_path)
            
            # Remove item from tree widget
            index = self.course_files_list.indexOfTopLevelItem(current_item)
            if index >= 0:
                self.course_files_list.takeTopLevelItem(index)
            
            # Quick recalculation of totals (no file reprocessing needed)
            self.recalculate_totals()
    
    def clear_all_course_files(self):
        """Clear all course material files"""
        reply = QMessageBox.question(self, "Clear All Files", 
                                   "Are you sure you want to remove all course material files?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            self.course_materials.clear()
            self.course_files_list.clear()
            self.course_materials_tokens = 0
            self.course_materials_token_counts = {}
            self.course_materials_cost = 0.0
            
            # Clear processing queue
            if hasattr(self, 'processing_queue'):
                self.processing_queue.clear()
            
            self.update_token_display()
            self.course_materials_cost = 0.0
            self.course_materials_token_counts = {}  # Clear per-file counts
            self.update_token_display()
    
    def get_budget_impact_icon(self, file_tokens):
        """Get budget impact icon based on file's percentage of total budget (legacy token-based method)"""
        if not file_tokens or self.token_budget <= 0:
            return "⚪"  # No data
        
        percentage = (file_tokens / self.token_budget) * 100
        
        if percentage <= 5:
            return "🟢"  # Green - minimal impact (≤5%)
        elif percentage <= 15:
            return "🟡"  # Yellow - moderate impact (5-15%)
        else:
            return "🔴"  # Red - high impact (>15%)
    
    def get_budget_impact_icon_by_cost(self, cost):
        """Get budget impact icon based on cost percentage of total USD budget"""
        if not cost or cost <= 0:
            return "⚪"  # No data
        
        # Get budget amount in USD
        try:
            if self.budget_mode == "USD":
                budget_usd = self.usd_budget
            else:  # Token mode - convert to USD
                budget_usd = self.calculate_cost(self.token_budget)
        except (ValueError, AttributeError):
            return "⚪"  # No budget data
        
        if budget_usd <= 0:
            return "⚪"  # Invalid budget
        
        percentage = (cost / budget_usd) * 100
        
        if percentage <= 5:
            return "🟢"  # Green - minimal impact (≤5%)
        elif percentage <= 15:
            return "🟡"  # Yellow - moderate impact (5-15%)
        else:
            return "🔴"  # Red - high impact (>15%)
    
    def remove_selected_submission(self):
        """Remove the selected downloaded submission"""
        current_item = self.downloaded_submissions_list.currentItem()
        if current_item:
            # Check if this is the placeholder item
            if current_item.text(0) == "No submissions downloaded yet":
                QMessageBox.information(self, "No Selection", 
                                      "No submissions are available to remove.")
                return
            
            reply = QMessageBox.question(self, "Remove Submission", 
                                       f"Are you sure you want to remove the submission for '{current_item.text(0)}'?",
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                # Get the student name and remove from downloaded_submission_data
                student_name = current_item.text(0)
                
                # Remove from submission data if it exists
                if hasattr(self, 'downloaded_submission_data'):
                    self.downloaded_submission_data = [
                        sub for sub in self.downloaded_submission_data 
                        if sub['name'] != student_name
                    ]
                
                # Get the index and remove the item from UI
                index = self.downloaded_submissions_list.indexOfTopLevelItem(current_item)
                if index >= 0:
                    self.downloaded_submissions_list.takeTopLevelItem(index)
                    
                    # If no more items, show placeholder
                    if self.downloaded_submissions_list.topLevelItemCount() == 0:
                        placeholder_item = QTreeWidgetItem(self.downloaded_submissions_list)
                        placeholder_item.setText(0, "No submissions downloaded yet")
                        placeholder_item.setText(1, "Use Step 1 to download")
                        placeholder_item.setText(2, "-")
                        placeholder_item.setText(3, "-")
                        placeholder_item.setText(4, "-")
                        
                        # Reset step 2 button state
                        self.step2_button.setEnabled(False)
                        self.step2_status.setText("Ready (Complete Step 1 first)")
                        self.step2_status.setStyleSheet("color: #6c757d;")
                    
                    # Update budget summary with remaining submissions
                    self.update_budget_summary()
    
    def clear_all_submissions(self):
        """Clear all downloaded submissions"""
        if self.downloaded_submissions_list.topLevelItemCount() == 0:
            QMessageBox.information(self, "No Submissions", 
                                  "No submissions are available to clear.")
            return
        
        # Check if only placeholder exists
        if (self.downloaded_submissions_list.topLevelItemCount() == 1 and 
            self.downloaded_submissions_list.topLevelItem(0).text(0) == "No submissions downloaded yet"):
            QMessageBox.information(self, "No Submissions", 
                                  "No submissions are available to clear.")
            return
        
        reply = QMessageBox.question(self, "Clear All Submissions", 
                                   "Are you sure you want to remove all downloaded submissions?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            self.downloaded_submissions_list.clear()
            
            # Add placeholder back
            placeholder_item = QTreeWidgetItem(self.downloaded_submissions_list)
            placeholder_item.setText(0, "No submissions downloaded yet")
            placeholder_item.setText(1, "Use Step 1 to download")
            placeholder_item.setText(2, "-")
            placeholder_item.setText(3, "-")
            placeholder_item.setText(4, "-")
            
            # Reset step states
            self.step2_button.setEnabled(False)
            self.step2_status.setText("Ready (Complete Step 1 first)")
            self.step2_status.setStyleSheet("color: #6c757d;")
            
            # Update budget summary
            self.update_budget_summary()
    
    def update_budget_summary(self):
        """Update the budget summary based on current submissions"""
        total_tokens = 0
        total_cost = 0.0
        
        # Calculate totals from downloaded submission data
        if hasattr(self, 'downloaded_submission_data') and self.downloaded_submission_data:
            for sub_data in self.downloaded_submission_data:
                total_tokens += sub_data.get('tokens', 0)
                total_cost += sub_data.get('cost', 0.0)
        
        # Get budget amount
        try:
            budget_amount = float(self.budget_amount_input.text())
        except (ValueError, AttributeError):
            budget_amount = 1.0  # Default fallback
        
        remaining = budget_amount - total_cost
        within_budget = "✅ Within Budget" if remaining >= 0 else "⚠️ Over Budget"
        
        # Update the budget summary label
        if hasattr(self, 'budget_summary_label'):
            self.budget_summary_label.setText(
                f"📊 Total Usage: {total_tokens:,} tokens (${total_cost:.4f}) | "
                f"Budget: ${budget_amount:.2f} | Remaining: ${remaining:.4f} | {within_budget}"
            )
    
    def update_file_list_display(self):
        """Update the file list display with token counts and costs"""
        if not hasattr(self, 'course_files_list'):
            return
            
        # Get current model for cost calculation
        print(f"DEBUG: Updating file list display")
        
        # Update each item in the tree widget
        for i in range(self.course_files_list.topLevelItemCount()):
            item = self.course_files_list.topLevelItem(i)
            file_path = item.data(0, Qt.ItemDataRole.UserRole)
            
            if file_path in self.course_materials_token_counts:
                tokens = self.course_materials_token_counts[file_path]
                cost = self.calculate_cost(tokens)
                impact_icon = self.get_budget_impact_icon_by_cost(cost)  # Use cost-based method
                
                # Update item columns with token and cost info
                filename = os.path.basename(file_path)
                item.setText(0, filename)  # File name
                if tokens > 0:
                    item.setText(1, f"{tokens:,}")  # Tokens
                    item.setText(2, f"${cost:.4f}")  # Cost
                    item.setText(3, impact_icon)  # Budget Impact
                    print(f"DEBUG: Updated {filename}: {tokens:,} tokens, ${cost:.4f}")
                else:
                    item.setText(1, "processing...")  # Tokens
                    item.setText(2, "processing...")  # Cost
                    item.setText(3, "⚪")  # Processing impact
            else:
                # No token count yet
                filename = os.path.basename(file_path)
                item.setText(0, filename)  # File name
                item.setText(1, "calculating...")  # Tokens
                item.setText(2, "calculating...")  # Cost
    
    def calculate_tokens_async(self):
        """Calculate tokens for all course materials asynchronously with progressive updates"""
        if not self.course_materials:
            self.course_materials_tokens = 0
            self.course_materials_cost = 0.0
            self.update_token_display()
            return
        
        # Initialize processing state
        if not hasattr(self, 'processing_queue'):
            self.processing_queue = []
            self.processing_index = 0
        
        # Add new files to processing queue (avoid reprocessing)
        for file_path in self.course_materials:
            if file_path not in self.course_materials_token_counts and file_path not in self.processing_queue:
                self.processing_queue.append(file_path)
        
        # Start processing if not already running
        if self.processing_queue:
            QTimer.singleShot(10, self.process_next_file)
    
    def process_next_file(self):
        """Process the next file in the queue"""
        if not self.processing_queue:
            # All files processed, finalize
            self.finalize_token_calculation()
            return
        
        file_path = self.processing_queue.pop(0)
        
        try:
            print(f"DEBUG: Processing {os.path.basename(file_path)}...")
            
            # Extract text content from file (this is the heavy operation)
            content = self.extract_file_content(file_path)
            if content:
                # Estimate tokens
                tokens = self.estimate_tokens(content)
                self.course_materials_token_counts[file_path] = tokens
                print(f"DEBUG: Calculated {tokens:,} tokens for {os.path.basename(file_path)}")
            else:
                self.course_materials_token_counts[file_path] = 0
                print(f"DEBUG: No content extracted from {os.path.basename(file_path)}")
            
            # Update this specific file in the UI immediately
            self.update_single_file_display(file_path)
            
            # Update totals
            self.recalculate_totals()
            
        except Exception as e:
            print(f"ERROR: Error processing {file_path}: {e}")
            self.course_materials_token_counts[file_path] = 0
            self.update_single_file_display(file_path)
        
        # Process next file after a short delay to keep UI responsive
        if self.processing_queue:
            QTimer.singleShot(10, self.process_next_file)
        else:
            self.finalize_token_calculation()
    
    def update_single_file_display(self, file_path):
        """Update display for a single file that has been processed"""
        filename = os.path.basename(file_path)
        tokens = self.course_materials_token_counts.get(file_path, 0)
        cost = self.calculate_cost(tokens)
        impact_icon = self.get_budget_impact_icon(tokens)
        
        # Find the item in the tree widget and update it
        for i in range(self.course_files_list.topLevelItemCount()):
            item = self.course_files_list.topLevelItem(i)
            item_file_path = item.data(0, Qt.ItemDataRole.UserRole)
            
            if item_file_path == file_path:
                item.setText(1, f"{tokens:,}" if tokens > 0 else "0")
                item.setText(2, f"${cost:.4f}")
                item.setText(3, impact_icon)
                print(f"DEBUG: Updated {filename}: {tokens:,} tokens, ${cost:.4f}")
                break
    
    def recalculate_totals(self):
        """Recalculate total tokens and cost including course materials and submissions"""
        # Course materials tokens
        course_tokens = sum(self.course_materials_token_counts.values())
        self.course_materials_tokens = course_tokens
        self.course_materials_cost = self.calculate_cost(course_tokens)
        
        # Submission tokens (if available)
        submission_tokens = getattr(self, 'submission_tokens', 0)
        submission_cost = getattr(self, 'submission_cost', 0.0)
        
        # Combined totals for budget display
        self.total_tokens = course_tokens + submission_tokens
        self.total_cost = self.course_materials_cost + submission_cost
        
        print(f"DEBUG: Budget totals - Course: {course_tokens:,} tokens (${self.course_materials_cost:.4f}), Submissions: {submission_tokens:,} tokens (${submission_cost:.4f}), Total: {self.total_tokens:,} tokens (${self.total_cost:.4f})")
        
        self.update_token_display()
    
    def finalize_token_calculation(self):
        """Finalize token calculation after all files are processed"""
        print(f"DEBUG: Finalized token calculation for {len(self.course_materials)} files")
        total_tokens = sum(self.course_materials_token_counts.values())
        self.course_materials_tokens = total_tokens
        self.course_materials_cost = self.calculate_cost(total_tokens)
        self.update_token_display()
        
        # Clear processing state
        self.processing_queue = []
    
    def extract_file_content(self, file_path):
        """Calculate tokens for file based on full document processing (not just text extraction)"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path)
            
            # For token calculation, we estimate based on the full document being passed to AI
            # This matches the grading agent's behavior of passing complete files when possible
            
            if file_ext in ['.txt', '.md', '.py', '.java', '.cpp', '.c', '.js', '.html', '.css']:
                # Plain text files - read actual content for accurate token count
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
            elif file_ext == '.pdf':
                # PDF files - estimate tokens for full document (not just extracted text)
                # The grading agent will process the entire PDF with vision or enhanced text extraction
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        # Calculate based on full document processing, including formatting/structure
                        page_count = len(reader.pages)
                        # Base estimate: ~400-800 tokens per page, plus overhead for structure/formatting
                        estimated_tokens = max(800, page_count * 600 + 200)
                        # Cap at reasonable maximum while considering file size
                        max_tokens = max(estimated_tokens, min(file_size // 4, 50000))
                        return f"[PDF_TOKEN_ESTIMATE:{max_tokens}]"
                except ImportError:
                    # Fallback based on file size for full document processing
                    estimated_tokens = max(800, min(file_size // 4, 20000))
                    return f"[PDF_TOKEN_ESTIMATE:{estimated_tokens}]"
                    
            elif file_ext == '.docx':
                # Word documents - estimate for full document with formatting
                try:
                    from docx import Document
                    doc = Document(file_path)
                    # Get text content but add overhead for formatting/structure
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    # Add 30% overhead for formatting, tables, styles that the AI will process
                    base_tokens = len(text) // 4
                    full_doc_tokens = int(base_tokens * 1.3)
                    return f"[DOCX_TOKEN_ESTIMATE:{full_doc_tokens}]{text}"
                except ImportError:
                    # Fallback based on file size for full document
                    estimated_tokens = max(400, min(file_size // 8, 15000))
                    return f"[DOCX_TOKEN_ESTIMATE:{estimated_tokens}]"
                    
            elif file_ext == '.odt':
                # ODT files - estimate for full document (grading agent handles these)
                estimated_tokens = max(400, min(file_size // 10, 12000))
                return f"[ODT_TOKEN_ESTIMATE:{estimated_tokens}]"
                    
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                # Image files - vision processing tokens (much higher than text)
                # Vision models consume significant tokens for image analysis
                estimated_tokens = max(500, min(2000, file_size // 1000))
                return f"[IMAGE_TOKEN_ESTIMATE:{estimated_tokens}]"
                
            elif file_ext in ['.xlsx', '.xls', '.csv']:
                # Spreadsheet files - estimate for full data processing
                estimated_tokens = max(300, min(file_size // 15, 8000))
                return f"[SPREADSHEET_TOKEN_ESTIMATE:{estimated_tokens}]"
                
            elif file_ext in ['.pptx', '.ppt']:
                # Presentation files - estimate for full slide processing
                estimated_tokens = max(500, min(file_size // 12, 10000))
                return f"[PRESENTATION_TOKEN_ESTIMATE:{estimated_tokens}]"
                
            else:
                # Unsupported file type - conservative estimate for file processing
                estimated_tokens = max(200, min(file_size // 25, 5000))
                return f"[UNKNOWN_TOKEN_ESTIMATE:{estimated_tokens}]"
                
        except Exception as e:
            # Fallback estimate
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 1000
            estimated_tokens = max(300, min(file_size // 20, 3000))
            return f"[ERROR_TOKEN_ESTIMATE:{estimated_tokens}]"
    
    def estimate_tokens(self, content):
        """Estimate token count for content or file processing"""
        if not content:
            return 0
            
        # Check if this is a token estimate marker from file processing
        if isinstance(content, str) and content.startswith('[') and '_TOKEN_ESTIMATE:' in content:
            # Extract the token estimate from the marker
            try:
                start_marker = content.find('_TOKEN_ESTIMATE:') + len('_TOKEN_ESTIMATE:')
                end_marker = content.find(']', start_marker)
                if end_marker > start_marker:
                    estimated_tokens = int(content[start_marker:end_marker])
                    print(f"DEBUG: Using file-based token estimate: {estimated_tokens}")
                    return estimated_tokens
            except (ValueError, IndexError):
                pass
            
            # If extraction failed, fall through to text-based estimation
            # Remove the marker and process any remaining text
            marker_end = content.find(']') + 1
            if marker_end > 0 and marker_end < len(content):
                content = content[marker_end:]
            else:
                content = ""
        
        # Standard text-based token estimation
        if TIKTOKEN_AVAILABLE:
            try:
                # Use tiktoken for accurate counting
                encoding = tiktoken.encoding_for_model("gpt-4")
                return len(encoding.encode(str(content)))
            except Exception:
                # Fallback to approximation
                pass
        
        # Rough approximation: ~4 characters per token for English text
        return len(str(content)) // 4
    
    def calculate_submission_tokens_and_populate_list(self):
        """Calculate tokens for each downloaded submission and populate the budget UI list"""
        if not hasattr(self, 'downloaded_submission_data') or not self.downloaded_submission_data:
            # Initialize submission tracking variables if no data
            self.submission_tokens = 0
            self.submission_cost = 0.0
            return
        
        print(f"DEBUG: Calculating tokens for {len(self.downloaded_submission_data)} submissions")
        
        # Clear the current submissions list (except placeholder)
        self.downloaded_submissions_list.clear()
        
        total_tokens = 0
        total_cost = 0.0
        
        for sub_data in self.downloaded_submission_data:
            student_name = sub_data.get('name', 'Unknown Student')
            files = sub_data.get('files', [])
            
            # Calculate tokens for this submission
            submission_tokens = 0
            file_count = len(files)
            
            for file_path in files:
                if file_path and os.path.exists(file_path):
                    try:
                        # Extract content and estimate tokens
                        content = self.extract_file_content(file_path)
                        file_tokens = self.estimate_tokens(content)
                        submission_tokens += file_tokens
                        print(f"DEBUG: {student_name} - {os.path.basename(file_path)}: {file_tokens} tokens")
                    except Exception as e:
                        print(f"DEBUG: Error processing {file_path}: {e}")
                        # Fallback estimate
                        submission_tokens += 500  # Conservative estimate
            
            # Calculate cost for this submission
            submission_cost = self.calculate_cost(submission_tokens)
            
            # Update submission data with token info
            sub_data['tokens'] = submission_tokens
            sub_data['cost'] = submission_cost
            
            # Add to UI list
            item = QTreeWidgetItem(self.downloaded_submissions_list)
            item.setText(0, student_name)  # Student name
            item.setText(1, f"{file_count} files")  # File count
            item.setText(2, f"{submission_tokens:,}")  # Tokens
            item.setText(3, f"${submission_cost:.4f}")  # Cost
            
            # Budget impact icon
            impact_icon = self.get_budget_impact_icon_by_cost(submission_cost)
            item.setText(4, impact_icon)  # Budget Impact
            
            total_tokens += submission_tokens
            total_cost += submission_cost
        
        print(f"DEBUG: Total submission tokens: {total_tokens:,}, cost: ${total_cost:.4f}")
        
        # Store submission totals for budget calculations
        self.submission_tokens = total_tokens
        self.submission_cost = total_cost
        
        # Update budget display with combined totals
        self.recalculate_totals()
        
        self.log_two_step(f"💰 Calculated token costs: {total_tokens:,} tokens (${total_cost:.4f}) for {len(self.downloaded_submission_data)} submissions")
    
    def get_selected_model_name(self):
        """Get the currently selected model name for pricing"""
        current_index = self.model_combo.currentIndex()
        if current_index >= 0:
            model_data = self.model_combo.itemData(current_index)
            if model_data:
                model_id = model_data.get('base_model', model_data.get('name', 'gpt-4'))
                print(f"DEBUG: Selected model: {model_id} (index {current_index})")
                return model_id
        print(f"DEBUG: No model selected, using fallback: gpt-4")
        return "gpt-4"  # Default fallback
    
    def calculate_cost(self, tokens, model_name=None):
        """Calculate estimated cost based on token count and selected model's dynamic pricing"""
        # Get the current model data from the combo box
        current_index = self.model_combo.currentIndex()
        if current_index >= 0:
            model_data = self.model_combo.itemData(current_index)
            if model_data and isinstance(model_data, dict):
                # Use the dynamic pricing from the model manager
                input_price = model_data.get('input_price', 0.03)  # Default to GPT-4 pricing
                cost = tokens * (input_price / 1000)  # Convert from per-1K-tokens to per-token
                model_id = model_data.get('base_model', model_data.get('name', 'unknown'))
                print(f"DEBUG: Calculating cost for {tokens} tokens with model '{model_id}'")
                print(f"DEBUG: Using dynamic pricing {input_price:.6f} per 1K tokens = ${cost:.6f}")
                return cost
        
        # Fallback pricing if no model data available
        fallback_price = 0.03 / 1000  # GPT-4 pricing per token
        default_cost = tokens * fallback_price
        print(f"DEBUG: No model data available, using GPT-4 fallback pricing = ${default_cost:.6f}")
        return default_cost
    
    def clean_feedback_for_excel(self, feedback):
        """Clean feedback text for Excel export by removing unwanted sections"""
        if not feedback or not isinstance(feedback, str):
            return feedback
        
        lines = feedback.split('\n')
        cleaned_lines = []
        skip_next_lines = 0
        
        for i, line in enumerate(lines):
            # Skip lines if we're in a section to remove
            if skip_next_lines > 0:
                skip_next_lines -= 1
                continue
            
            # Remove the detailed rubric breakdown header section
            if line.strip() == "=" * 50 and i + 1 < len(lines) and "DETAILED RUBRIC BREAKDOWN" in lines[i + 1]:
                # Skip the current line, the header, and the next separator line
                skip_next_lines = 2
                continue
            
            # Remove grading method lines
            if line.strip().startswith("Grading Method:"):
                continue
            
            # Remove formatting preservation messages
            if "Original document formatting was preserved" in line:
                continue
            
            # Keep all other lines
            cleaned_lines.append(line)
        
        # Join back and clean up extra blank lines
        cleaned_feedback = '\n'.join(cleaned_lines)
        
        # Remove excessive blank lines (more than 2 consecutive)
        import re
        cleaned_feedback = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_feedback)
        
        return cleaned_feedback.strip()
    
    def update_token_display(self):
        """Update the token counter and cost display with combined totals"""
        # Use combined totals if available, otherwise fall back to course materials only
        tokens = getattr(self, 'total_tokens', self.course_materials_tokens)
        cost = getattr(self, 'total_cost', self.course_materials_cost)
        budget = self.token_budget
        percentage = (tokens / budget * 100) if budget > 0 else 0
        
        # Update counter
        self.token_counter_label.setText(f"{tokens:,} / {budget:,} tokens ({percentage:.1f}%)")
        
        # Update cost
        self.cost_estimate_label.setText(f"Estimated cost: ${cost:.4f}")
        
        # Update status with color coding
        if percentage <= 70:
            status = "🟢 Within budget"
            color = "#28a745"
        elif percentage <= 100:
            status = "🟡 Approaching limit"
            color = "#ffc107"
        else:
            status = "🔴 Over budget!"
            color = "#dc3545"
        
        self.token_status_label.setText(status)
        self.token_status_label.setStyleSheet(f"color: {color}; font-weight: bold;")
        
        # Update the integrated budget summary label with the status included
        budget_amount = f"${self.usd_budget:.4f}" if hasattr(self, 'usd_budget') else "$1.0000"
        remaining_cost = max(0, float(budget_amount.replace('$', '')) - cost)
        
        # Build detailed summary text showing breakdown
        course_tokens = self.course_materials_tokens
        submission_tokens = getattr(self, 'submission_tokens', 0)
        
        if submission_tokens > 0:
            summary_text = f"📊 Total: {tokens:,} tokens (${cost:.4f}) | Budget: {budget_amount} | Remaining: ${remaining_cost:.4f} | {status}"
        else:
            summary_text = f"📊 Total Usage: {tokens:,} tokens (${cost:.4f}) | Budget: {budget_amount} | Remaining: ${remaining_cost:.4f} | {status}"
        
        self.budget_summary_label.setText(summary_text)
        
        # Update counter color
        if percentage > 100:
            self.token_counter_label.setStyleSheet("color: #dc3545; font-weight: bold;")
        elif percentage > 90:
            self.token_counter_label.setStyleSheet("color: #ffc107; font-weight: bold;")
        else:
            self.token_counter_label.setStyleSheet("color: #28a745; font-weight: bold;")
        
        # Update single-step materials status display
        if hasattr(self, 'single_materials_status'):
            total_files = len(self.course_materials)
            if total_files > 0:
                status_text = f"📄 {total_files} file{'s' if total_files != 1 else ''} uploaded ({course_tokens:,} tokens)"
                self.single_materials_status.setText(status_text)
                # Update color based on budget status
                self.single_materials_status.setStyleSheet(f"color: {color}; font-size: 10pt;")
            else:
                self.single_materials_status.setText("No course materials uploaded")
                self.single_materials_status.setStyleSheet("color: #6c757d; font-size: 10pt;")

    def save_configuration(self):
        """Save unified configuration (OpenAI API key + Canvas credentials)"""
        try:
            from secure_key_manager import APIKeyManager
            
            # Get all configuration values
            openai_key = self.openai_key_entry.text().strip()
            canvas_url = self.canvas_url_entry.text().strip()
            canvas_token = self.canvas_token_entry.text().strip()
            
            # Validate inputs
            if not openai_key and not (canvas_url and canvas_token):
                QMessageBox.warning(self, "Warning", 
                                  "Please enter at least an OpenAI API key or Canvas credentials to save.")
                return
            
            # Don't save if OpenAI key is masked
            if openai_key and "*" in openai_key and len(openai_key.replace("*", "")) < 10:
                QMessageBox.warning(self, "Warning", 
                                  "Cannot save a masked OpenAI key. Please enter your full API key or leave it blank.")
                return
            
            # Don't save if Canvas token is masked
            if canvas_token and "*" in canvas_token and len(canvas_token.replace("*", "")) < 10:
                QMessageBox.warning(self, "Warning", 
                                  "Cannot save a masked Canvas token. Please enter your full API token or leave it blank.")
                return
            
            # Show password dialog for encryption
            password, ok = QInputDialog.getText(
                self, 
                "Set Master Password", 
                "Enter a master password to encrypt your configuration:\n\n"
                "This password will be required to load your saved settings.\n"
                "Please remember this password - it cannot be recovered!",
                QLineEdit.EchoMode.Password
            )
            
            if not ok or not password:
                QMessageBox.information(self, "Cancelled", "Configuration save cancelled by user.")
                return
            
            if len(password) < 4:
                QMessageBox.warning(self, "Password Too Short", 
                                  "Please enter a password with at least 4 characters for security.")
                return
            
            # Confirm password
            confirm_password, ok = QInputDialog.getText(
                self, 
                "Confirm Password", 
                "Please confirm your master password:",
                QLineEdit.EchoMode.Password
            )
            
            if not ok or confirm_password != password:
                QMessageBox.warning(self, "Password Mismatch", 
                                  "Passwords do not match. Please try again.")
                return
            
            # Create manager and define password callback
            def password_callback(action="unlock"):
                return password
            
            manager = APIKeyManager()
            success_messages = []
            
            # Save OpenAI API key if provided
            if openai_key:
                success = manager.save_openai_key(openai_key, password_callback)
                if success:
                    success_messages.append("✅ OpenAI API key saved")
                    self.openai_status.setText("🟢 OpenAI API configured")
                    self.openai_status.setStyleSheet("color: green; font-weight: bold;")
                    # Cache the newly saved key for this session
                    self._cached_openai_key = openai_key
                    # Don't clear the field - keep the manually entered key visible
                    # Only clear if it was already masked
                    # self.openai_key_entry.clear()
                else:
                    QMessageBox.warning(self, "Error", "Failed to save OpenAI API key.")
                    return
            
            # Save Canvas credentials if provided
            if canvas_url and canvas_token:
                success = manager.save_canvas_credentials(canvas_url, canvas_token, password_callback)
                if success:
                    success_messages.append("✅ Canvas credentials saved")
                    # Cache the newly saved credentials for this session
                    self._cached_canvas_url = canvas_url
                    self._cached_canvas_token = canvas_token
                else:
                    QMessageBox.warning(self, "Error", "Failed to save Canvas credentials.")
                    return
            
            # Show success message
            if success_messages:
                QMessageBox.information(self, "Configuration Saved", 
                                      "Configuration saved successfully!\n\n" + 
                                      "\n".join(success_messages) + 
                                      "\n\nYour settings are encrypted with your master password.\n" +
                                      "Click 'Load Configuration' to restore these settings later.")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error saving configuration: {str(e)}")
    
    def load_configuration(self):
        """Load unified configuration (OpenAI API key + Canvas credentials)"""
        try:
            from secure_key_manager import APIKeyManager
            
            # Create manager first to check if config exists
            temp_manager = APIKeyManager()
            
            if not temp_manager.key_manager.has_config():
                QMessageBox.information(self, "No Configuration", 
                                      "No saved configuration found. Please save configuration first.")
                return
            
            # Show password dialog immediately
            password, ok = QInputDialog.getText(
                self, 
                "Enter Master Password", 
                "Enter your master password to decrypt the saved configuration:",
                QLineEdit.EchoMode.Password
            )
            
            if not ok or not password:
                QMessageBox.information(self, "Cancelled", "Operation cancelled by user.")
                return
            
            # Define password callback function for GUI prompt
            def password_callback(action="unlock"):
                return password
            
            # Create manager
            manager = APIKeyManager()
            
            try:
                loaded_items = []
                
                # Try to load OpenAI API key
                if manager.has_openai_key():
                    try:
                        api_key = manager.get_openai_key(password_callback)
                        if api_key:
                            # Cache the real API key for later use
                            self._cached_openai_key = api_key
                            
                            # Show masked key in the field
                            if len(api_key) > 8:
                                masked_key = api_key[:4] + "*" * (len(api_key) - 8) + api_key[-4:]
                            else:
                                masked_key = "*" * len(api_key)
                            
                            self.openai_key_entry.setText(masked_key)
                            self.openai_status.setText("🟢 OpenAI API configured")
                            self.openai_status.setStyleSheet("color: green; font-weight: bold;")
                            loaded_items.append("✅ OpenAI API key loaded")
                            
                            # Load available models with the same password
                            self.refresh_models(password)
                    except Exception:
                        loaded_items.append("⚠️ OpenAI API key could not be loaded")
                
                # Try to load Canvas credentials
                try:
                    canvas_creds = manager.get_canvas_credentials(password_callback)
                    if canvas_creds.get('canvas_url') and canvas_creds.get('canvas_api_token'):
                        # Cache the real credentials for later use
                        self._cached_canvas_url = canvas_creds['canvas_url']
                        self._cached_canvas_token = canvas_creds['canvas_api_token']
                        
                        self.canvas_url_entry.setText(canvas_creds['canvas_url'])
                        # Show masked token
                        token = canvas_creds['canvas_api_token']
                        if len(token) > 8:
                            masked_token = token[:4] + "*" * (len(token) - 8) + token[-4:]
                        else:
                            masked_token = "*" * len(token)
                        self.canvas_token_entry.setText(masked_token)
                        loaded_items.append("✅ Canvas credentials loaded")
                except Exception:
                    loaded_items.append("⚠️ Canvas credentials could not be loaded")
                
                # Show results
                if loaded_items:
                    QMessageBox.information(self, "Configuration Loaded", 
                                          "Configuration loaded successfully!\n\n" + 
                                          "\n".join(loaded_items) + 
                                          "\n\nCredentials are masked for security.")
                else:
                    QMessageBox.warning(self, "No Data", 
                                      "No configuration data could be loaded.")
                    
            except Exception as decrypt_error:
                QMessageBox.warning(self, "Decryption Error", 
                                  f"Failed to decrypt configuration. Password may be incorrect.\n\nError: {str(decrypt_error)}")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error loading configuration: {str(e)}")
    
    def model_supports_file_uploads(self, model_id):
        """
        Check if a model supports file uploads for document grading
        
        Args:
            model_id (str): The model ID to check
            
        Returns:
            bool: True if model supports file uploads, False otherwise
        """
        # Convert to lowercase for case-insensitive matching
        model_lower = model_id.lower()
        
        # Check if the model is in our blocked list
        for blocked_model in MODELS_WITHOUT_FILE_SUPPORT:
            if blocked_model.lower() in model_lower:
                return False
        
        return True

    def refresh_models(self, existing_password=None):
        """Refresh available OpenAI models"""
        try:
            from secure_key_manager import APIKeyManager
            from openai_model_manager import OpenAIModelManager
            
            temp_manager = APIKeyManager()
            
            if not temp_manager.has_openai_key():
                QMessageBox.warning(self, "Warning", "Please configure OpenAI API key first.")
                return
            
            # Use existing password if provided, otherwise try cached key, then show password dialog
            api_key = None
            if existing_password:
                password = existing_password
            elif self._cached_openai_key:
                # Use cached API key directly
                api_key = self._cached_openai_key
            else:
                # Show password dialog immediately
                password, ok = QInputDialog.getText(
                    self, 
                    "Enter Master Password", 
                    "Enter your master password to access the API key:",
                    QLineEdit.EchoMode.Password
                )
                
                if not ok or not password:
                    QMessageBox.information(self, "Cancelled", "Operation cancelled by user.")
                    return
            
            # Get the API key securely with password callback or from cache
            try:
                if not api_key:
                    # Define password callback function
                    def password_callback(action="unlock"):
                        return password
                    
                    # Create manager
                    manager = APIKeyManager()
                    
                    api_key = manager.get_openai_key(password_callback)
                    if api_key:
                        # Cache the key for future use
                        self._cached_openai_key = api_key
                
                if not api_key:
                    QMessageBox.warning(self, "Warning", "Could not retrieve OpenAI API key.")
                    return
                
                # Create model manager and fetch models (force refresh to get latest)
                model_manager = OpenAIModelManager(api_key)
                all_models = model_manager.get_available_models(force_refresh=True)
                
                # Filter out models that don't support file uploads
                compatible_models = []
                filtered_count = 0
                
                for model in all_models:
                    model_id = model.get('base_model', model['name'])
                    if self.model_supports_file_uploads(model_id):
                        compatible_models.append(model)
                    else:
                        filtered_count += 1
                
                # Update the combo box with compatible models only
                self.model_combo.clear()
                current_model = None
                for model in compatible_models:
                    display_text = model.get('display_text', model['name'])
                    model_id = model.get('base_model', model['name'])  # Use base_model for API calls
                    # Store the complete model data (including pricing) instead of just model_id
                    self.model_combo.addItem(display_text, model)
                    
                    # Prioritize gpt-4o-mini as default, then fallback to others
                    if model_id == 'gpt-4o-mini':
                        current_model = display_text  # Always prefer gpt-4o-mini
                    elif current_model is None and model_id in ['gpt-4o', 'gpt-4']:
                        current_model = display_text  # Only use as fallback if gpt-4o-mini not found
                
                # Set default selection - prioritize gpt-4o-mini
                if current_model:
                    index = self.model_combo.findText(current_model)
                    if index >= 0:
                        self.model_combo.setCurrentIndex(index)
                        # Trigger model change event to update budget display
                        self.on_model_changed(index)
                
                # Update initial budget display now that models are loaded
                QTimer.singleShot(100, self.update_initial_budget_display)
                
                # Show success message with filtering information
                success_msg = f"Successfully loaded {len(compatible_models)} compatible models."
                if filtered_count > 0:
                    success_msg += f"\n\nFiltered out {filtered_count} older models that don't support document uploads."
                    success_msg += f"\n(Only models with document reading capability are shown)"
                
                QMessageBox.information(self, "Models Updated", success_msg)
                        
            except Exception as decrypt_error:
                QMessageBox.warning(self, "Decryption Error", 
                                  f"Failed to decrypt API key. Password may be incorrect.\n\nError: {str(decrypt_error)}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error refreshing models: {str(e)}")
    
    def test_connection(self):
        """Test Canvas connection with real API call"""
        try:
            url = self.canvas_url_entry.text().strip()
            token = self.canvas_token_entry.text().strip()
            
            if not url or not token:
                QMessageBox.warning(self, "Warning", "Please enter both Canvas URL and API token.")
                return
            
            # Check if token is masked - if so, get the real token from cache or secure storage
            if "*" in token and len(token.replace("*", "")) < 10:
                # Token appears to be masked, try to use cached credentials first
                if self._cached_canvas_url and self._cached_canvas_token:
                    # Use cached credentials - no password prompt needed
                    url = self._cached_canvas_url
                    token = self._cached_canvas_token
                else:
                    # No cached credentials, get from secure storage
                    try:
                        from secure_key_manager import APIKeyManager
                        
                        manager = APIKeyManager()
                        if not manager.key_manager.has_config():
                            QMessageBox.warning(self, "Warning", 
                                              "No saved configuration found. Please enter your full Canvas API token.")
                            return
                        
                        # Show password dialog to get real credentials
                        password, ok = QInputDialog.getText(
                            self, 
                            "Enter Master Password", 
                            "Enter your master password to access saved Canvas credentials:",
                            QLineEdit.EchoMode.Password
                        )
                        
                        if not ok or not password:
                            QMessageBox.information(self, "Cancelled", "Operation cancelled by user.")
                            return
                        
                        # Define password callback function
                        def password_callback(action="unlock"):
                            return password
                        
                        # Get real credentials
                        canvas_creds = manager.get_canvas_credentials(password_callback)
                        if canvas_creds.get('canvas_url') and canvas_creds.get('canvas_api_token'):
                            url = canvas_creds['canvas_url']
                            token = canvas_creds['canvas_api_token']
                            # Cache for future use
                            self._cached_canvas_url = url
                            self._cached_canvas_token = token
                        else:
                            QMessageBox.warning(self, "Warning", 
                                              "Could not retrieve Canvas credentials. Please enter them manually.")
                            return
                            
                    except Exception as cred_error:
                        QMessageBox.warning(self, "Error", 
                                          f"Failed to retrieve saved credentials: {str(cred_error)}\n\n"
                                          f"Please enter your full Canvas API token manually.")
                        return
            
            # Basic URL validation
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
                self.canvas_url_entry.setText(url)
            
            # Import Canvas API class
            try:
                from canvas_integration import CanvasAPI
            except ImportError:
                QMessageBox.critical(self, "Error", "Canvas integration module not found. Please ensure canvas_integration.py is available.")
                return
            
            # Create Canvas API instance and test connection
            try:
                canvas_api = CanvasAPI(url, token)
                
                # Test connection with actual API call
                courses = canvas_api.get_courses()
                
                # Success! Show connection details
                QMessageBox.information(self, "Connection Successful", 
                                      f"✅ Successfully connected to Canvas!\n\n"
                                      f"URL: {url}\n"
                                      f"Token: {'*' * (len(token) - 4) + token[-4:] if len(token) > 4 else '****'}\n"
                                      f"Courses found: {len(courses)}\n\n"
                                      f"Privacy protection: Student names will be anonymized for AI processing.")
                
                self.connection_status.setText("🟢 Connected to Canvas")
                self.connection_status.setStyleSheet("color: green; font-weight: bold;")
                
                # Store the Canvas API instance for later use
                self.canvas_api = canvas_api
                
                # Populate course dropdowns with the retrieved courses
                self.populate_course_dropdowns(courses)
                
                # Enable buttons that require Canvas connection
                self.step1_download_button.setEnabled(True)
                self.single_grade_button.setEnabled(True)
                
            except Exception as api_error:
                # Connection failed - show specific error
                error_msg = str(api_error)
                
                if "401" in error_msg or "unauthorized" in error_msg.lower():
                    QMessageBox.critical(self, "Connection Failed", 
                                       f"❌ Authentication failed.\n\n"
                                       f"Please check your API token:\n"
                                       f"• Token may be invalid or expired\n"
                                       f"• Token may not have required permissions\n\n"
                                       f"Error: {error_msg}")
                elif "404" in error_msg or "not found" in error_msg.lower():
                    QMessageBox.critical(self, "Connection Failed", 
                                       f"❌ Canvas URL not found.\n\n"
                                       f"Please check your Canvas URL:\n"
                                       f"• Ensure the URL is correct\n"
                                       f"• Include the full domain (e.g., yourschool.instructure.com)\n\n"
                                       f"Error: {error_msg}")
                elif "timeout" in error_msg.lower() or "connection" in error_msg.lower():
                    QMessageBox.critical(self, "Connection Failed", 
                                       f"❌ Network connection failed.\n\n"
                                       f"Please check:\n"
                                       f"• Your internet connection\n"
                                       f"• Canvas server availability\n"
                                       f"• Firewall settings\n\n"
                                       f"Error: {error_msg}")
                else:
                    QMessageBox.critical(self, "Connection Failed", 
                                       f"❌ Canvas connection failed.\n\n"
                                       f"Error: {error_msg}")
                
                self.connection_status.setText("� Canvas connection failed")
                self.connection_status.setStyleSheet("color: red; font-weight: bold;")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error during connection test: {str(e)}")
            self.connection_status.setText("🔴 Connection test error")
            self.connection_status.setStyleSheet("color: red; font-weight: bold;")
    
    def populate_course_dropdowns(self, courses):
        """Populate course dropdowns with retrieved courses from Canvas"""
        try:
            # Clear existing items from course dropdowns
            self.course_id_combo.clear()
            
            # Check if we have courses
            if not courses:
                self.course_id_combo.addItem("No courses found")
                return
            
            # Add placeholder for course selection
            self.course_id_combo.addItem("Please select a course", None)
            
            # Sort courses by name for easier selection
            sorted_courses = sorted(courses, key=lambda c: c.get('name', ''))
            
            # Add courses to dropdown
            for course in sorted_courses:
                course_id = course.get('id')
                course_name = course.get('name', 'Unnamed Course')
                
                # Create display text: "Course Name (ID: 12345)"
                display_text = f"{course_name} (ID: {course_id})"
                
                # Add item with display text, store course_id as data
                self.course_id_combo.addItem(display_text, course_id)
            
            # Set current index to the placeholder
            self.course_id_combo.setCurrentIndex(0)
            
            print(f"🦆 Populated course dropdown with {len(sorted_courses)} courses")
            
        except Exception as e:
            print(f"Error populating course dropdowns: {str(e)}")
            self.course_id_combo.clear()
            self.course_id_combo.addItem("Error loading courses")
    
    def on_course_selected(self):
        """Handle course selection and auto-populate assignments"""
        try:
            # Check if Canvas is connected
            if not hasattr(self, 'canvas_api') or not self.canvas_api:
                return
            
            # Get the selected course ID
            current_index = self.course_id_combo.currentIndex()
            course_id = self.course_id_combo.itemData(current_index)
            
            # Skip if no valid course selected (placeholder or None)
            if course_id is None or current_index == 0:
                self.assignment_combo.clear()
                self.assignment_combo.addItem("Select a course first")
                return
            
            # Populate assignments for the selected course
            self.populate_assignments_for_course(course_id)
            
        except Exception as e:
            print(f"Error in course selection: {str(e)}")
    
    def populate_assignments_for_course(self, course_id):
        """Populate assignments dropdown for a specific course"""
        try:
            # Clear assignments dropdown
            self.assignment_combo.clear()
            self.assignment_combo.addItem("Loading assignments...")
            
            # Get assignments from Canvas API
            assignments = self.canvas_api.get_assignments(course_id)
            
            # Clear and populate
            self.assignment_combo.clear()
            
            if not assignments:
                self.assignment_combo.addItem("No assignments found")
                return
            
            # Add placeholder
            self.assignment_combo.addItem("Please select an assignment", None)
            
            # Sort assignments by name
            sorted_assignments = sorted(assignments, key=lambda a: a.get('name', ''))
            
            # Add assignments to dropdown
            for assignment in sorted_assignments:
                assignment_id = assignment.get('id')
                assignment_name = assignment.get('name', 'Unnamed Assignment')
                
                # Create display text
                display_text = f"{assignment_name}"
                
                # Add item with display text, store assignment_id as data
                self.assignment_combo.addItem(display_text, assignment_id)
            
            # Set to placeholder
            self.assignment_combo.setCurrentIndex(0)
            
            print(f"🦆 Populated assignments dropdown with {len(sorted_assignments)} assignments")
            
        except Exception as e:
            print(f"Error populating assignments: {str(e)}")
            self.assignment_combo.clear()
            self.assignment_combo.addItem("Error loading assignments")
    
    def get_canvas_api(self):
        """Get Canvas API instance with current configuration"""
        try:
            canvas_url = self.canvas_url_entry.text().strip()
            canvas_api_key = self.canvas_token_entry.text().strip()
            
            if not canvas_url or not canvas_api_key:
                return None
            
            # Try to get API key from secure storage if input is masked
            if canvas_api_key.startswith('***'):
                try:
                    from secure_key_manager import SecureKeyManager
                    key_manager = SecureKeyManager()
                    canvas_api_key = key_manager.get_key('canvas_api_key')
                except:
                    return None
            
            if not canvas_api_key or canvas_api_key.startswith('***'):
                return None
            
            # Import Canvas API class
            from canvas_integration import CanvasAPI
            return CanvasAPI(canvas_url, canvas_api_key)
        except Exception as e:
            print(f"Error creating Canvas API: {e}")
            return None
    
    def set_status(self, message):
        """Set status message in the connection tab"""
        # Check if we have a status widget
        if hasattr(self, 'connection_status'):
            self.connection_status.setText(message)
        # Otherwise just print to console
        print(f"Status: {message}")
    
    def calculate_token_cost(self, tokens):
        """Calculate cost for given number of tokens based on current model"""
        try:
            # Get the current model data from the combo box (dynamic pricing)
            current_index = self.model_combo.currentIndex()
            if current_index >= 0:
                model_data = self.model_combo.itemData(current_index)
                if model_data and 'input_price' in model_data:
                    cost_per_1k = model_data['input_price'] / 1000.0  # Convert to per-token
                    cost = (tokens / 1000.0) * model_data['input_price']
                    print(f"DEBUG: Calculating cost for {tokens} tokens with model '{model_data.get('id', 'unknown')}'")
                    print(f"DEBUG: Using dynamic pricing {model_data['input_price']:.6f} per 1K tokens = ${cost:.6f}")
                    return cost
            
            # Fallback to static pricing if no model data available
            current_model = self.model_combo.currentText()
            
            # Token costs per 1K tokens (as of 2024) - fallback pricing
            model_costs = {
                'gpt-4o': 0.005,  # $5 per 1M tokens input
                'gpt-4o-mini': 0.00015,  # $0.15 per 1M tokens input
                'gpt-4-turbo': 0.01,  # $10 per 1M tokens input
                'gpt-4': 0.03,  # $30 per 1M tokens input
                'gpt-3.5-turbo': 0.0005,  # $0.5 per 1M tokens input
            }
            
            # Default to gpt-4o-mini pricing if model not found
            cost_per_1k = model_costs.get(current_model, model_costs['gpt-4o-mini'])
            cost = (tokens / 1000.0) * cost_per_1k
            print(f"DEBUG: Using fallback pricing for {current_model}: ${cost:.6f}")
            
            return cost
            
        except Exception as e:
            print(f"Error calculating token cost: {e}")
            return 0.001  # Default small cost

    def refresh_courses(self):
        """Refresh available courses from Canvas API"""
        try:
            # Get Canvas API instance
            canvas_api = self.get_canvas_api()
            if not canvas_api:
                QMessageBox.warning(self, "Canvas Connection", 
                                  "Canvas API not connected. Please configure Canvas API settings first.")
                return
            
            # Show loading message
            self.set_status("🔄 Loading courses from Canvas...")
            
            # Fetch courses from Canvas
            courses = canvas_api.get_courses()
            
            if not courses:
                QMessageBox.information(self, "No Courses", 
                                      "No courses found in your Canvas account.")
                return
            
            # Update course dropdown
            self.course_id_combo.clear()
            for course in courses:
                course_name = course.get('name', 'Unnamed Course')
                course_id = course.get('id', 'Unknown ID')
                display_text = f"{course_id} - {course_name}"
                self.course_id_combo.addItem(display_text, course)
            
            self.set_status(f"✅ Loaded {len(courses)} courses from Canvas")
            QMessageBox.information(self, "Courses Refreshed", 
                                  f"Successfully loaded {len(courses)} courses from Canvas API.\n\n"
                                  "You can now select a course to load its assignments.")
            
        except Exception as e:
            error_msg = f"Error refreshing courses: {str(e)}"
            self.set_status(f"❌ {error_msg}")
            QMessageBox.critical(self, "Error", error_msg)
    
    def enable_test_mode(self):
        """Enable test mode for grading buttons without Canvas connection"""
        reply = QMessageBox.question(self, "Enable Test Mode", 
                                   "Enable test mode to allow grading functionality without Canvas connection?\n\n"
                                   "⚠️ This is for testing purposes only. Actual grading will not work without "
                                   "proper Canvas connection.\n\n"
                                   "Enable test mode?", 
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            # Enable grading buttons
            self.step1_download_button.setEnabled(True)
            self.single_grade_button.setEnabled(True)
            
            # Update connection status
            self.connection_status.setText("🧪 Test mode enabled")
            self.connection_status.setStyleSheet("color: orange; font-weight: bold;")
            
            QMessageBox.information(self, "Test Mode Enabled", 
                                  "🧪 Test mode enabled!\n\n"
                                  "Grading buttons are now active for testing the interface.\n"
                                  "Remember: Actual Canvas functionality requires proper connection.")
    
    def refresh_assignments(self):
        """Refresh available assignments from Canvas"""
        try:
            # Check if Canvas is connected
            if not hasattr(self, 'canvas_api') or not self.canvas_api:
                QMessageBox.warning(self, "Warning", 
                                  "Please connect to Canvas first using the 'Connect' button.")
                return
            
            # Get course ID from the current tab's course field
            current_tab = self.tab_widget.currentIndex()
            
            if current_tab == 1:  # Two-Step Grading tab
                # Try to get course ID from current data (if dropdown selection)
                current_index = self.course_id_combo.currentIndex()
                if current_index >= 0 and self.course_id_combo.itemData(current_index):
                    course_id = str(self.course_id_combo.itemData(current_index))
                else:
                    # Fall back to manual entry text
                    course_id = self.course_id_combo.currentText().strip()
                    # Extract course ID from display text format "Course Name (ID: 12345)"
                    if "(ID: " in course_id and ")" in course_id:
                        try:
                            course_id = course_id.split("(ID: ")[1].split(")")[0]
                        except:
                            pass  # Keep original text if parsing fails
            elif current_tab == 2:  # Single-Step Grading tab
                course_id = self.single_course_id_entry.text().strip()
            else:
                # If called from connection tab or elsewhere, try two-step first
                current_index = self.course_id_combo.currentIndex()
                if current_index >= 0 and self.course_id_combo.itemData(current_index):
                    course_id = str(self.course_id_combo.itemData(current_index))
                else:
                    course_id = self.course_id_combo.currentText().strip()
                    # Extract course ID from display text format "Course Name (ID: 12345)"
                    if "(ID: " in course_id and ")" in course_id:
                        try:
                            course_id = course_id.split("(ID: ")[1].split(")")[0]
                        except:
                            pass  # Keep original text if parsing fails
                if not course_id:
                    course_id = getattr(self, 'single_course_id_entry', None)
                    if course_id:
                        course_id = course_id.text().strip()
            
            if not course_id or course_id in ["No courses loaded", "No courses found", "Error loading courses", "Please select a course", "Select a course first"]:
                QMessageBox.warning(self, "Warning", 
                                  "Please select a Course from the dropdown or enter a Course ID first.")
                return
            
            # Validate course ID is numeric
            if not course_id.isdigit():
                QMessageBox.warning(self, "Warning", 
                                  "Course ID must be a number. You can find this in your Canvas course URL.")
                return
            
            course_id = int(course_id)
            
            try:
                # Fetch real assignments from Canvas
                assignments = self.canvas_api.get_assignments(course_id)
                
                if not assignments:
                    QMessageBox.information(self, "No Assignments", 
                                          f"No assignments found for Course {course_id}.")
                    return
                
                # Clear and populate assignment dropdowns
                self.assignment_combo.clear()
                if hasattr(self, 'single_assignment_combo'):
                    self.single_assignment_combo.clear()
                
                # Add placeholder
                self.assignment_combo.addItem("Please select an assignment", None)
                if hasattr(self, 'single_assignment_combo'):
                    self.single_assignment_combo.addItem("Please select an assignment", None)
                
                assignment_details_list = []
                
                # Sort assignments by name
                sorted_assignments = sorted(assignments, key=lambda a: a.get('name', ''))
                
                for assignment in sorted_assignments:
                    assignment_name = assignment.get('name', 'Unnamed Assignment')
                    assignment_id = assignment.get('id', '')
                    due_at = assignment.get('due_at', 'No due date')
                    points_possible = assignment.get('points_possible', 'No points set')
                    
                    # Add to dropdowns with assignment_id as data
                    self.assignment_combo.addItem(assignment_name, assignment_id)
                    if hasattr(self, 'single_assignment_combo'):
                        self.single_assignment_combo.addItem(assignment_name, assignment_id)
                    
                    # Collect assignment details
                    assignment_details_list.append({
                        'name': assignment_name,
                        'id': assignment_id,
                        'due_at': due_at,
                        'points': points_possible,
                        'description': assignment.get('description', 'No description available')[:200] + '...' if assignment.get('description') else 'No description available'
                    })
                
                # Set to placeholder
                self.assignment_combo.setCurrentIndex(0)
                if hasattr(self, 'single_assignment_combo'):
                    self.single_assignment_combo.setCurrentIndex(0)
                
                # Update assignment info displays with the first assignment
                if assignment_details_list:
                    first_assignment = assignment_details_list[0]
                    assignment_details = (
                        f"Found {len(assignments)} assignments for Course {course_id}\n\n"
                        f"First assignment details:\n"
                        f"• Name: {first_assignment['name']}\n"
                        f"• ID: {first_assignment['id']}\n"
                        f"• Due: {first_assignment['due_at']}\n"
                        f"• Points: {first_assignment['points']}\n"
                        f"• Description: {first_assignment['description']}"
                    )
                else:
                    assignment_details = f"No assignment details available for Course {course_id}"
                
                # Assignment details functionality removed since it's not implemented
                # self.assignment_info.setText(assignment_details)
                # if hasattr(self, 'single_assignment_info'):
                #     self.single_assignment_info.setText(assignment_details)
                
                QMessageBox.information(self, "Assignments Loaded", 
                                      f"✅ Successfully loaded {len(assignments)} assignments for Course {course_id}!")
                
            except Exception as api_error:
                error_msg = str(api_error)
                
                if "404" in error_msg:
                    QMessageBox.critical(self, "Course Not Found", 
                                       f"❌ Course {course_id} not found.\n\n"
                                       f"Please check:\n"
                                       f"• Course ID is correct\n"
                                       f"• You have access to this course\n"
                                       f"• Course is published and active")
                elif "403" in error_msg or "unauthorized" in error_msg.lower():
                    QMessageBox.critical(self, "Access Denied", 
                                       f"❌ Access denied to Course {course_id}.\n\n"
                                       f"Please check:\n"
                                       f"• You are enrolled in this course\n"
                                       f"• Your API token has sufficient permissions")
                else:
                    QMessageBox.critical(self, "Error Loading Assignments", 
                                       f"❌ Failed to load assignments for Course {course_id}.\n\n"
                                       f"Error: {error_msg}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error refreshing assignments: {str(e)}")
    
    def browse_rubric(self):
        """Browse for rubric file"""
        import os
        default_dir = "rubrics" if os.path.exists("rubrics") else ""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "Select Rubric File", default_dir, "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.rubric_path_entry.setText(file_path)
    
    def browse_instructor_config(self):
        """Browse for instructor config file"""
        import os
        default_dir = "config" if os.path.exists("config") else ""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "Select Instructor Config", default_dir, "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.instructor_config_entry.setText(file_path)
    
    def browse_single_instructor_config(self):
        """Browse for instructor config file for single-step tab"""
        import os
        default_dir = "config" if os.path.exists("config") else ""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, 
            "Select Instructor Config File", 
            default_dir, 
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            self.single_instructor_entry.setText(file_path)
    
    def create_new_instructor_config(self):
        """Open the instructor configuration builder dialog"""
        builder = InstructorConfigBuilder(self)
        if builder.exec() == QDialog.DialogCode.Accepted:
            # If a config was saved, update the path in the appropriate combo box
            if hasattr(builder, 'saved_config_path'):
                # Check which tab is currently active to update the right field
                current_tab = self.tab_widget.currentIndex()
                if current_tab == 0:  # Two-Step Grading tab
                    self.instructor_config_entry.setText(builder.saved_config_path)
                elif current_tab == 1:  # Single-Step Grading tab
                    self.single_instructor_entry.setText(builder.saved_config_path)
    
    def on_rubric_source_changed(self):
        """Handle rubric source selection change for two-step grading"""
        if self.local_rubric_radio.isChecked():
            # Show local rubric widget, hide canvas rubric widget
            self.canvas_rubric_widget.hide()
            self.local_rubric_widget.show()
        else:
            # Hide local rubric widget, show canvas rubric widget
            self.local_rubric_widget.hide()
            self.canvas_rubric_widget.show()
    
    def on_single_rubric_source_changed(self):
        """Handle rubric source selection change for single-step grading"""
        if self.single_local_rubric_radio.isChecked():
            # Show local rubric widget, hide canvas rubric widget
            self.single_canvas_rubric_widget.hide()
            self.single_local_rubric_widget.show()
        else:
            # Hide local rubric widget, show canvas rubric widget
            self.single_local_rubric_widget.hide()
            self.single_canvas_rubric_widget.show()
    
    def browse_single_rubric(self):
        """Browse for rubric file in single-step grading"""
        import os
        default_dir = "rubrics" if os.path.exists("rubrics") else ""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "Select Rubric File", default_dir, "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.single_rubric_entry.setText(file_path)
    
    def start_step1_download(self):
        """Start Step 1: Download submissions and create folder structure"""
        if not self.validate_step1_inputs():
            return
        
        # Use the properly threaded implementation instead of blocking the GUI
        self.two_step_log.append("=== Step 1: Downloading Submissions ===")
        self.two_step_log.append("📥 Starting download of all submissions...")
        
        # Force GUI update
        from PyQt6.QtWidgets import QApplication
        QApplication.processEvents()
        
        # Call the threaded Step 1 implementation to avoid GUI blocking
        self.start_step1()
    
    def start_step3(self):
        """Start Step 3: Upload results (renamed from old start_step2)"""
        # Call the original start_step2 logic
        self.start_step2()
    
    def start_step2_grading(self):
        """Start Step 2: Grade the downloaded submissions (threaded)"""
        if not hasattr(self, 'downloaded_submission_data') or not self.downloaded_submission_data:
            QMessageBox.warning(self, "Step 2 Error", 
                              "No downloaded submissions found. Please complete Step 1 first.")
            return
        
        # Disable button during processing
        self.step2_button.setEnabled(False)
        self.step2_button.setText("🔄 Running Step 2...")
        
        # Get grading configuration
        rubric_path = self.rubric_path_entry.text().strip()
        instructor_config_path = self.instructor_config_entry.text().strip()
        use_canvas_rubric = self.canvas_rubric_radio.isChecked()
        
        if not use_canvas_rubric and not rubric_path:
            QMessageBox.warning(self, "Step 2 Error", 
                              "Please select a rubric file or enable Canvas rubric.")
            self.reset_step2_ui()
            return
        
        # Get OpenAI key and model
        if hasattr(self, '_cached_openai_key') and self._cached_openai_key:
            openai_key = self._cached_openai_key
            self.log_two_step("🔑 Using cached OpenAI API key")
        else:
            openai_key = self.openai_key_entry.text().strip()
            self.log_two_step("🔑 Using API key from text field")
        
        if not openai_key:
            QMessageBox.warning(self, "Step 2 Error", 
                              "OpenAI API key not configured. Please enter your API key.")
            self.reset_step2_ui()
            return
        
        # Get selected model
        current_index = self.model_combo.currentIndex()
        if current_index >= 0:
            model_data = self.model_combo.itemData(current_index)
            if isinstance(model_data, dict) and 'id' in model_data:
                selected_model = model_data['id']
            elif isinstance(model_data, str):
                selected_model = model_data
            else:
                selected_model = "gpt-4o-mini"
        else:
            selected_model = "gpt-4o-mini"
        
        # Get assignment/course info
        assignment_name = self.assignment_combo.currentText()
        course_name = self.course_id_combo.currentText()
        
        self.log_two_step("=== Step 2: AI Grading ===")
        self.log_two_step("🤖 Starting AI grading of downloaded submissions...")
        self.log_two_step(f"🤖 Using model: {selected_model}")
        self.log_two_step(f"📝 Rubric: {'Canvas rubric' if use_canvas_rubric else Path(rubric_path).name if rubric_path else 'None'}")
        if instructor_config_path and Path(instructor_config_path).exists():
            self.log_two_step(f"⚙️ Using instructor config: {Path(instructor_config_path).name}")
        
        # Get course materials data
        course_materials_files = self.course_materials.copy() if hasattr(self, 'course_materials') else []
        course_materials_instructions = self.course_materials_instructions.toPlainText().strip() if hasattr(self, 'course_materials_instructions') else ""
        
        # Log course materials information
        if course_materials_files:
            self.log_two_step(f"📚 Using {len(course_materials_files)} course material file(s)")
            if course_materials_instructions:
                self.log_two_step(f"📝 Course material instructions: {course_materials_instructions[:100]}..." if len(course_materials_instructions) > 100 else f"📝 Course material instructions: {course_materials_instructions}")
            else:
                self.log_two_step("📝 No specific instructions for course materials - using default guidance")
        else:
            self.log_two_step("📚 No course materials provided")
        
        # Create worker and thread
        self.step2_worker = GradingWorker(
            self.downloaded_submission_data, 
            self.download_folder,
            rubric_path,
            instructor_config_path,
            use_canvas_rubric,
            openai_key,
            selected_model,
            assignment_name,
            course_name,
            self.canvas_api,
            course_materials_files,
            course_materials_instructions,
            self.additional_grading_instructions.toPlainText().strip()
        )
        
        self.step2_thread = threading.Thread(target=self.step2_worker.run, daemon=True)
        
        # Connect worker signals to GUI update methods
        self.step2_worker.progress_updated.connect(self.update_step2_progress)
        self.step2_worker.log_message.connect(self.log_two_step)
        self.step2_worker.completed.connect(self.handle_step2_completion)
        self.step2_worker.error_occurred.connect(self.handle_step2_error)
        
        # Start the thread
        self.step2_thread.start()
    
    def update_step2_progress(self, percent, description):
        """Update Step 2 progress bar and description"""
        self.progress_step2.setValue(percent)
        self.progress_desc_step2.setText(description)
    
    def handle_step2_completion(self, results):
        """Handle Step 2 (Grading) completion"""
        if results['success']:
            self.update_step2_progress(100, "Step 2 grading complete!")
            
            graded_submissions = results.get('graded_submissions', [])
            grading_folder = results.get('grading_folder', '')
            total_submissions = results.get('total_submissions', 0)
            successful_submissions = results.get('successful_submissions', 0)
            
            self.log_two_step("✅ Step 2 completed successfully!")
            self.log_two_step(f"📁 Grading results saved to: {grading_folder}")
            self.log_two_step(f"📊 {successful_submissions}/{total_submissions} submissions graded successfully")
            self.log_two_step(f"🔒 Student identities remain anonymized in file system")
            
            # Update Step 2 status
            self.step2_status.setText("Completed")
            self.step2_status.setStyleSheet("color: green;")
            
            # Enable Step 3
            self.step3_button.setEnabled(True)
            self.step3_status.setText("Ready to upload")
            self.step3_status.setStyleSheet("color: green;")
            
            # Try to create Excel spreadsheet for review
            try:
                review_file = results.get('review_file')
                if review_file and Path(review_file).exists():
                    self.log_two_step(f"📊 Review spreadsheet created: {Path(review_file).name}")
                    
                    # Show Review Tab with the generated spreadsheet
                    try:
                        self.show_review_tab(review_file, grading_folder)
                        self.log_two_step("� Review Tab opened for editing grades")
                    except Exception as review_error:
                        self.log_two_step(f"⚠️ Could not open Review Tab: {review_error}")
                        
                else:
                    self.log_two_step("⚠️ No review spreadsheet was created")
                    
            except Exception as excel_error:
                self.log_two_step(f"⚠️ Error handling Excel file: {excel_error}")
            
            # Show completion message
            QMessageBox.information(self, "Step 2 Complete", 
                                  f"Grading completed!\n\n"
                                  f"Successfully graded: {successful_submissions}/{total_submissions} submissions\n"
                                  f"Results saved to: {grading_folder}\n"
                                  f"📊 Review spreadsheet available in Review tab")
            
        else:
            error_msg = results.get('message', 'Unknown error')
            self.log_two_step(f"Step 2 failed: {error_msg}")
            QMessageBox.critical(self, "Step 2 Error", 
                               f"Step 2 failed:\n{error_msg}")
        
        self.reset_step2_ui()
    
    def handle_step2_error(self, error_msg):
        """Handle Step 2 error"""
        self.log_two_step(f"ERROR in Step 2: {error_msg}")
        QMessageBox.critical(self, "Step 2 Error", f"Step 2 failed:\n{error_msg}")
        self.reset_step2_ui()
    
    def reset_step2_ui(self):
        """Reset Step 2 UI elements"""
        self.step2_button.setEnabled(True)
        self.step2_button.setText("🎯 Start Step 2: AI Grading")
        self.progress_step2.setValue(0)
        self.progress_desc_step2.setText("Ready to start...")
    

    def start_step1(self):
        """Step 1: Download submissions only (no grading) - now separated from grading"""
        if not self.validate_step1_inputs():
            return
        
        # Hide Review Tab if visible from previous session
        self.hide_review_tab()
        
        # Disable button during processing
        self.step1_download_button.setEnabled(False)
        self.step1_download_button.setText("🔄 Running Step 1...")
        
        # Get all the parameters needed for the worker
        course_index = self.course_id_combo.currentIndex()
        course_id = self.course_id_combo.itemData(course_index)
        course_name = self.course_id_combo.currentText()
        
        assignment_index = self.assignment_combo.currentIndex()
        assignment_id = self.assignment_combo.itemData(assignment_index)
        assignment_name = self.assignment_combo.currentText()
        
        # Store assignment/course info for Review Tab
        self.current_assignment_name = assignment_name
        self.current_course_name = course_name
        
        # Determine rubric settings
        use_canvas_rubric = self.canvas_rubric_radio.isChecked()
        rubric_path = None if use_canvas_rubric else self.rubric_path_entry.text().strip()
        
        # Get instructor configuration (optional)
        instructor_config_path = self.instructor_config_entry.text().strip()
        if not instructor_config_path or not Path(instructor_config_path).exists():
            instructor_config_path = None
        
        # Get API key and model
        # Use cached key if available (real key), otherwise fall back to text field
        if hasattr(self, '_cached_openai_key') and self._cached_openai_key:
            openai_key = self._cached_openai_key
            self.log_two_step("🔑 Using cached OpenAI API key")
        else:
            openai_key = self.openai_key_entry.text().strip()
            self.log_two_step("🔑 Using API key from text field")
            
        # Get the actual model ID (not display text) from combo box data
        current_index = self.model_combo.currentIndex()
        if current_index >= 0:
            model_data = self.model_combo.itemData(current_index)
            if isinstance(model_data, dict) and 'id' in model_data:
                selected_model = model_data['id']
            elif isinstance(model_data, str):
                selected_model = model_data
            else:
                selected_model = "gpt-4o-mini"
        else:
            selected_model = "gpt-4o-mini"
        
        # Debug: Check API key format (don't log the actual key)
        if openai_key and len(openai_key) > 20:
            self.log_two_step(f"🔍 API key format: {openai_key[:10]}...{openai_key[-10:]}")
        else:
            self.log_two_step(f"🔍 API key length: {len(openai_key) if openai_key else 0}")
        
        if not openai_key:
            QMessageBox.warning(self, "Missing API Key", "Please enter your OpenAI API key.")
            self.reset_step1_ui()
            return
        
        # Get course materials data
        course_materials_files = self.course_materials.copy() if hasattr(self, 'course_materials') else []
        course_materials_instructions = self.course_materials_instructions.toPlainText().strip() if hasattr(self, 'course_materials_instructions') else ""
        
        # Log course materials info
        if course_materials_files:
            self.log_two_step(f"📚 Using {len(course_materials_files)} course material file(s)")
            self.log_two_step(f"💰 Course materials token count: {self.course_materials_tokens:,}")
            self.log_two_step(f"💵 Estimated additional cost: ${self.course_materials_cost:.3f}")
            if course_materials_instructions:
                self.log_two_step(f"📝 Course material instructions: {course_materials_instructions[:100]}..." if len(course_materials_instructions) > 100 else f"📝 Course material instructions: {course_materials_instructions}")
            else:
                self.log_two_step("📝 No specific instructions for course materials - using default guidance")
        else:
            self.log_two_step("📚 No course materials provided")
        
        # Create worker and thread - use DownloadOnlyWorker for Step 1
        self.step1_worker = DownloadOnlyWorker(
            self.canvas_api, course_id, course_name, assignment_id, assignment_name,
            use_canvas_rubric, rubric_path, instructor_config_path, openai_key, selected_model,
            course_materials_files, course_materials_instructions
        )
        
        self.step1_thread = threading.Thread(target=self.step1_worker.run, daemon=True)
        
        # Connect worker signals to GUI update methods
        self.step1_worker.progress_updated.connect(self.update_step1_progress)
        self.step1_worker.log_message.connect(self.log_two_step)
        self.step1_worker.completed.connect(self.handle_step1_completion)
        self.step1_worker.error_occurred.connect(self.handle_step1_error)
        
        # Start the thread
        self.step1_thread.start()
    
    def handle_step1_error(self, error_msg):
        """Handle Step 1 error"""
        self.log_two_step(f"ERROR in Step 1: {error_msg}")
        QMessageBox.critical(self, "Step 1 Error", f"Step 1 failed:\n{error_msg}")
        self.reset_step1_ui()
    
    def validate_step1_inputs(self):
        """Validate inputs before starting Step 1"""
        # Check Canvas connection
        if not hasattr(self, 'canvas_api') or not self.canvas_api:
            QMessageBox.warning(self, "Validation Error", "Please connect to Canvas first.")
            return False
        
        # Check course selection
        current_index = self.course_id_combo.currentIndex()
        if current_index <= 0 or not self.course_id_combo.itemData(current_index):
            QMessageBox.warning(self, "Validation Error", "Please select a course.")
            return False
        
        # Check assignment selection
        assignment_index = self.assignment_combo.currentIndex()
        if assignment_index <= 0 or not self.assignment_combo.itemData(assignment_index):
            QMessageBox.warning(self, "Validation Error", "Please select an assignment.")
            return False
        
        # Validate rubric source selection
        if self.local_rubric_radio.isChecked():
            rubric_path = self.rubric_path_entry.text().strip()
            if not rubric_path:
                QMessageBox.warning(self, "Validation Error", "Please specify a rubric file path.")
                return False
            if not Path(rubric_path).exists():
                QMessageBox.warning(self, "Validation Error", f"Rubric file not found: {rubric_path}")
                return False
        
        return True
    
    def update_step1_progress(self, percent, description):
        """Update Step 1 progress bar and description"""
        self.progress_step1_download.setValue(percent)
        self.progress_desc_step1_download.setText(description)
    
    def handle_step1_completion(self, results):
        """Handle Step 1 (Download Only) completion"""
        if results['success']:
            self.update_step1_progress(100, "Step 1 download complete!")
            
            # Store downloaded submission data for Step 2
            if 'submission_data' in results:
                self.downloaded_submission_data = results['submission_data']
                self.download_folder = Path(results['folder_path'])
                
                # Calculate token costs for downloaded submissions and populate the UI
                self.calculate_submission_tokens_and_populate_list()
                
                self.log_two_step("✓ Step 1 (Download) completed successfully!")
                self.log_two_step(f"Downloaded {results.get('submission_count', 0)} submissions")
                self.log_two_step(f"Download folder: {results['folder_path']}")
                self.log_two_step("📋 Rubric and instructor config saved for Step 2")
                self.log_two_step("🔒 Student mapping saved (identities anonymized)")
                
                # Enable Step 2 grading
                self.step2_button.setEnabled(True)
                self.step2_status.setText("Ready to grade")
                self.step2_status.setStyleSheet("color: green;")
                
                # Update Step 1 status
                self.step1_download_status.setText("Completed")
                self.step1_download_status.setStyleSheet("color: green;")
                
                self.log_two_step("Next: Run Step 2 to grade the downloaded submissions")
                
                # Store results for reference
                self.current_step1_results = results
                
                # Show simple completion notification
                QMessageBox.information(self, "Step 1 Complete", 
                                      f"✅ Step 1 (Download) completed successfully!\n\n"
                                      f"• Downloaded {results.get('submission_count', 0)} submissions\n"
                                      f"• Folder: {results['folder_path']}\n"
                                      f"• Rubric and configuration saved\n"
                                      f"• Student identities anonymized\n\n"
                                      f"You can now run Step 2 to grade the submissions.")
            else:
                self.log_two_step("⚠️ No submission data found in results")
                QMessageBox.warning(self, "Step 1 Warning", 
                                  "Step 1 completed but no submission data was found.")
        else:
            error_msg = results.get('message', 'Unknown error')
            self.log_two_step(f"Step 1 failed: {error_msg}")
            QMessageBox.critical(self, "Step 1 Error", 
                               f"Step 1 failed:\n{error_msg}")
        
        self.reset_step1_ui()
    
    def show_completion_dialog(self, results):
        """Show completion dialog with option to open review folder"""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton
        import os
        import subprocess
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Step 1 Complete")
        dialog.setModal(True)
        dialog.resize(400, 200)
        
        layout = QVBoxLayout()
        
        # Success message
        message = QLabel(f"✅ Step 1 completed successfully!\n\n"
                        f"• Graded {len(results.get('grading_results', {}))} submissions\n"
                        f"• Results saved to: {results['folder_path']}\n"
                        f"• Review spreadsheet created: {results.get('review_file', 'N/A')}\n"
                        f"• Review Tab opened for easy editing\n\n"
                        f"Use the Review Tab to view submissions and edit scores/comments.\n"
                        f"When finished, run Step 2 to upload grades.")
        message.setWordWrap(True)
        layout.addWidget(message)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        open_folder_btn = QPushButton("Open Review Folder Now")
        open_folder_btn.clicked.connect(lambda: self.open_review_folder(results['folder_path']))
        button_layout.addWidget(open_folder_btn)
        
        ok_btn = QPushButton("OK")
        ok_btn.clicked.connect(dialog.accept)
        ok_btn.setDefault(True)
        button_layout.addWidget(ok_btn)
        
        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec()
    
    def open_review_folder(self, folder_path, checked=False):
        """Open the review folder in Windows Explorer"""
        import os
        import subprocess
        
        try:
            if os.path.exists(folder_path):
                # Use Windows Explorer to open the folder
                subprocess.run(['explorer', folder_path])
                self.log_two_step(f"📁 Opened review folder: {folder_path}")
            else:
                QMessageBox.warning(self, "Folder Not Found", f"Could not find folder: {folder_path}")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to open folder: {str(e)}")
    
    def reset_step1_ui(self):
        """Reset Step 1 UI elements"""
        self.step1_download_button.setEnabled(True)
        self.step1_download_button.setText("🔽 Start Step 1: Download Submissions")
        self.progress_step1_download.setValue(0)
        self.progress_desc_step1_download.setText("Ready to start...")
    
    def log_two_step(self, message):
        """Add message to two-step grading log"""
        current_text = self.two_step_log.toPlainText()
        if current_text:
            new_text = current_text + "\n" + message
        else:
            new_text = message
        self.two_step_log.setPlainText(new_text)
        # Scroll to bottom
        scrollbar = self.two_step_log.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
        # Force GUI update to ensure log display updates immediately
        from PyQt6.QtWidgets import QApplication
        QApplication.processEvents()
    
    def refresh_gui(self):
        """Force GUI refresh - call this during long operations to keep UI responsive"""
        from PyQt6.QtWidgets import QApplication
        QApplication.processEvents()
        
        # Also update the scrollbar to ensure log stays at bottom
        if hasattr(self, 'two_step_log'):
            scrollbar = self.two_step_log.verticalScrollBar()
            scrollbar.setValue(scrollbar.maximum())
    
    def start_step2(self):
        """Start Step 3: Upload results to Canvas"""
        # Validate that Step 2 has been completed
        if not self.step2_button.isEnabled() or self.step2_status.text() != "Ready to grade":
            QMessageBox.warning(self, "Step 3 Not Ready", 
                              "Step 3 requires Step 2 (AI Grading) to be completed first.\n\n"
                              "Please run Step 2 to generate grades before uploading to Canvas.")
            return
        
        # Check for review results
        reply = QMessageBox.question(self, "Upload Grades to Canvas", 
                                   "⚠️ Step 3: Upload Results to Canvas\n\n"
                                   "This will:\n"
                                   "• Upload final grades to Canvas\n"
                                   "• Add AI-generated feedback as comments\n"
                                   "• Mark grading as complete\n\n"
                                   "Important: Make sure you have reviewed the grades from Step 2.\n\n"
                                   "Proceed with upload?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        self.two_step_log.append("=== Step 3: Uploading Results to Canvas ===")
        
        try:
            # Get Canvas API instance
            canvas_api = self.get_canvas_api()
            if not canvas_api:
                QMessageBox.critical(self, "Canvas Error", 
                                   "Canvas API not available. Please check your Canvas configuration.")
                return
            
            # Get course and assignment info
            current_assignment = self.assignment_combo.itemData(self.assignment_combo.currentIndex())
            if not current_assignment:
                QMessageBox.critical(self, "Upload Error", "No assignment selected.")
                return
                
            assignment_id = current_assignment.get('id')
            course_id = self.course_id_combo.currentText().split(' - ')[0]
            
            self.two_step_log.append(f"📚 Assignment: {current_assignment.get('name', 'Unknown')}")
            self.two_step_log.append(f"🏫 Course ID: {course_id}")
            
            # Update Step 3 UI
            self.step3_button.setEnabled(False)
            self.step3_status.setText("Uploading...")
            self.step3_status.setStyleSheet("color: orange;")
            
            # TODO: Implement actual grade upload logic
            # This would need to integrate with the existing grading results
            # For now, show a completion message
            
            self.two_step_log.append("⚠️ Upload functionality requires integration with Step 2 results")
            self.two_step_log.append("This feature will be completed when grading workflow is fully integrated")
            
            # Show success message
            QMessageBox.information(self, "Upload Status", 
                                  "Upload functionality is being implemented.\n\n"
                                  "This will integrate with the existing grading workflow "
                                  "to upload grades and feedback to Canvas.")
            
            # Reset UI
            self.step3_button.setEnabled(True)
            self.step3_status.setText("Ready")
            self.step3_status.setStyleSheet("color: green;")
            
        except Exception as e:
            self.two_step_log.append(f"❌ Error in Step 3: {str(e)}")
            QMessageBox.critical(self, "Upload Error", f"Error uploading to Canvas: {str(e)}")
            self.step3_button.setEnabled(True)
            self.step3_status.setText("Error")
            self.step3_status.setStyleSheet("color: red;")
    
    def start_single_grading(self):
        """Start single-step grading"""
        # Validate rubric source selection
        if self.single_local_rubric_radio.isChecked():
            rubric_path = self.single_rubric_entry.text().strip()
            if not rubric_path:
                QMessageBox.warning(self, "Validation Error", "Please specify a rubric file path.")
                return
            if not Path(rubric_path).exists():
                QMessageBox.warning(self, "Validation Error", f"Rubric file not found: {rubric_path}")
                return
            rubric_info = f"Using local rubric: {rubric_path}"
        else:
            rubric_info = "Using Canvas rubric (will be downloaded automatically)"
        
        # Course materials info
        total_files = len(self.course_materials)
        materials_info = f"Course materials: {total_files} file(s) uploaded ({self.course_materials_tokens:,} tokens)" if total_files > 0 else "No course materials uploaded"
        
        # Show confirmation dialog with course materials info
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Icon.Information)
        msg.setWindowTitle("Single Grading")
        msg.setText("⚠️ Single-Step Grading (Legacy Mode)\n\n"
                   "This interface is now optimized for the 1-2-3 Grading workflow.\n\n"
                   "For the best experience, please use the Two-Step Grading tab:\n"
                   "• Step 1: Download & Calculate Costs\n"
                   "• Step 2: AI Grading with Review\n"
                   "• Step 3: Upload Results\n\n"
                   "Configuration would have been:\n"
                   f"• {rubric_info}\n"
                   f"• {materials_info}\n\n"
                   "Single-step mode bypasses cost calculation and review steps.")
        
        # Add course materials context display if materials are available
        if total_files > 0:
            msg.setDetailedText("Course Materials Context:\n" + 
                              "\n".join([f"• {os.path.basename(path)}" for path in self.course_materials]))
        
        msg.exec()
    
    # Placeholder methods for button click handlers
    def browse_rubric(self):
        """Browse for rubric file"""
        import os
        default_dir = "rubrics" if os.path.exists("rubrics") else ""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, 
            "Select Rubric File", 
            default_dir, 
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            self.rubric_path_entry.setText(file_path)
    
    def browse_instructor_config(self):
        """Browse for instructor config file"""
        import os
        default_dir = "config" if os.path.exists("config") else ""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, 
            "Select Instructor Config File", 
            default_dir, 
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            self.instructor_config_entry.setText(file_path)
    
    # ========================================
    # Review Tab Functionality Methods
    # ========================================
    
    def show_review_tab(self, review_spreadsheet_path, submission_folder_path):
        """Show the Review Tab after Step 1 completion"""
        # Store paths as instance attributes for later use
        self.review_spreadsheet_path = review_spreadsheet_path
        self.submission_folder_path = submission_folder_path
        
        # Try to extract assignment and course info if not already set
        if not hasattr(self, 'current_assignment_name') or not hasattr(self, 'current_course_name'):
            print("DEBUG: Trying to extract assignment/course info from UI or file path")
            
            # Try to get from UI dropdowns first
            if hasattr(self, 'assignment_combo') and self.assignment_combo.currentIndex() > 0:
                self.current_assignment_name = self.assignment_combo.currentText()
                print(f"DEBUG: Got assignment from UI: {self.current_assignment_name}")
            else:
                # Try to extract from file path
                import os
                file_name = os.path.basename(review_spreadsheet_path)
                if "_REVIEW.xlsx" in file_name:
                    self.current_assignment_name = file_name.replace("_REVIEW.xlsx", "").replace("_", " ")
                    print(f"DEBUG: Extracted assignment from filename: {self.current_assignment_name}")
                else:
                    self.current_assignment_name = "Review Assignment"
                    print("DEBUG: Using fallback assignment name")
            
            if hasattr(self, 'course_id_combo') and self.course_id_combo.currentIndex() > 0:
                self.current_course_name = self.course_id_combo.currentText()
                print(f"DEBUG: Got course from UI: {self.current_course_name}")
            else:
                self.current_course_name = "Review Course"
                print("DEBUG: Using fallback course name")
        
        if not self.review_tab_visible:
            # Create the review tab
            self.review_tab = self.create_review_tab(review_spreadsheet_path, submission_folder_path)
            
            # Add the review tab after the two-step tab (position 2)
            self.tab_widget.insertTab(2, self.review_tab, "👁️ Review")
            self.review_tab_visible = True
            
            # Debug: Verify combo box exists after tab creation
            print(f"DEBUG: After tab creation - hasattr combo box: {hasattr(self, 'review_submission_combo')}")
            if hasattr(self, 'review_submission_combo'):
                print(f"DEBUG: Combo box type: {type(self.review_submission_combo)}")
                print(f"DEBUG: Combo box is None: {self.review_submission_combo is None}")
            
            # Switch to the review tab first
            self.tab_widget.setCurrentIndex(2)
            
            # Force a GUI update to ensure the tab is fully rendered
            from PyQt6.QtWidgets import QApplication
            QApplication.processEvents()
            
            # Now that the tab is created, rendered, and added, load the data
            self.load_review_data()
            
            self.log_two_step("✅ Review Tab opened - you can now review and edit the graded submissions")
    
    def hide_review_tab(self):
        """Hide the Review Tab when starting a new grading session"""
        if self.review_tab_visible:
            # Find and remove the review tab
            for i in range(self.tab_widget.count()):
                if self.tab_widget.tabText(i) == "👁️ Review":
                    self.tab_widget.removeTab(i)
                    break
            
            self.review_tab_visible = False
            self.review_tab = None
    
    def load_review_data(self):
        """Load data from the review spreadsheet (primary) or graded results (fallback)"""
        try:
            import pandas as pd
            import os
            import json
            
            # Primary: Try to load from Excel spreadsheet first
            if hasattr(self, 'review_spreadsheet_path') and os.path.exists(self.review_spreadsheet_path):
                print("DEBUG: Using Excel spreadsheet as primary data source")
                self.load_review_data_from_spreadsheet()
                return
            
            # Fallback 1: Load from saved graded results if available in memory
            if hasattr(self, 'graded_results') and self.graded_results:
                print("DEBUG: Fallback to graded results from memory")
                graded_submissions = self.graded_results
                self.load_review_data_from_graded_results(graded_submissions)
                return
            
            # Fallback 2: Try to load from saved graded results JSON
            graded_results_path = None
            if hasattr(self, 'submission_folder_path') and self.submission_folder_path:
                results_folder = Path(self.submission_folder_path)
                graded_results_path = results_folder / "graded_results.json"
                print(f"DEBUG: Looking for graded results at: {graded_results_path}")
                
                if graded_results_path and graded_results_path.exists():
                    print("DEBUG: Fallback to graded results from JSON file")
                    with open(graded_results_path, 'r', encoding='utf-8') as f:
                        graded_submissions = json.load(f)
                    self.load_review_data_from_graded_results(graded_submissions)
                    return
            
            # No data source available
            QMessageBox.warning(self, "File Not Found", 
                              f"No review data found.\n"
                              f"Expected spreadsheet: {getattr(self, 'review_spreadsheet_path', 'Not set')}\n"
                              f"Or graded results JSON: {graded_results_path or 'Not available'}")
            return

        except Exception as e:
            print(f"DEBUG: Error loading review data: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"Failed to load review data: {str(e)}")

    def load_review_data_from_spreadsheet(self):
        """Load review data from Excel spreadsheet (primary data source)"""
        try:
            import pandas as pd
            import os
            import json
            
            print("DEBUG: Loading data from Excel spreadsheet")
            # Read the spreadsheet
            df = pd.read_excel(self.review_spreadsheet_path)
            
            # Debug: Print column names and first few rows
            print(f"DEBUG: Spreadsheet columns: {list(df.columns)}")
            print(f"DEBUG: First few rows:")
            print(df.head())
            
            # Load student name mapping (real_name -> anon_name) for file access
            student_name_to_anon_mapping = {}
            mapping_file = Path(self.submission_folder_path).parent / "student_mapping.json"
            if mapping_file.exists():
                try:
                    with open(mapping_file, 'r', encoding='utf-8') as f:
                        mapping_data = json.load(f)
                    # Handle both old and new mapping file formats
                    name_map = mapping_data.get('name_map', mapping_data) if isinstance(mapping_data, dict) else mapping_data
                    # Create reverse mapping: real_name -> anon_name
                    for anon_name, data in name_map.items():
                        if isinstance(data, dict) and 'real_name' in data:
                            real_name = data['real_name']
                            student_name_to_anon_mapping[real_name] = anon_name
                    print(f"DEBUG: Loaded student name mapping with {len(student_name_to_anon_mapping)} entries")
                except Exception as mapping_error:
                    print(f"DEBUG: Could not load student mapping: {mapping_error}")
            else:
                print(f"DEBUG: Student mapping file not found: {mapping_file}")
            
            # Extract assignment information from first row if available
            if len(df) > 0:
                assignment_name = getattr(self, 'current_assignment_name', 'Unknown Assignment')
                course_name = getattr(self, 'current_course_name', 'Unknown Course')
                self.review_assignment_info.setText(
                    f"Assignment: {assignment_name} | Course: {course_name} | "
                    f"Total Submissions: {len(df)}"
                )
            
            # Store the data
            self.review_data = {}
            self.review_original_data = {}
            
            for index, row in df.iterrows():
                # Map to actual spreadsheet column names based on debug output
                student_id = str(row.get('Canvas_User_ID', row.get('Student ID', row.get('student_id', row.get('ID', f'student_{index}')))))
                
                # Get student name from actual column (this is the REAL name now)
                real_student_name = str(row.get('Student_Name', 
                                         row.get('Student Name', 
                                               row.get('student_name', 
                                                     row.get('Name', f'Student {index + 1}')))))
                
                # Map real name to anonymized name for finding files
                anon_student_name = student_name_to_anon_mapping.get(real_student_name, real_student_name)
                print(f"DEBUG: Real name '{real_student_name}' -> Anon name '{anon_student_name}'")
                
                # Get score from AI_Score column
                score = row.get('AI_Score', row.get('Score', row.get('score', row.get('Points', 0))))
                
                # Get comments from AI_Comments column
                comments = str(row.get('AI_Comments', 
                                     row.get('Comments', 
                                           row.get('comments', 
                                                 row.get('Feedback', 
                                                       row.get('feedback', ''))))))
                
                # Find submission files using anonymized name
                submission_files = self.find_submission_files(anon_student_name, student_id)
                
                print(f"DEBUG: Row {index}: ID={student_id}, Real Name={real_student_name}, Anon Name={anon_student_name}, Score={score}, Files={len(submission_files) if submission_files else 0}")
                
                # Store review data using REAL name as key (for display in Excel)
                self.review_data[real_student_name] = {
                    'student_id': student_id,
                    'real_name': real_student_name,
                    'anon_name': anon_student_name,  # Store both names
                    'score': score,
                    'comments': comments,
                    'submission_files': submission_files,
                    'changed': False
                }
                
                # Store original data for comparison
                self.review_original_data[real_student_name] = {
                    'score': score,
                    'comments': comments
                }
            
            # Populate the student list in the UI
            self.populate_review_student_list()
            
            print(f"DEBUG: Successfully loaded review data for {len(self.review_data)} students")
            
        except Exception as e:
            print(f"DEBUG: Error loading review data from spreadsheet: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"Failed to load review data from spreadsheet: {str(e)}")

    def load_review_data_from_graded_results(self, graded_submissions):
        """Load review data directly from graded submissions data"""
        try:
            import json
            from pathlib import Path
            
            print("DEBUG: Loading review data from graded submissions")
            
            # Load student mapping to get real names
            student_mapping = {}
            if hasattr(self, 'submission_folder_path') and self.submission_folder_path:
                # Look for mapping files in parent folder
                root_folder = Path(self.submission_folder_path).parent
                mapping_file = root_folder / "student_mapping.json"
                submission_data_file = root_folder / "submission_data.json"
                
                if mapping_file.exists():
                    with open(mapping_file, 'r', encoding='utf-8') as f:
                        mapping_data = json.load(f)
                    name_map = mapping_data.get('name_map', mapping_data) if isinstance(mapping_data, dict) else mapping_data
                    for anon_name, data in name_map.items():
                        if isinstance(data, dict) and 'real_name' in data:
                            student_mapping[anon_name] = {
                                'real_name': data['real_name'],
                                'user_id': data.get('user_id', '')
                            }
                elif submission_data_file.exists():
                    # Fallback to submission data
                    with open(submission_data_file, 'r', encoding='utf-8') as f:
                        submission_data = json.load(f)
                    for sub in submission_data:
                        anon_name = sub.get('name', '')
                        student_mapping[anon_name] = {
                            'real_name': sub.get('real_name', anon_name),
                            'user_id': sub.get('user_id', '')
                        }
                
                print(f"DEBUG: Loaded student mapping with {len(student_mapping)} entries")
            
            # Store the data for review
            self.review_data = {}
            self.review_original_data = {}
            
            # Get assignment info
            assignment_name = getattr(self, 'current_assignment_name', getattr(self, 'assignment_name', 'Unknown Assignment'))
            course_name = getattr(self, 'current_course_name', getattr(self, 'course_name', 'Unknown Course'))
            
            successful_submissions = [sub for sub in graded_submissions if sub.get('status') == 'graded']
            print(f"DEBUG: Found {len(successful_submissions)} successfully graded submissions")
            
            for sub in successful_submissions:
                anon_name = sub.get('name', '')
                mapping_info = student_mapping.get(anon_name, {})
                real_name = mapping_info.get('real_name', anon_name)
                user_id = mapping_info.get('user_id', sub.get('id', ''))
                
                score = sub.get('score', sub.get('grading_result', {}).get('overall_score', 0))
                comments = sub.get('feedback', sub.get('grading_result', {}).get('overall_feedback', ''))
                
                # Find submission files using the refactored method
                submission_files = self.find_submission_files(anon_name, user_id)
                
                print(f"DEBUG: Processing {real_name} (anon: {anon_name}): score={score}, files={len(submission_files) if submission_files else 0}")
                
                # Store review data using REAL name as key
                self.review_data[real_name] = {
                    'student_id': user_id,
                    'real_name': real_name,
                    'anon_name': anon_name,
                    'score': score,
                    'comments': comments,
                    'submission_files': submission_files,
                    'changed': False
                }
                
                # Store original data for comparison
                self.review_original_data[real_name] = {
                    'score': score,
                    'comments': comments
                }
            
            # Set assignment info
            self.review_assignment_info.setText(
                f"Assignment: {assignment_name} | Course: {course_name} | "
                f"Total Submissions: {len(self.review_data)}"
            )
            
            # Populate the student list in the UI
            self.populate_review_student_list()
            
            print(f"DEBUG: Successfully loaded review data from graded results for {len(self.review_data)} students")
            
        except Exception as e:
            print(f"DEBUG: Error loading review data from graded results: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"Failed to load review data from graded results: {str(e)}")

    def populate_review_student_list(self):
        """Populate the student list in the review tab"""
        try:
            # Debug: Check widget existence
            print(f"DEBUG: populate_review_student_list called")
            print(f"DEBUG: hasattr(self, 'review_submission_combo'): {hasattr(self, 'review_submission_combo')}")
            if hasattr(self, 'review_submission_combo'):
                print(f"DEBUG: self.review_submission_combo is None: {self.review_submission_combo is None}")
                print(f"DEBUG: self.review_submission_combo type: {type(self.review_submission_combo)}")
                print(f"DEBUG: bool(self.review_submission_combo): {bool(self.review_submission_combo)}")
                
                # Check if widget is valid/accessible
                try:
                    print(f"DEBUG: Widget objectName: {self.review_submission_combo.objectName()}")
                    print(f"DEBUG: Widget isVisible: {self.review_submission_combo.isVisible()}")
                    print(f"DEBUG: Widget count (should be 0 initially): {self.review_submission_combo.count()}")
                except Exception as widget_error:
                    print(f"DEBUG: Error accessing widget properties: {widget_error}")
            
            # More specific check
            if not hasattr(self, 'review_submission_combo'):
                print("DEBUG: review_submission_combo attribute not found")
                return
            
            if self.review_submission_combo is None:
                print("DEBUG: review_submission_combo is None")
                return
            
            # Try to access the widget to see if it's valid
            try:
                current_count = self.review_submission_combo.count()
                print(f"DEBUG: Current combo box count: {current_count}")
            except Exception as access_error:
                print(f"DEBUG: Cannot access combo box, might be deleted: {access_error}")
                return
                
            self.review_submission_combo.clear()
            print(f"DEBUG: Successfully cleared combo box")
            
            if not hasattr(self, 'review_data') or not self.review_data:
                print("DEBUG: No review data available")
                return
            
            # Add students to the combo box
            items_added = 0
            for student_name in sorted(self.review_data.keys()):
                student_data = self.review_data[student_name]
                student_id = student_data.get('student_id', student_name)
                
                # Create display text with name and score
                score = student_data.get('score', 0)
                display_text = f"{student_name} (Score: {score})"
                
                # Add item with student name as display and student_id as data
                self.review_submission_combo.addItem(display_text, student_id)
                items_added += 1
            
            print(f"DEBUG: Added {items_added} items to combo box")
            print(f"DEBUG: Combo box now has {self.review_submission_combo.count()} items")
            print(f"DEBUG: Populated student combo box with {len(self.review_data)} students")
            
            # Select first student if available
            if self.review_submission_combo.count() > 0:
                self.review_submission_combo.setCurrentIndex(0)
                print("DEBUG: Set initial selection to index 0")
                # Manually trigger the load since setting index 0 might not fire signal
                self.load_current_submission()
                print("DEBUG: Manually triggered load_current_submission")
            else:
                print("DEBUG: No items in combo box to select")
                
        except Exception as e:
            print(f"DEBUG: Error populating student list: {e}")
            import traceback
            traceback.print_exc()
                
        except Exception as e:
            print(f"DEBUG: Error populating student list: {e}")
            import traceback
            traceback.print_exc()
    
    def find_submission_files(self, student_name, student_id):
        """Find submission files for a student using stored submission data"""
        try:
            import json
            from pathlib import Path
            
            # Load submission data from the download process
            if not hasattr(self, '_submission_data_cache'):
                self._submission_data_cache = {}
                # Try to load from submission_data.json in the root folder
                root_folder = Path(self.submission_folder_path).parent if Path(self.submission_folder_path).name == 'results' else Path(self.submission_folder_path).parent
                submission_data_file = root_folder / "submission_data.json"
                print(f"DEBUG: find_submission_files - Loading submission data from: {submission_data_file}")
                
                if submission_data_file.exists():
                    with open(submission_data_file, 'r', encoding='utf-8') as f:
                        submission_data = json.load(f)
                    
                    # Index by anonymized name for fast lookup
                    for sub in submission_data:
                        anon_name = sub.get('name', '')
                        self._submission_data_cache[anon_name] = sub
                    print(f"DEBUG: find_submission_files - Loaded submission data for {len(self._submission_data_cache)} students")
                else:
                    print(f"DEBUG: find_submission_files - submission_data.json not found at: {submission_data_file}")
            
            # Look up the student data directly
            student_data = self._submission_data_cache.get(student_name, {})
            if student_data:
                files = student_data.get('files', [])
                print(f"DEBUG: find_submission_files - Found {len(files)} files for {student_name} from submission data")
                return files
            else:
                print(f"DEBUG: find_submission_files - No submission data found for {student_name}")
                return []
                
        except Exception as e:
            print(f"DEBUG: find_submission_files - Error loading submission data: {e}")
            return []


    
    def load_current_submission(self):
        """Load the currently selected submission"""
        print("DEBUG: load_current_submission called")
        
        if not self.review_data:
            print("DEBUG: No review data available")
            return
        
        # Get current selection
        current_index = self.review_submission_combo.currentIndex()
        print(f"DEBUG: Current combo box index: {current_index}")
        
        if current_index < 0:
            print("DEBUG: No valid selection")
            return
        
        student_id = self.review_submission_combo.itemData(current_index)
        print(f"DEBUG: Selected student ID: {student_id}")
        
        # Find the student data by ID (since review_data is keyed by student name)
        student_data = None
        student_name_key = None
        for name, data in self.review_data.items():
            if str(data.get('student_id', '')) == str(student_id):
                student_data = data
                student_name_key = name
                break
        
        if not student_data:
            print(f"DEBUG: No student data found for ID: {student_id}")
            print(f"DEBUG: Available student IDs: {[data.get('student_id', 'N/A') for data in self.review_data.values()]}")
            return
        
        print(f"DEBUG: Found student data for: {student_name_key}")
        
        # Update score
        score = student_data.get('score', 0)
        print(f"DEBUG: Setting score to: {score}")
        if hasattr(self, 'review_score_entry'):
            self.review_score_entry.setText(str(int(score)))
        
        # Update comments (block signals to avoid triggering change detection)
        comments = student_data.get('comments', '')
        print(f"DEBUG: Setting comments (length: {len(comments)})")
        if hasattr(self, 'review_comments_editor'):
            self.review_comments_editor.blockSignals(True)
            self.review_comments_editor.setPlainText(comments)
            self.review_comments_editor.blockSignals(False)
        
        # Load submission content
        submission_files = student_data.get('submission_files', [])
        print(f"DEBUG: Loading submission files: {len(submission_files)} files")
        if submission_files:
            self.load_submission_content_from_files(submission_files)
        else:
            print("DEBUG: No submission files found")
            if hasattr(self, 'review_submission_viewer'):
                self.review_submission_viewer.setPlainText(f"No submission files found for {student_name_key}")
        
        # Update navigation buttons
        total_submissions = len(self.review_data)
        current_pos = current_index + 1
        
        # With wraparound navigation, enable buttons when there are multiple submissions
        if hasattr(self, 'review_prev_btn'):
            self.review_prev_btn.setEnabled(total_submissions > 1)
        if hasattr(self, 'review_next_btn'):
            self.review_next_btn.setEnabled(total_submissions > 1)
        if hasattr(self, 'review_progress_label'):
            self.review_progress_label.setText(f"{current_pos} of {total_submissions}")
        
        # Update current index
        self.review_current_index = current_index
        
        # Enable the "View In Directory" button
        if hasattr(self, 'review_view_directory_btn'):
            self.review_view_directory_btn.setEnabled(True)
        
        # Check if this submission has unsaved changes
        if hasattr(self, 'update_save_button_state'):
            self.update_save_button_state()
            
        print("DEBUG: load_current_submission completed")
    
    def load_submission_content_from_files(self, submission_files):
        """Load submission content from a list of files"""
        print(f"DEBUG: load_submission_content_from_files called with {len(submission_files)} files")
        
        try:
            if not submission_files:
                if hasattr(self, 'review_submission_viewer'):
                    self.review_submission_viewer.setPlainText("No submission files available")
                if hasattr(self, 'review_file_info'):
                    self.review_file_info.setText("No files")
                return
            
            # Store files for navigation
            self.current_submission_files = submission_files
            self.current_submission_index = 0
            
            # Load the first file
            self.load_current_submission_file()
            
            # Update file navigation if there are multiple files
            if len(submission_files) > 1 and hasattr(self, 'review_file_nav_label'):
                self.review_file_nav_label.setText(f"File 1 of {len(submission_files)}")
            
        except Exception as e:
            print(f"DEBUG: Error in load_submission_content_from_files: {e}")
            if hasattr(self, 'review_submission_viewer'):
                self.review_submission_viewer.setPlainText(f"Error loading submission files: {str(e)}")
    
    def on_document_load_finished(self, success, web_scrollbar_js):
        """Handle document load completion with error checking."""
        print(f"DEBUG: Document load finished. Success: {success}")
        if success:
            # Inject scrollbar styling on successful load
            self.review_document_viewer.page().runJavaScript(web_scrollbar_js)
        else:
            print("DEBUG: Document load failed")
    
    def load_submission_content(self, submission_file):
        """Load and display submission content from student's specific submission folder"""
        try:
            # Get current student info
            current_index = self.review_submission_combo.currentIndex()
            student_id = self.review_submission_combo.itemData(current_index) if current_index >= 0 else None
            current_student_data = self.review_data.get(student_id, {}) if student_id else {}
            student_name = current_student_data.get('student_name', '')
            
            print(f"DEBUG: Loading submissions for student: {student_name} (ID: {student_id})")
            print(f"DEBUG: Base submission folder: '{self.review_submission_folder}'")
            
            # Find the student's specific submission folder using the Student_XXX_YYYY pattern
            submissions_folder = os.path.join(self.review_submission_folder, 'submissions')
            print(f"DEBUG: Looking for student folders in: {submissions_folder}")
            
            student_submission_folder = None
            if os.path.exists(submissions_folder):
                # Look for folders matching Student_XXX_YYYY pattern containing this student's ID
                for folder_name in os.listdir(submissions_folder):
                    folder_path = os.path.join(submissions_folder, folder_name)
                    if os.path.isdir(folder_path):
                        # Check if this folder contains the student ID
                        if f"_{student_id}" in folder_name or student_id in folder_name:
                            student_submission_folder = folder_path
                            print(f"DEBUG: Found student submission folder: {student_submission_folder}")
                            break
            
            if not student_submission_folder:
                # Try alternative: look for folders with student name patterns
                if os.path.exists(submissions_folder):
                    name_parts = student_name.lower().replace(' ', '_').split('_')
                    for folder_name in os.listdir(submissions_folder):
                        folder_path = os.path.join(submissions_folder, folder_name)
                        if os.path.isdir(folder_path):
                            folder_lower = folder_name.lower()
                            # Check if any part of the student name is in the folder name
                            if any(part in folder_lower for part in name_parts if len(part) > 2):
                                student_submission_folder = folder_path
                                print(f"DEBUG: Found student submission folder by name pattern: {student_submission_folder}")
                                break
            
            if not student_submission_folder:
                self.review_file_info.setText(f"No submission folder found for {student_name}")
                self.review_submission_viewer.setPlainText(f"No submission folder found for {student_name} (ID: {student_id}).\n\nLooked in: {submissions_folder}\n\nExpected pattern: Student_XXX_{student_id}")
                print(f"DEBUG: No submission folder found for student {student_name}")
                return
            
            # Get all submission files in the student's folder
            submission_files = []
            supported_extensions = ['.pdf', '.docx', '.odt', '.txt', '.doc', '.py', '.java', '.cpp', '.c', '.js', '.html', '.css', '.md', '.rtf']
            
            for file_name in os.listdir(student_submission_folder):
                file_path = os.path.join(student_submission_folder, file_name)
                if os.path.isfile(file_path):
                    file_ext = os.path.splitext(file_name)[1].lower()
                    if file_ext in supported_extensions:
                        submission_files.append(file_path)
            
            print(f"DEBUG: Found {len(submission_files)} submission files: {[os.path.basename(f) for f in submission_files]}")
            
            # Initialize submission navigation if not already done
            if not hasattr(self, 'current_submission_files'):
                self.current_submission_files = []
                self.current_submission_index = 0
            
            self.current_submission_files = submission_files
            self.current_submission_index = 0
            
            # Load the first file or show message if no files
            if submission_files:
                self.load_current_submission_file()
            else:
                # No files found - check for text submission in submission data
                self.load_text_submission_fallback(student_submission_folder)
                
        except Exception as e:
            self.review_file_info.setText(f"Error loading submissions for {student_name}")
            self.review_submission_viewer.setPlainText(f"Error loading submissions: {str(e)}")
            print(f"DEBUG: Error in load_submission_content: {e}")
    
    def load_current_submission_file(self):
        """Load the currently selected submission file"""
        print(f"DEBUG: load_current_submission_file called")
        
        if not hasattr(self, 'current_submission_files') or not self.current_submission_files:
            print("DEBUG: No current submission files")
            return
            
        if self.current_submission_index >= len(self.current_submission_files):
            print("DEBUG: Invalid submission index")
            return
            
        current_file = self.current_submission_files[self.current_submission_index]
        file_name = os.path.basename(current_file)
        
        print(f"DEBUG: Loading file: {file_name}")
        
        try:
            # Clear previous content first
            if hasattr(self, 'review_submission_viewer'):
                self.review_submission_viewer.clear()
            
            # Clear web viewer if available
            if hasattr(self, 'review_document_viewer') and hasattr(self, 'WEB_ENGINE_AVAILABLE'):
                if self.WEB_ENGINE_AVAILABLE and self.review_document_viewer:
                    self.review_document_viewer.setHtml("")
            
            # Update file info 
            total_files = len(self.current_submission_files)
            file_size = os.path.getsize(current_file)
            file_info_text = f"File: {file_name} ({file_size:,} bytes)"
            
            if hasattr(self, 'review_file_info'):
                self.review_file_info.setText(file_info_text)
            
            print(f"DEBUG: File info: {file_info_text}")
            
            # Update navigation controls if they exist
            if total_files > 1:
                if hasattr(self, 'review_prev_file_btn'):
                    self.review_prev_file_btn.show()
                    self.review_prev_file_btn.setEnabled(self.current_submission_index > 0)
                if hasattr(self, 'review_next_file_btn'):
                    self.review_next_file_btn.show()
                    self.review_next_file_btn.setEnabled(self.current_submission_index < total_files - 1)
                if hasattr(self, 'review_file_counter'):
                    self.review_file_counter.show()
                    self.review_file_counter.setText(f"{self.current_submission_index + 1} of {total_files}")
            else:
                # Hide navigation for single files
                if hasattr(self, 'review_prev_file_btn'):
                    self.review_prev_file_btn.hide()
                if hasattr(self, 'review_next_file_btn'):
                    self.review_next_file_btn.hide()
                if hasattr(self, 'review_file_counter'):
                    self.review_file_counter.hide()
            
            # Determine file type and load content
            file_ext = os.path.splitext(file_name)[1].lower()
            
            print(f"DEBUG: File extension: {file_ext}")
            
            # Determine file rendering capabilities
            can_render_natively = file_ext in ['.pdf'] and WEB_ENGINE_AVAILABLE
            can_convert_to_html = file_ext in ['.docx', '.odt'] and WEB_ENGINE_AVAILABLE
            
            print(f"DEBUG: can_render_natively={can_render_natively}, file_ext={file_ext}")
            print(f"DEBUG: can_convert_to_html={can_convert_to_html}")
            print(f"DEBUG: WEB_ENGINE_AVAILABLE={WEB_ENGINE_AVAILABLE}")
            
            # Update view mode buttons availability
            if hasattr(self, 'view_mode_rendered_btn') and self.view_mode_rendered_btn:
                self.view_mode_rendered_btn.setEnabled(can_render_natively or can_convert_to_html)
                if not (can_render_natively or can_convert_to_html) and self.submission_viewer_stack.currentIndex() == 1:
                    # Force to text mode if rendered view isn't available and we're in rendered mode
                    self.switch_view_mode("text")
            
            # Determine current view mode preference
            should_use_rendered = False
            if hasattr(self, 'view_mode_rendered_btn') and self.view_mode_rendered_btn and self.view_mode_rendered_btn.isChecked():
                if can_render_natively or can_convert_to_html:
                    should_use_rendered = True
                    print("DEBUG: Using rendered view mode")
                else:
                    print("DEBUG: Rendered view requested but file cannot be rendered, using text mode")
            else:
                print("DEBUG: Using text view mode")
            
            # Load content based on determined view mode
            if should_use_rendered and WEB_ENGINE_AVAILABLE:
                # Set to rendered view
                self.submission_viewer_stack.setCurrentIndex(1)
                print(f"DEBUG: Set stack to rendered view (index 1)")
                
                if can_render_natively and file_ext == '.pdf':
                    # Load PDF using PDF.js
                    print(f"DEBUG: Loading PDF file natively: {current_file}")
                    self.load_pdf_in_viewer(current_file)
                elif can_convert_to_html and file_ext in ['.docx', '.odt']:
                    # Convert to HTML and display
                    print(f"DEBUG: Converting {file_ext} file to HTML")
                    html_content = self.convert_document_to_html(current_file, file_ext)
                    if hasattr(self, 'review_document_viewer') and self.review_document_viewer:
                        self.review_document_viewer.setHtml(html_content)
                        print(f"DEBUG: HTML content set in QWebEngineView")
                    else:
                        print(f"DEBUG: ERROR - review_document_viewer not available")
                else:
                    print(f"DEBUG: ERROR - Should not reach this case in rendered mode")
            else:
                # Use text view
                self.submission_viewer_stack.setCurrentIndex(0)
                print(f"DEBUG: Set stack to text view (index 0)")
                
                try:
                    content = self.read_file_content(current_file, file_ext)
                    if hasattr(self, 'review_submission_viewer') and content:
                        self.review_submission_viewer.setPlainText(content)
                        print(f"DEBUG: Loaded text content ({len(content)} characters)")
                        print(f"DEBUG: Content preview: {content[:200]}...")
                    elif hasattr(self, 'review_submission_viewer'):
                        self.review_submission_viewer.setPlainText(f"Could not read content from {file_name}")
                        print("DEBUG: Could not read file content")
                    else:
                        print("DEBUG: No review_submission_viewer available")
                except Exception as content_error:
                    print(f"DEBUG: Error reading file content: {content_error}")
                    if hasattr(self, 'review_submission_viewer'):
                        self.review_submission_viewer.setPlainText(f"Error reading {file_name}: {str(content_error)}")
                    
        except Exception as e:
            print(f"DEBUG: Error in load_current_submission_file: {e}")
            if hasattr(self, 'review_submission_viewer'):
                self.review_submission_viewer.setPlainText(f"Error loading file: {str(e)}")
            if hasattr(self, 'review_file_info'):
                self.review_file_info.setText(f"Error: {file_name}")
                # Hide navigation controls for single files
                self.review_prev_file_btn.hide()
                self.review_next_file_btn.hide()
                self.review_file_counter.hide()
        except Exception as e:
            print(f"DEBUG: Error in load_current_submission_file: {e}")
            if hasattr(self, 'review_submission_viewer'):
                self.review_submission_viewer.setPlainText(f"Error loading file: {str(e)}")
            if hasattr(self, 'review_file_info'):
                self.review_file_info.setText(f"Error: {file_name}")
                # Hide navigation controls for single files
                if hasattr(self, 'review_prev_file_btn'):
                    self.review_prev_file_btn.hide()
                if hasattr(self, 'review_next_file_btn'):
                    self.review_next_file_btn.hide()
                if hasattr(self, 'review_file_counter'):
                    self.review_file_counter.hide()
    
    def load_pdf_in_viewer(self, pdf_file_path):
        """Load PDF file in the web viewer using PDF.js"""
        try:
            if not hasattr(self, 'review_document_viewer') or not self.review_document_viewer:
                print("DEBUG: ERROR - review_document_viewer not available")
                return
                
            print(f"DEBUG: Loading PDF file: {pdf_file_path}")
            
            # Read PDF file and encode as base64
            import base64
            with open(pdf_file_path, 'rb') as f:
                pdf_data = f.read()
                pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
                
            file_name = os.path.basename(pdf_file_path)
            
            # Create HTML with PDF.js viewer
            html_with_pdf = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>PDF Viewer</title>
                <style>
                    body {{ 
                        margin: 0; 
                        padding: 0; 
                        font-family: Arial, sans-serif;
                        background-color: #f0f0f0;
                        display: flex;
                        flex-direction: column;
                        height: 100vh;
                    }}
                    .pdf-header {{
                        background-color: #333;
                        color: white;
                        padding: 10px;
                        font-size: 14px;
                        flex-shrink: 0;
                    }}
                    .pdf-container {{
                        flex-grow: 1;
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        background-color: #525659;
                    }}
                    canvas {{
                        border: 1px solid #ccc;
                        background-color: white;
                        max-width: 100%;
                        max-height: 100%;
                    }}
                    .controls {{
                        background-color: #444;
                        color: white;
                        padding: 8px;
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        gap: 10px;
                        flex-shrink: 0;
                    }}
                    button {{
                        background-color: #666;
                        color: white;
                        border: none;
                        padding: 5px 10px;
                        border-radius: 3px;
                        cursor: pointer;
                    }}
                    button:hover {{
                        background-color: #777;
                    }}
                    button:disabled {{
                        background-color: #555;
                        color: #999;
                        cursor: not-allowed;
                    }}
                    .page-info {{
                        color: #ccc;
                    }}
                    .fallback {{
                        padding: 20px;
                        text-align: center;
                        background-color: white;
                        margin: 20px;
                        border-radius: 8px;
                    }}
                    .download-link {{
                        display: inline-block;
                        background-color: #007ACC;
                        color: white;
                        padding: 10px 20px;
                        text-decoration: none;
                        border-radius: 4px;
                        margin-top: 10px;
                    }}
                </style>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
            </head>
            <body>
                <div class="pdf-header">
                    📄 {file_name}
                </div>
                <div class="controls">
                    <button id="prev-page">◀ Previous</button>
                    <span class="page-info">
                        Page <span id="page-num">1</span> of <span id="page-count">-</span>
                    </span>
                    <button id="next-page">Next ▶</button>
                    <button id="zoom-in">🔍+</button>
                    <button id="zoom-out">🔍-</button>
                    <span class="page-info">Zoom: <span id="zoom-level">100%</span></span>
                </div>
                <div class="pdf-container">
                    <canvas id="pdf-canvas"></canvas>
                </div>
                
                <script>
                    // PDF.js setup
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
                    
                    const pdfData = 'data:application/pdf;base64,{pdf_base64}';
                    let pdfDoc = null;
                    let pageNum = 1;
                    let pageRendering = false;
                    let pageNumPending = null;
                    let scale = 1.0;
                    const canvas = document.getElementById('pdf-canvas');
                    const ctx = canvas.getContext('2d');
                    
                    // Load the PDF
                    pdfjsLib.getDocument(pdfData).promise.then(function(pdfDoc_) {{
                        pdfDoc = pdfDoc_;
                        document.getElementById('page-count').textContent = pdfDoc.numPages;
                        renderPage(pageNum);
                    }}).catch(function(error) {{
                        console.error('Error loading PDF:', error);
                        document.querySelector('.pdf-container').innerHTML = `
                            <div class="fallback">
                                <h3>PDF Loading Error</h3>
                                <p>Unable to display PDF file using PDF.js</p>
                                <p>Error: ${{error.message}}</p>
                                <a href="data:application/pdf;base64,{pdf_base64}" 
                                   download="{file_name}" 
                                   class="download-link">Download PDF</a>
                            </div>
                        `;
                    }});
                    
                    function renderPage(num) {{
                        pageRendering = true;
                        pdfDoc.getPage(num).then(function(page) {{
                            const viewport = page.getViewport({{scale: scale}});
                            canvas.height = viewport.height;
                            canvas.width = viewport.width;
                            
                            const renderContext = {{
                                canvasContext: ctx,
                                viewport: viewport
                            }};
                            
                            const renderTask = page.render(renderContext);
                            renderTask.promise.then(function() {{
                                pageRendering = false;
                                if (pageNumPending !== null) {{
                                    renderPage(pageNumPending);
                                    pageNumPending = null;
                                }}
                            }});
                        }});
                        
                        document.getElementById('page-num').textContent = num;
                        updateButtons();
                    }}
                    
                    function queueRenderPage(num) {{
                        if (pageRendering) {{
                            pageNumPending = num;
                        }} else {{
                            renderPage(num);
                        }}
                    }}
                    
                    function updateButtons() {{
                        document.getElementById('prev-page').disabled = pageNum <= 1;
                        document.getElementById('next-page').disabled = pageNum >= pdfDoc.numPages;
                        document.getElementById('zoom-level').textContent = Math.round(scale * 100) + '%';
                    }}
                    
                    // Event listeners
                    document.getElementById('prev-page').addEventListener('click', function() {{
                        if (pageNum <= 1) return;
                        pageNum--;
                        queueRenderPage(pageNum);
                    }});
                    
                    document.getElementById('next-page').addEventListener('click', function() {{
                        if (pageNum >= pdfDoc.numPages) return;
                        pageNum++;
                        queueRenderPage(pageNum);
                    }});
                    
                    document.getElementById('zoom-in').addEventListener('click', function() {{
                        scale *= 1.2;
                        queueRenderPage(pageNum);
                    }});
                    
                    document.getElementById('zoom-out').addEventListener('click', function() {{
                        scale /= 1.2;
                        queueRenderPage(pageNum);
                    }});
                </script>
            </body>
            </html>
            """
            
            print("DEBUG: Setting HTML with PDF.js viewer")
            self.review_document_viewer.setHtml(html_with_pdf)
            print("DEBUG: PDF.js HTML set successfully")
            
        except Exception as e:
            print(f"DEBUG: Error loading PDF: {e}")
            # Fallback error message
            error_html = f"""
            <html><body style="font-family: Arial; padding: 20px; text-align: center;">
                <h3 style="color: #d32f2f;">PDF Loading Error</h3>
                <p>Unable to display PDF file: <strong>{os.path.basename(pdf_file_path)}</strong></p>
                <p style="color: #666;">File path: {pdf_file_path}</p>
                <p style="color: #666;">Error: {str(e)}</p>
                <p>Please ensure the file exists and is a valid PDF document.</p>
            </body></html>
            """
            if hasattr(self, 'review_document_viewer') and self.review_document_viewer:
                self.review_document_viewer.setHtml(error_html)

    def convert_document_to_html(self, file_path, file_ext):
        """Convert Word/ODT documents to HTML for web rendering"""
        try:
            if file_ext == '.docx':
                # Convert Word document to HTML
                try:
                    from docx import Document
                    doc = Document(file_path)
                    
                    html_content = """
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <style>
                            body { font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }
                            .paragraph { margin-bottom: 12px; }
                            .heading { font-weight: bold; font-size: 1.1em; margin-top: 20px; margin-bottom: 10px; }
                        </style>
                    </head>
                    <body>
                    """
                    
                    for paragraph in doc.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            style_class = "heading" if paragraph.style.name.startswith('Heading') else "paragraph"
                            html_content += f'<div class="{style_class}">{text}</div>\n'
                    
                    html_content += "</body></html>"
                    return html_content
                    
                except ImportError:
                    return "Word document viewing requires python-docx. Please install it to view .docx content."
                except Exception as e:
                    return f"Error converting Word document: {str(e)}"
                    
            elif file_ext == '.odt':
                # Convert ODT document to HTML
                try:
                    from odf.opendocument import load
                    from odf.text import P
                    
                    doc = load(file_path)
                    
                    html_content = """
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <style>
                            body { font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }
                            .paragraph { margin-bottom: 12px; }
                        </style>
                    </head>
                    <body>
                    """
                    
                    paragraphs = doc.getElementsByType(P)
                    for paragraph in paragraphs:
                        text = ""
                        for node in paragraph.childNodes:
                            if hasattr(node, 'data'):
                                text += node.data
                            elif hasattr(node, 'firstChild') and node.firstChild:
                                text += node.firstChild.data if hasattr(node.firstChild, 'data') else str(node.firstChild)
                        
                        text = text.strip()
                        if text:
                            html_content += f'<div class="paragraph">{text}</div>\n'
                    
                    html_content += "</body></html>"
                    return html_content
                    
                except ImportError:
                    return "OpenDocument viewing requires odfpy. Please install it to view .odt content."
                except Exception as e:
                    return f"Error converting OpenDocument: {str(e)}"
            
            else:
                return f"Cannot convert {file_ext} files to HTML"
                
        except Exception as e:
            return f"Error in document conversion: {str(e)}"
    
    def load_text_submission_fallback(self, student_folder):
        """Load text submission data if no files are found"""
        try:
            # Look for submission data files (JSON, etc.) that might contain text submissions
            data_files = [f for f in os.listdir(student_folder) 
                         if f.lower().endswith(('.json', '.txt')) and 'submission' in f.lower()]
            
            if data_files:
                # Try to read submission data
                data_file = os.path.join(student_folder, data_files[0])
                with open(data_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                self.review_file_info.setText(f"Text submission data: {data_files[0]}")
                self.review_submission_viewer.setPlainText(f"Text Submission Data:\n\n{content}")
                print(f"DEBUG: Loaded text submission data from {data_files[0]}")
            else:
                self.review_file_info.setText("No submission files found")
                self.review_submission_viewer.setPlainText("No submission files or text submissions found for this student.")
                print(f"DEBUG: No submission files or data found")
                
        except Exception as e:
            self.review_file_info.setText("No submissions found")
            self.review_submission_viewer.setPlainText(f"No submission content available. Error checking for text submissions: {str(e)}")
    
    def read_file_content(self, file_path, file_ext):
        """Read and return file content based on file type"""
        try:
            if file_ext in ['.txt', '.py', '.java', '.cpp', '.c', '.js', '.html', '.css', '.md']:
                # Plain text files
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError:
                    with open(file_path, 'r', encoding='latin1') as f:
                        return f.read()
                        
            elif file_ext == '.pdf':
                # PDF files - extract text with improved formatting
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text_content = ""
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            if page_text:
                                # Improve text formatting by adding proper line breaks
                                lines = page_text.split('\n')
                                formatted_lines = []
                                for line in lines:
                                    line = line.strip()
                                    if line:
                                        # Add space if line doesn't end with punctuation and next char is letter
                                        if (formatted_lines and 
                                            not formatted_lines[-1].endswith(('.', '!', '?', ':', ';', '-')) and 
                                            line[0].islower()):
                                            formatted_lines[-1] += ' ' + line
                                        else:
                                            formatted_lines.append(line)
                                
                                text_content += '\n'.join(formatted_lines)
                                if page_num < len(pdf_reader.pages) - 1:
                                    text_content += "\n\n--- Page Break ---\n\n"
                    
                    if text_content.strip():
                        return text_content
                    else:
                        return "PDF content could not be extracted or is empty."
                except ImportError:
                    return "PDF viewing requires PyPDF2. Please install it to view PDF content."
                except Exception as e:
                    return f"Error reading PDF: {str(e)}"
                    
            elif file_ext == '.docx':
                # Word documents
                try:
                    import docx
                    doc = docx.Document(file_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    return content if content.strip() else "Word document appears to be empty."
                except ImportError:
                    return "Word document viewing requires python-docx. Please install it to view .docx content."
                except Exception as e:
                    return f"Error reading Word document: {str(e)}"
                    
            elif file_ext == '.odt':
                # OpenDocument Text files
                try:
                    from odf.opendocument import load
                    from odf.text import P
                    
                    doc = load(file_path)
                    content = ""
                    paragraphs = doc.getElementsByType(P)
                    for paragraph in paragraphs:
                        para_text = ""
                        for node in paragraph.childNodes:
                            if hasattr(node, 'data'):
                                para_text += node.data
                            elif hasattr(node, 'childNodes'):
                                for child in node.childNodes:
                                    if hasattr(child, 'data'):
                                        para_text += child.data
                        if para_text.strip():
                            content += para_text + "\n"
                    
                    return content if content.strip() else "OpenDocument file appears to be empty."
                except ImportError:
                    return "OpenDocument viewing requires odfpy. Please install it to view .odt content.\nInstall with: pip install odfpy"
                except Exception as e:
                    return f"Error reading OpenDocument file: {str(e)}"
                    
            else:
                # Unsupported file type
                return (f"File type '{file_ext}' is not directly viewable.\n\n"
                       f"Supported formats: .txt, .py, .java, .cpp, .c, .js, .html, .css, .md, .pdf, .docx, .odt\n\n"
                       f"File path: {file_path}")
                
        except Exception as e:
            return f"Error reading file: {str(e)}"
    
    def review_previous_submission_file(self):
        """Navigate to previous submission file for current student"""
        if hasattr(self, 'current_submission_files') and self.current_submission_files:
            if self.current_submission_index > 0:
                self.current_submission_index -= 1
                self.load_current_submission_file()
    
    def review_next_submission_file(self):
        """Navigate to next submission file for current student"""
        if hasattr(self, 'current_submission_files') and self.current_submission_files:
            if self.current_submission_index < len(self.current_submission_files) - 1:
                self.current_submission_index += 1
                self.load_current_submission_file()
    
    def switch_view_mode(self, mode):
        """Switch between text and rendered view modes"""
        print(f"DEBUG: Switching to {mode} view mode")
        if mode == "text":
            self.submission_viewer_stack.setCurrentIndex(0)
            self.view_mode_text_btn.setChecked(True)
            if self.view_mode_rendered_btn:
                self.view_mode_rendered_btn.setChecked(False)
            print(f"DEBUG: Switched to text view (stack index 0)")
            # Reload current file in text mode
            self.load_current_submission_file()
        elif mode == "rendered" and WEB_ENGINE_AVAILABLE:
            self.submission_viewer_stack.setCurrentIndex(1)
            self.view_mode_text_btn.setChecked(False)
            if self.view_mode_rendered_btn:
                self.view_mode_rendered_btn.setChecked(True)
            print(f"DEBUG: Switched to rendered view (stack index 1)")
            # Reload current file in rendered mode
            self.load_current_submission_file()
        else:
            print(f"DEBUG: Cannot switch to {mode} - WEB_ENGINE_AVAILABLE={WEB_ENGINE_AVAILABLE}")
    
    def review_previous_submission(self):
        """Navigate to previous submission with wraparound"""
        current_index = self.review_submission_combo.currentIndex()
        total = self.review_submission_combo.count()
        if total > 1:
            if current_index > 0:
                self.review_submission_combo.setCurrentIndex(current_index - 1)
            else:
                # Wrap around to last submission
                self.review_submission_combo.setCurrentIndex(total - 1)
    
    def review_next_submission(self):
        """Navigate to next submission with wraparound"""
        current_index = self.review_submission_combo.currentIndex()
        total = self.review_submission_combo.count()
        if total > 1:
            if current_index < total - 1:
                self.review_submission_combo.setCurrentIndex(current_index + 1)
            else:
                # Wrap around to first submission
                self.review_submission_combo.setCurrentIndex(0)
    
    def review_submission_changed(self):
        """Handle submission selection change"""
        self.load_current_submission()
    
    def review_score_changed(self):
        """Handle score change"""
        self.mark_current_as_changed()
        self.update_save_button_state()
    
    def review_comments_changed(self):
        """Handle comments change"""
        self.mark_current_as_changed()
        self.update_save_button_state()
    
    def update_score_focus_style(self, has_focus, event):
        """Update the score field styling when it gains/loses focus"""
        from PyQt6.QtWidgets import QLineEdit
        
        # Call the original focus event method
        if has_focus:
            QLineEdit.focusInEvent(self.review_score_entry, event)
        else:
            QLineEdit.focusOutEvent(self.review_score_entry, event)
        
        # Update styling based on focus state
        self.update_score_field_styling(has_focus)
    
    def update_score_field_styling(self, has_focus=False):
        """Update the score field styling - simple standalone box design"""
        if not hasattr(self, 'review_score_entry'):
            return
            
        # Simple standalone styling - no complex two-part design needed
        focus_border = "#80bdff" if has_focus else "#ced4da"
        focus_shadow = "0px 0px 0px 0.2rem rgba(0, 123, 255, 0.25)" if has_focus else "none"
        
        self.review_score_entry.setStyleSheet(f"""
            QLineEdit {{
                border: 2px solid {focus_border};
                border-radius: 6px;
                padding: 8px 10px;
                background-color: white;
                font-size: 11pt;
                font-weight: normal;
                color: #495057;
                min-height: 16px;
                max-height: 16px;
            }}
            QLineEdit:focus {{
                border-color: #80bdff;
                box-shadow: {focus_shadow};
                outline: 0;
                background-color: #fff;
            }}
        """)
    
    def review_open_submission_directory(self):
        """Open the current submission's directory in file explorer"""
        if not hasattr(self, 'review_data') or not self.review_data:
            QMessageBox.warning(self, "Warning", "No submissions loaded.")
            return
        
        current_index = self.review_submission_combo.currentIndex()
        if current_index < 0:
            QMessageBox.warning(self, "Warning", "No submission selected.")
            return
        
        # Get student data
        student_id = self.review_submission_combo.itemData(current_index)
        student_data = None
        student_name_key = None
        for name, data in self.review_data.items():
            if str(data.get('student_id', '')) == str(student_id):
                student_data = data
                student_name_key = name
                break
        
        if not student_data:
            QMessageBox.warning(self, "Warning", "Student data not found.")
            return
        
        # Use the same submission data lookup as find_submission_files method
        anon_student_name = student_data.get('anon_name', student_name_key)
        real_student_name = student_data.get('real_name', student_name_key)
        
        print(f"DEBUG: Looking for directory for student: {real_student_name} (anon: {anon_student_name}, ID: {student_id})")
        
        # Use find_submission_files to get actual file paths stored during download
        submission_files = self.find_submission_files(anon_student_name, student_id)
        
        student_folder = None
        if submission_files:
            # Get the directory from the first file path
            from pathlib import Path
            first_file_path = Path(submission_files[0])
            student_folder = first_file_path.parent
            print(f"DEBUG: Found student folder from submission files: {student_folder}")
        else:
            print(f"DEBUG: No submission files found for {anon_student_name}")
            
            # Fallback: try to find by direct folder search as before
            from pathlib import Path
            submission_path = Path(self.submission_folder_path)
            
            if submission_path.exists():
                # Try exact name matches first
                for name_to_try in [anon_student_name, real_student_name]:
                    test_folder = submission_path / name_to_try
                    if test_folder.exists() and test_folder.is_dir():
                        student_folder = test_folder
                        print(f"DEBUG: Found exact folder match: {student_folder}")
                        break
                
                # If no exact match, try pattern matching
                if not student_folder:
                    print(f"DEBUG: No exact folder match, searching for partial matches...")
                    search_terms = [
                        real_student_name.lower(),
                        anon_student_name.lower(),
                        str(student_id).lower()
                    ]
                    
                    for folder in submission_path.iterdir():
                        if folder.is_dir():
                            folder_name = folder.name.lower()
                            if any(term in folder_name for term in search_terms):
                                student_folder = folder
                                print(f"DEBUG: Found partial match: {student_folder}")
                                break
        
        # Open the directory if found
        if student_folder and student_folder.exists():
            import subprocess
            import os
            try:
                folder_path = str(student_folder)
                print(f"DEBUG: Attempting to open folder: {folder_path}")
                
                # Method 1: Use os.startfile (Windows-specific) - most reliable for folders
                try:
                    os.startfile(folder_path)
                    print(f"DEBUG: Successfully opened directory with os.startfile: {folder_path}")
                    return
                except Exception as startfile_error:
                    print(f"DEBUG: os.startfile method failed: {startfile_error}")
                
                # Method 2: Use explorer directly on the folder
                try:
                    subprocess.run(['explorer', folder_path], check=True)
                    print(f"DEBUG: Successfully opened directory with explorer: {folder_path}")
                    return
                except subprocess.CalledProcessError as explorer_error:
                    print(f"DEBUG: Direct explorer method failed: {explorer_error}")
                
                # Method 3: Use explorer to open parent folder and show the target folder
                try:
                    parent_folder = str(student_folder.parent)
                    subprocess.run(['explorer', parent_folder], check=True)
                    print(f"DEBUG: Opened parent directory as fallback: {parent_folder}")
                    QMessageBox.information(self, "Directory Opened", 
                                          f"Opened parent folder. Look for: {student_folder.name}")
                    return
                except subprocess.CalledProcessError as parent_error:
                    print(f"DEBUG: Parent folder method failed: {parent_error}")
                
                # If all methods fail, show error
                QMessageBox.warning(self, "Error", 
                                  f"Could not open folder: {folder_path}\n"
                                  f"You can manually navigate to:\n{folder_path}")
                
            except Exception as e:
                print(f"DEBUG: Unexpected error opening directory: {e}")
                QMessageBox.warning(self, "Error", f"Failed to open directory: {str(e)}")
        else:
            QMessageBox.warning(self, "Warning", f"Submission directory not found for student: {real_student_name}")
            print(f"DEBUG: Could not find directory for student: {real_student_name} (anon: {anon_student_name}, ID: {student_id})")
            
            # Show more debugging info for troubleshooting
            from pathlib import Path
            submission_path = Path(self.submission_folder_path)
            if submission_path.exists():
                available_folders = [f.name for f in submission_path.iterdir() if f.is_dir()]
                print(f"DEBUG: Searched in: {submission_path}")
                print(f"DEBUG: Available folders: {available_folders}")
            else:
                print(f"DEBUG: Submission path does not exist: {submission_path}")
    
    def mark_current_as_changed(self):
        """Mark the current submission as having unsaved changes"""
        current_index = self.review_submission_combo.currentIndex()
        if current_index >= 0:
            student_id = self.review_submission_combo.itemData(current_index)
            if student_id:
                self.review_unsaved_changes.add(student_id)
                self.update_changes_label()
    
    def update_save_button_state(self):
        """Update the save button enabled state"""
        current_index = self.review_submission_combo.currentIndex()
        if current_index >= 0:
            student_id = self.review_submission_combo.itemData(current_index)
            has_changes = student_id in self.review_unsaved_changes
            self.review_save_btn.setEnabled(has_changes)
    
    def update_changes_label(self):
        """Update the changes indicator label"""
        num_changes = len(self.review_unsaved_changes)
        if num_changes == 0:
            self.review_changes_label.setText("No unsaved changes")
            self.review_changes_label.setStyleSheet("color: #666; font-style: italic;")
        else:
            self.review_changes_label.setText(f"{num_changes} submission(s) with unsaved changes")
            self.review_changes_label.setStyleSheet("color: #dc3545; font-weight: bold;")
    
    def review_clear_comments(self):
        """Clear all comments for current submission"""
        reply = QMessageBox.question(self, "Clear Comments", 
                                   "Are you sure you want to clear all comments for this submission?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            self.review_comments_editor.clear()
    
    def review_restore_ai_comments(self):
        """Restore original AI comments for current submission"""
        current_index = self.review_submission_combo.currentIndex()
        if current_index < 0:
            return
        
        # Get student data using the same logic as load_current_submission
        student_id = self.review_submission_combo.itemData(current_index)
        student_data = None
        student_name_key = None
        for name, data in self.review_data.items():
            if str(data.get('student_id', '')) == str(student_id):
                student_data = data
                student_name_key = name
                break
        
        if not student_data or student_name_key not in self.review_original_data:
            QMessageBox.warning(self, "Warning", "Original data not found for this student.")
            return
        
        reply = QMessageBox.question(self, "Restore AI Comments", 
                                   "Are you sure you want to restore the original AI comments and score?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            # Get original data using student name as key
            original_data = self.review_original_data[student_name_key]
            
            # Restore score
            if hasattr(self, 'review_score_entry'):
                self.review_score_entry.setText(str(int(original_data['score'])))
            
            # Restore comments
            if hasattr(self, 'review_comments_editor'):
                self.review_comments_editor.blockSignals(True)
                self.review_comments_editor.setPlainText(original_data['comments'])
                self.review_comments_editor.blockSignals(False)
            
            # Mark as changed since we're restoring to original (but actually back to original state)
            # Update the review_data to reflect the restored values
            self.review_data[student_name_key]['score'] = original_data['score']
            self.review_data[student_name_key]['comments'] = original_data['comments']
            self.review_data[student_name_key]['changed'] = False  # Reset to unchanged since we restored original
            
            # Update save button state
            if hasattr(self, 'update_save_button_state'):
                self.update_save_button_state()
            
            print(f"DEBUG: Restored original AI comments for {student_name_key}")
    
    def review_save_current(self):
        """Save changes to current submission"""
        current_index = self.review_submission_combo.currentIndex()
        if current_index < 0:
            return
        
        student_id = self.review_submission_combo.itemData(current_index)
        if not student_id or student_id not in self.review_data:
            return
        
        # Update the in-memory data
        try:
            score_text = self.review_score_entry.text().strip()
            score_value = int(score_text) if score_text else 0
            self.review_data[student_id]['score'] = score_value
        except ValueError:
            # If score text is not a valid integer, default to 0
            self.review_data[student_id]['score'] = 0
        self.review_data[student_id]['comments'] = self.review_comments_editor.toPlainText()
        
        # Remove from unsaved changes
        self.review_unsaved_changes.discard(student_id)
        
        # Update UI
        self.update_save_button_state()
        self.update_changes_label()
        
        # Show confirmation
        student_name = self.review_data[student_id]['student_name']
        QMessageBox.information(self, "Saved", f"Changes saved for {student_name}")
    
    def review_save_all_changes(self):
        """Save all changes back to the spreadsheet"""
        try:
            import pandas as pd
            
            # Read current spreadsheet
            df = pd.read_excel(self.review_spreadsheet_path)
            
            # Update the dataframe with our changes
            for index, row in df.iterrows():
                student_id = str(row.get('Student ID', f'student_{index}'))
                
                if student_id in self.review_data:
                    df.at[index, 'Score'] = self.review_data[student_id]['score']
                    df.at[index, 'Comments'] = self.review_data[student_id]['comments']
            
            # Save back to spreadsheet
            df.to_excel(self.review_spreadsheet_path, index=False)
            
            # Clear all unsaved changes
            self.review_unsaved_changes.clear()
            self.update_save_button_state()
            self.update_changes_label()
            
            QMessageBox.information(self, "All Changes Saved", 
                                  f"All changes have been saved to:\n{self.review_spreadsheet_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Save Error", 
                               f"Failed to save changes to spreadsheet:\n{str(e)}")


def main():
    """Main function to run the Canvas GUI"""
    app = QApplication(sys.argv)
    
    # Set application properties
    app.setApplicationName("DuckGrade Canvas Integration")
    app.setApplicationVersion("2.0.0")
    app.setOrganizationName("DuckWorks Educational Automation")
    
    # Create and show the main window
    window = DuckGradeCanvasGUI()
    window.showMaximized()  # Start in fullscreen/maximized mode
    
    print("🦆 DuckGrade Canvas Integration (PyQt6) started successfully!")
    print("Professional interface now matches Tkinter structure exactly")
    
    return app.exec()



class InstructorConfigBuilder(QDialog):
    """A user-friendly GUI for building instructor configuration JSON files"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Instructor Configuration Builder")
        self.setWindowIcon(QIcon("assets/icons8-flying-duck-48.png") if Path("assets/icons8-flying-duck-48.png").exists() else QIcon())
        self.setModal(True)
        self.resize(900, 600)  # Reduced height to 600 for better screen fit
        self.setFixedSize(900, 600)  # Force the dialog to stay at this exact size
        
        self.parent_widget = parent  # Store parent for later centering
        self.setup_ui()
        
        # Center the dialog after UI is set up - use a timer to ensure proper initialization
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(0, self.center_dialog)
    
    def center_dialog(self):
        """Center the dialog on screen or parent after UI is fully set up"""
        # Ensure the dialog size is properly set
        self.updateGeometry()
        
        if self.parent_widget and self.parent_widget.isVisible():
            # Center relative to parent window
            parent_geo = self.parent_widget.frameGeometry()
            parent_center = parent_geo.center()
            
            # Get this dialog's frame geometry
            dialog_geo = self.frameGeometry()
            dialog_geo.moveCenter(parent_center)
            
            # Make sure the dialog doesn't go off screen
            screen = QApplication.primaryScreen().availableGeometry()
            if dialog_geo.left() < screen.left():
                dialog_geo.moveLeft(screen.left())
            if dialog_geo.right() > screen.right():
                dialog_geo.moveRight(screen.right())
            if dialog_geo.top() < screen.top():
                dialog_geo.moveTop(screen.top())
            if dialog_geo.bottom() > screen.bottom():
                dialog_geo.moveBottom(screen.bottom())
            
            self.move(dialog_geo.topLeft())
        else:
            # Center on screen if no parent or parent not visible
            screen = QApplication.primaryScreen().availableGeometry()
            dialog_geo = self.frameGeometry()
            dialog_geo.moveCenter(screen.center())
            self.move(dialog_geo.topLeft())
    
    def setup_ui(self):
        """Setup the user interface"""
        layout = QVBoxLayout(self)
        
        # Header
        header_label = QLabel("🎓 Instructor Configuration Builder")
        header_label.setStyleSheet("font-size: 16pt; font-weight: bold; color: #2c3e50; padding: 10px;")
        header_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(header_label)
        
        # Description
        desc_label = QLabel("Create a personalized instructor configuration to customize AI grading behavior, "
                           "feedback style, and subject expertise. This will help the AI provide more relevant "
                           "and consistent grading based on your teaching preferences.")
        desc_label.setWordWrap(True)
        desc_label.setStyleSheet("color: #666; padding: 0 10px 10px 10px; font-size: 10pt;")
        layout.addWidget(desc_label)
        
        # Create scroll area for the form
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Form widget
        form_widget = QWidget()
        form_layout = QVBoxLayout(form_widget)
        
        # Basic Information Group
        basic_group = QGroupBox("📋 Basic Information")
        basic_layout = QFormLayout(basic_group)
        
        self.instructor_name = QLineEdit()
        self.instructor_name.setPlaceholderText("e.g., Dr. Smith")
        basic_layout.addRow("Instructor Name:", self.instructor_name)
        
        self.institution = QLineEdit()
        self.institution.setPlaceholderText("e.g., University of Education")
        basic_layout.addRow("Institution:", self.institution)
        
        form_layout.addWidget(basic_group)
        
        # Course Context Group
        course_group = QGroupBox("🎯 Course Context")
        course_layout = QFormLayout(course_group)
        
        self.course_level = ScrollFriendlyComboBox()
        self.course_level.addItems(["High School", "Undergraduate", "Graduate", "Professional"])
        self.course_level.setStyleSheet("""
            QComboBox {
                border: 2px solid #ced4da;
                border-radius: 6px;
                padding: 8px 10px;
                background-color: white;
                font-size: 10pt;
                min-height: 20px;
            }
            QComboBox:focus {
                border-color: #80bdff;
                background-color: #fff;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        course_layout.addRow("Course Level:", self.course_level)
        
        self.course_type = ScrollFriendlyComboBox()
        self.course_type.addItems(["General Education", "Major Requirement", "Elective", "Capstone", "Seminar", "Laboratory"])
        self.course_type.setStyleSheet("""
            QComboBox {
                border: 2px solid #ced4da;
                border-radius: 6px;
                padding: 8px 10px;
                background-color: white;
                font-size: 10pt;
                min-height: 20px;
            }
            QComboBox:focus {
                border-color: #80bdff;
                background-color: #fff;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        course_layout.addRow("Course Type:", self.course_type)
        
        self.student_background = ScrollFriendlyComboBox()
        self.student_background.addItems(["Mixed levels", "Beginner", "Intermediate", "Advanced", "Mixed with support needs"])
        self.student_background.setStyleSheet("""
            QComboBox {
                border: 2px solid #ced4da;
                border-radius: 6px;
                padding: 8px 10px;
                background-color: white;
                font-size: 10pt;
                min-height: 20px;
            }
            QComboBox:focus {
                border-color: #80bdff;
                background-color: #fff;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ced4da;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: url(assets/down-arrow_gray.png);
                width: 12px;
                height: 8px;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #ced4da;
                background-color: white;
                selection-background-color: #007bff;
                selection-color: white;
            }
        """)
        course_layout.addRow("Student Background:", self.student_background)
        
        # Learning objectives
        objectives_label = QLabel("Learning Objectives:")
        self.learning_objectives = QTextEdit()
        self.learning_objectives.setMaximumHeight(80)
        self.learning_objectives.setPlaceholderText("Enter key learning objectives, one per line...")
        course_layout.addRow(objectives_label, self.learning_objectives)
        
        form_layout.addWidget(course_group)
        
        # Grading Philosophy Group
        philosophy_group = QGroupBox("💭 Grading Philosophy")
        philosophy_layout = QVBoxLayout(philosophy_group)
        
        philosophy_label = QLabel("How would you describe your grading approach?")
        philosophy_layout.addWidget(philosophy_label)
        
        self.philosophy_encouraging = QRadioButton("Encouraging - Focus on positive reinforcement and growth")
        self.philosophy_balanced = QRadioButton("Balanced - Mix of praise and constructive criticism")
        self.philosophy_rigorous = QRadioButton("Rigorous - High standards with detailed feedback")
        self.philosophy_custom = QRadioButton("Custom:")
        
        self.philosophy_balanced.setChecked(True)  # Default
        
        philosophy_layout.addWidget(self.philosophy_encouraging)
        philosophy_layout.addWidget(self.philosophy_balanced)
        philosophy_layout.addWidget(self.philosophy_rigorous)
        philosophy_layout.addWidget(self.philosophy_custom)
        
        self.custom_philosophy = QLineEdit()
        self.custom_philosophy.setPlaceholderText("Describe your custom grading philosophy...")
        self.custom_philosophy.setEnabled(False)
        philosophy_layout.addWidget(self.custom_philosophy)
        
        # Connect custom radio button
        self.philosophy_custom.toggled.connect(lambda checked: self.custom_philosophy.setEnabled(checked))
        
        form_layout.addWidget(philosophy_group)
        
        # Subject Expertise Group
        expertise_group = QGroupBox("🎓 Subject Expertise")
        expertise_layout = QVBoxLayout(expertise_group)
        
        expertise_label = QLabel("Select your areas of expertise (check all that apply):")
        expertise_layout.addWidget(expertise_label)
        
        # Create expertise checkboxes in a grid
        expertise_grid = QGridLayout()
        
        self.expertise_checkboxes = {}
        expertise_areas = [
            "English Literature", "Writing and Composition", "Mathematics", "Science",
            "History", "Psychology", "Business", "Education", "Computer Science",
            "Art and Design", "Music", "Philosophy", "Political Science", "Economics",
            "Sociology", "Anthropology", "Environmental Studies", "Health Sciences"
        ]
        
        for i, area in enumerate(expertise_areas):
            checkbox = QCheckBox(area)
            self.expertise_checkboxes[area] = checkbox
            expertise_grid.addWidget(checkbox, i // 3, i % 3)
        
        expertise_layout.addLayout(expertise_grid)
        
        # Custom expertise
        custom_expertise_layout = QHBoxLayout()
        custom_expertise_layout.addWidget(QLabel("Other:"))
        self.custom_expertise = QLineEdit()
        self.custom_expertise.setPlaceholderText("Enter additional expertise areas separated by commas...")
        custom_expertise_layout.addWidget(self.custom_expertise)
        expertise_layout.addLayout(custom_expertise_layout)
        
        form_layout.addWidget(expertise_group)
        
        # Comment Preferences Group
        comments_group = QGroupBox("💬 Comment Preferences")
        comments_layout = QVBoxLayout(comments_group)
        
        self.include_strengths = QCheckBox("Always highlight what students did well")
        self.include_strengths.setChecked(True)
        comments_layout.addWidget(self.include_strengths)
        
        self.include_suggestions = QCheckBox("Provide specific suggestions for improvement")
        self.include_suggestions.setChecked(True)
        comments_layout.addWidget(self.include_suggestions)
        
        self.personal_touch = QCheckBox("Make comments personal and encouraging")
        self.personal_touch.setChecked(True)
        comments_layout.addWidget(self.personal_touch)
        
        self.detailed_feedback = QCheckBox("Provide detailed explanations for grades")
        self.detailed_feedback.setChecked(True)
        comments_layout.addWidget(self.detailed_feedback)
        
        form_layout.addWidget(comments_group)
        
        # Custom Instructions Group
        custom_group = QGroupBox("📝 Custom Instructions")
        custom_layout = QVBoxLayout(custom_group)
        
        custom_label = QLabel("Additional instructions for the AI grader (optional):")
        custom_layout.addWidget(custom_label)
        
        self.custom_instructions = QTextEdit()
        self.custom_instructions.setMaximumHeight(100)
        self.custom_instructions.setPlaceholderText("e.g., Pay special attention to citation format, Focus on critical thinking skills, etc.")
        custom_layout.addWidget(self.custom_instructions)
        
        form_layout.addWidget(custom_group)
        
        # Set the form widget in the scroll area
        scroll_area.setWidget(form_widget)
        layout.addWidget(scroll_area)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        preview_btn = QPushButton("👀 Preview JSON")
        preview_btn.clicked.connect(self.preview_config)
        button_layout.addWidget(preview_btn)
        
        button_layout.addStretch()
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        save_btn = QPushButton("💾 Save Configuration")
        save_btn.setStyleSheet("QPushButton { background-color: #28a745; }")
        save_btn.clicked.connect(self.save_config)
        button_layout.addWidget(save_btn)
        
        layout.addLayout(button_layout)
    
    def get_selected_philosophy(self):
        """Get the selected grading philosophy"""
        if self.philosophy_encouraging.isChecked():
            return "encouraging"
        elif self.philosophy_balanced.isChecked():
            return "balanced"
        elif self.philosophy_rigorous.isChecked():
            return "rigorous"
        else:
            return self.custom_philosophy.text().strip() or "balanced"
    
    def get_selected_expertise(self):
        """Get selected expertise areas"""
        expertise = []
        for area, checkbox in self.expertise_checkboxes.items():
            if checkbox.isChecked():
                expertise.append(area)
        
        # Add custom expertise
        custom = self.custom_expertise.text().strip()
        if custom:
            expertise.extend([area.strip() for area in custom.split(',') if area.strip()])
        
        return expertise
    
    def get_learning_objectives(self):
        """Get learning objectives as a list"""
        text = self.learning_objectives.toPlainText().strip()
        if not text:
            return []
        return [obj.strip() for obj in text.split('\n') if obj.strip()]
    
    def get_custom_instructions(self):
        """Get custom instructions as a list"""
        text = self.custom_instructions.toPlainText().strip()
        if not text:
            return []
        return [instr.strip() for instr in text.split('\n') if instr.strip()]
    
    def build_config(self):
        """Build the configuration dictionary"""
        config = {
            "instructor_info": {
                "name": self.instructor_name.text().strip(),
                "institution": self.institution.text().strip()
            },
            "course_context": {
                "course_level": self.course_level.currentText(),
                "course_type": self.course_type.currentText(),
                "student_background": self.student_background.currentText(),
                "learning_objectives": self.get_learning_objectives()
            },
            "grading_philosophy": self.get_selected_philosophy(),
            "subject_expertise": self.get_selected_expertise(),
            "comment_preferences": {
                "include_strengths": self.include_strengths.isChecked(),
                "include_suggestions": self.include_suggestions.isChecked(),
                "personal_touch": self.personal_touch.isChecked(),
                "detailed_feedback": self.detailed_feedback.isChecked()
            },
            "specific_instructions": self.get_custom_instructions()
        }
        
        # Add custom system message based on philosophy
        philosophy = self.get_selected_philosophy()
        system_messages = {
            "encouraging": "You are a supportive and encouraging academic grader. Focus on student growth, highlight positives, and provide constructive feedback that motivates learning.",
            "balanced": "You are an experienced academic grader who provides fair, balanced feedback. Acknowledge both strengths and areas for improvement with specific, actionable suggestions.",
            "rigorous": "You are a rigorous academic grader with high standards. Provide detailed, thorough feedback that challenges students to excel while maintaining fairness."
        }
        
        if philosophy in system_messages:
            config["custom_system_message"] = system_messages[philosophy]
        elif philosophy != "balanced":  # Custom philosophy
            config["custom_system_message"] = f"You are an academic grader with this approach: {philosophy}"
        
        return config
    
    def preview_config(self):
        """Show a preview of the generated JSON"""
        import json
        config = self.build_config()
        
        preview_dialog = QDialog(self)
        preview_dialog.setWindowTitle("Configuration Preview")
        preview_dialog.setModal(True)
        preview_dialog.resize(500, 400)
        
        layout = QVBoxLayout(preview_dialog)
        
        label = QLabel("Generated JSON Configuration:")
        label.setStyleSheet("font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(label)
        
        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setPlainText(json.dumps(config, indent=2))
        text_edit.setStyleSheet("font-family: 'Courier New', monospace; font-size: 9pt;")
        layout.addWidget(text_edit)
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(preview_dialog.accept)
        layout.addWidget(close_btn)
        
        preview_dialog.exec()
    
    def save_config(self):
        """Save the configuration to a JSON file"""
        import json
        from pathlib import Path
        
        # Validate required fields
        if not self.instructor_name.text().strip():
            QMessageBox.warning(self, "Validation Error", "Please enter an instructor name.")
            return
        
        config = self.build_config()
        
        # Ask where to save
        default_dir = Path("config")
        default_dir.mkdir(exist_ok=True)
        
        suggested_name = f"instructor_config_{self.instructor_name.text().strip().replace(' ', '_').lower()}.json"
        default_path = default_dir / suggested_name
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Instructor Configuration",
            str(default_path),
            "JSON Files (*.json);;All Files (*)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(config, f, indent=2, ensure_ascii=False)
                
                QMessageBox.information(
                    self, 
                    "Success", 
                    f"Instructor configuration saved successfully!\n\nFile: {Path(file_path).name}\n\n"
                    "You can now select this configuration in the main interface."
                )
                
                # Store the saved path for the parent to use
                self.saved_config_path = file_path
                self.accept()
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save configuration:\n{str(e)}")


if __name__ == "__main__":
    sys.exit(main())

